"""Optional LLM agent orchestration for contextual questionnaire filling."""

from __future__ import annotations

import asyncio
import csv
import datetime
import email.utils
import json
import random
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Any

import httpx
import numpy as np
import pdfplumber
from docx import Document as DocxDocument
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.config import settings
from app.models import QAPair
from app.services.parser import ExtractedItem
from app.utils.embeddings import bytes_to_embedding, compute_embeddings

AGENT_MODE_OFF = "off"
AGENT_MODE_AGENT = "agent"
AGENT_MODES = (AGENT_MODE_OFF, AGENT_MODE_AGENT)

# Runtime provider identifiers.
PROVIDER_OPENAI_COMPATIBLE = "openai-compatible"
PROVIDER_OPENAI = "openai"
PROVIDER_ANTHROPIC = "anthropic"

# Backward-compat aliases for old modes stored in DB or sent by API clients.
_MODE_ALIASES = {"assist": AGENT_MODE_AGENT, "full": AGENT_MODE_AGENT}
_PROVIDER_ALIASES = {
    "openai-compatible": PROVIDER_OPENAI_COMPATIBLE,
    "openai_compatible": PROVIDER_OPENAI_COMPATIBLE,
    "openai": PROVIDER_OPENAI,
    "anthropic": PROVIDER_ANTHROPIC,
    "claude": PROVIDER_ANTHROPIC,
}

AGENT_API_MAX_RETRIES = 5
AGENT_API_BASE_BACKOFF_SECONDS = 1.0
AGENT_API_MAX_BACKOFF_SECONDS = 30.0
AGENT_API_MAX_PARALLEL_CALLS = 8
_AGENT_API_CALL_SEMAPHORE = asyncio.Semaphore(AGENT_API_MAX_PARALLEL_CALLS)


@dataclass(frozen=True)
class AgentRuntimeConfig:
    """Provider/runtime values used for an individual agent run."""

    api_base: str
    api_key: str
    model: str
    provider: str = PROVIDER_OPENAI_COMPATIBLE


def _normalize_provider(provider: str | None) -> str:
    """Normalize provider names so UI aliases map to runtime behavior."""

    value = (provider or "").strip().lower()
    return _PROVIDER_ALIASES.get(value, value or PROVIDER_OPENAI_COMPATIBLE)


def default_agent_runtime_config() -> AgentRuntimeConfig:
    """Build runtime config from environment-backed application settings."""

    return AgentRuntimeConfig(
        api_base=settings.agent_api_base.strip(),
        api_key=settings.agent_api_key.strip(),
        model=settings.agent_model.strip(),
        provider=_normalize_provider(settings.agent_provider.strip()),
    )


def list_agent_modes() -> list[dict[str, str]]:
    """Return the UI-facing list of supported agent modes."""

    return [
        {
            "name": AGENT_MODE_OFF,
            "label": "Semantic Only",
            "description": "Use only knowledge-base semantic matching.",
        },
        {
            "name": AGENT_MODE_AGENT,
            "label": "Agent",
            "description": "AI-first mode. Agent decides all answers using document context + KB and flags uncertain fields (no semantic auto-match fallback).",
        },
    ]


def normalize_agent_mode(mode: str | None) -> str:
    """Normalize a requested agent mode and validate support."""

    normalized = (mode or settings.agent_default_mode or AGENT_MODE_OFF).strip().lower()
    normalized = _MODE_ALIASES.get(normalized, normalized)
    if normalized not in AGENT_MODES:
        raise ValueError(f"Unknown agent mode '{mode}'")
    return normalized


def is_agent_available(runtime_config: AgentRuntimeConfig | None = None) -> bool:
    """Return whether agent calls are fully configured and enabled."""

    config = runtime_config or default_agent_runtime_config()
    if runtime_config is None and not settings.agent_enabled:
        return False
    return bool(
        config.api_base.strip()
        and config.api_key.strip()
        and config.model.strip()
    )


def append_trace(
    trace: list[dict[str, Any]],
    step: str,
    status: str,
    message: str,
    data: dict[str, Any] | None = None,
) -> None:
    """Append a structured trace event for frontend rendering."""

    event: dict[str, Any] = {
        "timestamp": datetime.datetime.now(datetime.timezone.utc).replace(microsecond=0).isoformat(),
        "step": step,
        "status": status,
        "message": message,
    }
    if data:
        event["data"] = data
    trace.append(event)


def _chunked(items: list[Any], size: int) -> list[list[Any]]:
    """Split a list into fixed-size chunks."""

    if size <= 0:
        size = 1
    return [items[i:i + size] for i in range(0, len(items), size)]


def _estimate_tokens(text: str) -> int:
    """Rough token count estimate (~1 token per 4 chars for English text)."""
    return max(1, len(text) // 4)


def _extract_json_object(text: str) -> dict[str, Any]:
    """Parse LLM content into a JSON object, handling code fences."""

    cleaned = text.strip()
    if cleaned.startswith("```"):
        cleaned = re.sub(r"^```(?:json)?\s*", "", cleaned)
        cleaned = re.sub(r"\s*```$", "", cleaned)

    try:
        parsed = json.loads(cleaned)
        if isinstance(parsed, dict):
            return parsed
    except json.JSONDecodeError:
        pass

    start = cleaned.find("{")
    end = cleaned.rfind("}")
    if start >= 0 and end > start:
        try:
            parsed = json.loads(cleaned[start:end + 1])
            if isinstance(parsed, dict):
                return parsed
        except json.JSONDecodeError as exc:
            raise ValueError(
                f"Agent response contained braces but was not valid JSON: {exc}"
            ) from exc

    raise ValueError("Agent response was not valid JSON object text")


def _content_to_text(content: Any) -> str:
    """Normalize chat message content payloads to plain text."""

    if isinstance(content, str):
        return content
    if isinstance(content, list):
        parts: list[str] = []
        for item in content:
            if isinstance(item, dict):
                text = item.get("text")
                if isinstance(text, str) and text.strip():
                    parts.append(text.strip())
        return "\n".join(parts)
    return str(content)


def _response_json_or_none(response: httpx.Response) -> dict[str, Any] | None:
    """Best-effort parse of an HTTP response JSON body."""

    try:
        data = response.json()
    except ValueError:
        return None
    return data if isinstance(data, dict) else None


def _response_error_detail(response: httpx.Response) -> str:
    """Extract a human-friendly error message from provider responses."""

    payload = _response_json_or_none(response)
    if payload:
        error = payload.get("error")
        if isinstance(error, dict):
            message = str(error.get("message") or "").strip()
            code = str(error.get("code") or error.get("type") or "").strip()
            if message and code:
                return f"{message} ({code})"
            if message:
                return message
            if code:
                return code
        if isinstance(error, str) and error.strip():
            return error.strip()
        message = payload.get("message")
        if isinstance(message, str) and message.strip():
            return message.strip()

    text = (response.text or "").strip()
    if text:
        return text[:300]
    return f"HTTP {response.status_code}"


def _is_non_retriable_quota_error(response: httpx.Response) -> bool:
    """Detect provider errors where retrying immediately is unlikely to help."""

    payload = _response_json_or_none(response)
    if not payload:
        return False

    code_values: list[str] = []
    error = payload.get("error")
    if isinstance(error, dict):
        for key in ("code", "type", "message"):
            value = error.get(key)
            if isinstance(value, str) and value.strip():
                code_values.append(value.lower())
    elif isinstance(error, str) and error.strip():
        code_values.append(error.lower())

    message = payload.get("message")
    if isinstance(message, str) and message.strip():
        code_values.append(message.lower())

    joined = " ".join(code_values)
    return any(
        marker in joined
        for marker in (
            "insufficient_quota",
            "billing_hard_limit",
            "billing",
            "credit balance is too low",
            "quota exceeded",
        )
    )


def _is_retriable_response(response: httpx.Response) -> bool:
    """Return whether an HTTP response should be retried with backoff."""

    status = response.status_code
    if status == 429:
        return not _is_non_retriable_quota_error(response)
    return status in {408, 409} or 500 <= status <= 599


def _retry_delay_seconds(attempt: int, response: httpx.Response | None = None) -> float:
    """Compute retry delay using Retry-After header or exponential backoff."""

    if response is not None:
        retry_after = response.headers.get("retry-after")
        if retry_after:
            raw = retry_after.strip()
            try:
                seconds = float(raw)
                if seconds >= 0:
                    return min(seconds, AGENT_API_MAX_BACKOFF_SECONDS)
            except ValueError:
                try:
                    parsed_dt = email.utils.parsedate_to_datetime(raw)
                except (TypeError, ValueError):
                    parsed_dt = None
                if parsed_dt is not None:
                    now = datetime.datetime.now(datetime.timezone.utc)
                    delta = (parsed_dt - now).total_seconds()
                    if delta > 0:
                        return min(delta, AGENT_API_MAX_BACKOFF_SECONDS)

    base_delay = min(
        AGENT_API_MAX_BACKOFF_SECONDS,
        AGENT_API_BASE_BACKOFF_SECONDS * (2 ** max(attempt, 0)),
    )
    jitter = random.uniform(0, max(base_delay * 0.25, 0.05))
    return min(AGENT_API_MAX_BACKOFF_SECONDS, base_delay + jitter)


async def _post_json_with_retries(
    *,
    endpoint: str,
    headers: dict[str, str],
    payload: dict[str, Any],
    provider_label: str,
) -> dict[str, Any]:
    """POST JSON with retry/backoff for transient provider failures."""

    max_attempts = max(1, AGENT_API_MAX_RETRIES + 1)
    timeout = httpx.Timeout(settings.agent_timeout_seconds)

    async with httpx.AsyncClient(timeout=timeout) as client:
        for attempt in range(max_attempts):
            try:
                async with _AGENT_API_CALL_SEMAPHORE:
                    response = await client.post(endpoint, headers=headers, json=payload)
            except (httpx.TimeoutException, httpx.NetworkError, httpx.RemoteProtocolError) as exc:
                if attempt + 1 >= max_attempts:
                    raise RuntimeError(
                        f"{provider_label} request failed after {max_attempts} attempts: {exc}"
                    ) from exc
                await asyncio.sleep(_retry_delay_seconds(attempt))
                continue

            if response.status_code >= 400:
                if _is_retriable_response(response) and attempt + 1 < max_attempts:
                    await asyncio.sleep(_retry_delay_seconds(attempt, response))
                    continue
                raise RuntimeError(
                    f"{provider_label} API error {response.status_code}: {_response_error_detail(response)}"
                )

            data = _response_json_or_none(response)
            if data is None:
                raise ValueError(f"{provider_label} response was not valid JSON")
            return data

    raise RuntimeError(f"{provider_label} request failed after {max_attempts} attempts.")


async def _call_chat_json(
    messages: list[dict[str, str]],
    runtime_config: AgentRuntimeConfig,
    *,
    temperature: float = 0.1,
    token_usage: dict[str, int] | None = None,
) -> dict[str, Any]:
    """Call an LLM provider and parse a JSON object response."""

    # Estimate input tokens before call
    input_text = json.dumps(messages)
    input_est = _estimate_tokens(input_text)

    provider = _normalize_provider(runtime_config.provider)
    if provider == PROVIDER_ANTHROPIC:
        result = await _call_anthropic_json(messages, runtime_config, temperature=temperature)
    else:
        result = await _call_openai_compatible_json(messages, runtime_config, temperature=temperature)

    # Track token usage
    if token_usage is not None:
        output_est = _estimate_tokens(json.dumps(result))
        token_usage["input_tokens"] += input_est
        token_usage["output_tokens"] += output_est
        token_usage["llm_calls"] += 1

    return result


async def _call_openai_compatible_json(
    messages: list[dict[str, str]],
    runtime_config: AgentRuntimeConfig,
    *,
    temperature: float = 0.1,
) -> dict[str, Any]:
    """Call an OpenAI-compatible chat API and parse a JSON object response."""

    endpoint = runtime_config.api_base.rstrip("/") + "/chat/completions"
    headers = {
        "Authorization": f"Bearer {runtime_config.api_key}",
        "Content-Type": "application/json",
    }
    payload = {
        "model": runtime_config.model,
        "messages": messages,
        "temperature": temperature,
    }
    if settings.agent_structured_output:
        payload["response_format"] = {"type": "json_object"}

    data = await _post_json_with_retries(
        endpoint=endpoint,
        headers=headers,
        payload=payload,
        provider_label="OpenAI-compatible",
    )

    choices = data.get("choices") or []
    if not choices:
        raise ValueError("Agent response did not include choices")
    message = choices[0].get("message") or {}
    content = _content_to_text(message.get("content"))
    if not content.strip():
        raise ValueError("Agent response was empty")
    return _extract_json_object(content)


async def _call_anthropic_json(
    messages: list[dict[str, str]],
    runtime_config: AgentRuntimeConfig,
    *,
    temperature: float = 0.1,
) -> dict[str, Any]:
    """Call Anthropic Messages API and parse a JSON object response."""

    api_base = runtime_config.api_base.rstrip("/")
    endpoint = api_base if api_base.endswith("/messages") else f"{api_base}/messages"

    system_parts: list[str] = []
    converted_messages: list[dict[str, str]] = []
    for message in messages:
        role = str(message.get("role") or "user").strip().lower()
        content = str(message.get("content") or "").strip()
        if not content:
            continue
        if role == "system":
            system_parts.append(content)
            continue
        if role not in {"user", "assistant"}:
            role = "user"
        converted_messages.append({"role": role, "content": content})

    if not converted_messages:
        converted_messages = [{"role": "user", "content": "{}"}]

    payload: dict[str, Any] = {
        "model": runtime_config.model,
        "max_tokens": 4096,
        "temperature": temperature,
        "messages": converted_messages,
    }
    if system_parts:
        payload["system"] = "\n\n".join(system_parts)

    headers = {
        "x-api-key": runtime_config.api_key,
        "anthropic-version": "2023-06-01",
        "Content-Type": "application/json",
    }

    data = await _post_json_with_retries(
        endpoint=endpoint,
        headers=headers,
        payload=payload,
        provider_label="Anthropic",
    )

    content_blocks = data.get("content") or []
    text_parts: list[str] = []
    if isinstance(content_blocks, list):
        for block in content_blocks:
            if isinstance(block, dict) and block.get("type") == "text":
                text_value = block.get("text")
                if isinstance(text_value, str) and text_value.strip():
                    text_parts.append(text_value.strip())

    content = "\n".join(text_parts).strip()
    if not content:
        raise ValueError("Agent response was empty")
    return _extract_json_object(content)


def _extract_document_context(file_path: Path, items: list[ExtractedItem]) -> str:
    """Build a compact context string from source documents and extracted items."""

    lines: list[str] = []
    suffix = file_path.suffix.lower()

    try:
        if suffix == ".docx":
            doc = DocxDocument(str(file_path))
            for paragraph in doc.paragraphs[:120]:
                text = paragraph.text.strip()
                if text:
                    lines.append(text)
            for table in doc.tables[:10]:
                for row in table.rows[:40]:
                    row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_cells:
                        lines.append(" | ".join(row_cells))
        elif suffix == ".csv":
            with file_path.open("r", encoding="utf-8", errors="ignore", newline="") as handle:
                reader = csv.reader(handle)
                for idx, row in enumerate(reader):
                    if idx >= 120:
                        break
                    compact = [cell.strip() for cell in row if cell.strip()]
                    if compact:
                        lines.append(" | ".join(compact))
        elif suffix == ".pdf":
            with pdfplumber.open(str(file_path)) as pdf:
                for page in pdf.pages[:8]:
                    page_text = page.extract_text() or ""
                    for line in page_text.splitlines():
                        cleaned = line.strip()
                        if cleaned:
                            lines.append(cleaned)
    except Exception:
        # Fall back to extracted-question context below.
        pass

    if items:
        lines.append("Extracted question preview:")
        for item in items[:40]:
            question = (item.question_text or "").strip()
            if question:
                lines.append(f"- {question}")

    context = "\n".join(lines)
    max_chars = max(settings.agent_max_context_chars, 500)
    if len(context) > max_chars:
        context = context[:max_chars] + "\n...[truncated]"
    return context


async def _build_candidate_map(
    questions: list[str],
    db: AsyncSession,
    top_k: int = 3,
) -> dict[int, list[dict[str, Any]]]:
    """Build top-k KB candidates per question using existing embeddings."""

    if not questions:
        return {}

    result = await db.execute(select(QAPair).where(QAPair.embedding.isnot(None)).where(QAPair.deleted_at.is_(None)))
    qa_pairs = result.scalars().all()
    if not qa_pairs:
        return {idx: [] for idx in range(len(questions))}

    stored_embeddings = np.array([bytes_to_embedding(qa.embedding) for qa in qa_pairs])
    question_embeddings = compute_embeddings(questions)

    candidate_map: dict[int, list[dict[str, Any]]] = {}
    for idx, embedding in enumerate(question_embeddings):
        similarities = np.dot(stored_embeddings, embedding)
        if similarities.size == 0:
            candidate_map[idx] = []
            continue

        top_count = min(top_k, similarities.size)
        top_indices = np.argpartition(-similarities, top_count - 1)[:top_count]
        ordered = top_indices[np.argsort(-similarities[top_indices])]

        candidates: list[dict[str, Any]] = []
        cutoff = settings.agent_kb_similarity_cutoff
        for ranked_idx in ordered:
            sim = float(similarities[int(ranked_idx)])
            if sim < cutoff:
                continue  # Skip low-similarity noise to save tokens
            qa = qa_pairs[int(ranked_idx)]
            candidates.append(
                {
                    "id": qa.id,
                    "category": qa.category,
                    "question": qa.question,
                    "answer": qa.answer,
                    "similarity": round(sim, 4),
                }
            )
        candidate_map[idx] = candidates

    return candidate_map


def _coerce_confidence(value: Any) -> float | None:
    """Coerce arbitrary model confidence values to [0, 1]."""

    try:
        confidence = float(value)
    except (TypeError, ValueError):
        return None
    return max(0.0, min(1.0, confidence))


async def _summarize_document_context(
    raw_context: str,
    runtime_config: AgentRuntimeConfig,
    token_usage: dict[str, int] | None = None,
) -> str:
    """Use an LLM call to produce a structured document summary.

    Extracts key entities (company name, dates, contacts, certifications,
    document purpose) into a compact format that replaces raw truncation.
    Falls back to the raw context if the summarization call fails.
    """

    # Short documents don't need summarization
    if len(raw_context) <= settings.agent_max_context_chars:
        return raw_context
    # Very short documents: not worth the LLM call
    if len(raw_context) < 800:
        return raw_context

    try:
        messages = [
            {
                "role": "system",
                "content": (
                    "You are a document summarizer. Extract key structured information from the document text below. "
                    "Return a JSON object with these fields:\n"
                    "- document_type: what kind of document this is (questionnaire, RFP, compliance form, etc.)\n"
                    "- document_purpose: one-sentence summary of purpose\n"
                    "- requesting_entity: who is sending/requesting this questionnaire (if identifiable)\n"
                    "- target_entity: who should fill it in (if identifiable)\n"
                    "- key_dates: any dates mentioned\n"
                    "- contacts: any contact names/emails mentioned\n"
                    "- certifications_mentioned: any compliance standards/certs referenced (SOC2, ISO, GDPR, etc.)\n"
                    "- key_topics: list of 3-5 main topic areas covered\n"
                    "- notable_context: any other critical context a questionnaire filler would need\n"
                    "Be concise. Only include fields where you found information."
                ),
            },
            {
                "role": "user",
                "content": raw_context[:12000],
            },
        ]
        result = await _call_chat_json(messages, runtime_config, temperature=0.0, token_usage=token_usage)

        # Build compact summary string from structured result
        summary_parts: list[str] = ["=== DOCUMENT SUMMARY ==="]
        for key, value in result.items():
            if not value:
                continue
            label = key.replace("_", " ").title()
            if isinstance(value, list):
                summary_parts.append(f"{label}: {', '.join(str(v) for v in value)}")
            else:
                summary_parts.append(f"{label}: {value}")
        summary_parts.append("=== END SUMMARY ===")

        return "\n".join(summary_parts)
    except Exception:
        # Fall back to raw truncation
        max_chars = max(settings.agent_max_context_chars, 500)
        if len(raw_context) > max_chars:
            return raw_context[:max_chars] + "\n...[truncated]"
        return raw_context


async def _run_verification_stage(
    *,
    items: list[ExtractedItem],
    decisions: list[dict[str, Any]],
    document_summary: str,
    runtime_config: AgentRuntimeConfig,
    trace: list[dict[str, Any]],
    token_usage: dict[str, int] | None = None,
) -> list[dict[str, Any]]:
    """Verification stage: review low-confidence answers for consistency.

    Checks for internal contradictions, miscalibrated confidence, and
    answers that conflict with document context. Returns corrections.
    """

    if not settings.agent_verification_enabled:
        append_trace(trace, "verify", "skipped", "Verification stage disabled in settings.")
        return []

    # Collect answers to verify: all low-confidence + sample of high-confidence
    answers_to_verify: list[dict[str, Any]] = []
    flagged_questions: list[dict[str, Any]] = []

    for dec in decisions:
        if dec.get("action") == "answer":
            conf = dec.get("confidence") or 0.0
            answers_to_verify.append({
                "id": dec["id"],
                "question": dec.get("question", ""),
                "answer": dec.get("answer", ""),
                "confidence": conf,
                "reason": dec.get("reason", ""),
            })
        elif dec.get("action") == "flag":
            flagged_questions.append({
                "id": dec["id"],
                "question": dec.get("question", ""),
                "reason": dec.get("reason", ""),
            })

    if not answers_to_verify:
        append_trace(trace, "verify", "skipped", "No answers to verify.")
        return []

    # Cap to max verification questions
    max_q = settings.agent_verification_max_questions
    if len(answers_to_verify) > max_q:
        # Prioritize low-confidence answers
        answers_to_verify.sort(key=lambda x: x.get("confidence", 0.0))
        answers_to_verify = answers_to_verify[:max_q]

    append_trace(
        trace,
        "verify",
        "running",
        f"Verifying {len(answers_to_verify)} answer(s) for consistency.",
    )

    try:
        messages = [
            {
                "role": "system",
                "content": (
                    "You are VerifyAgent. Review the following answered questions for a questionnaire and check for:\n"
                    "1. INTERNAL CONSISTENCY: Are entity names (company, contact, dates) used consistently across answers?\n"
                    "2. DOCUMENT CONTRADICTIONS: Does any answer contradict information in the document summary?\n"
                    "3. CONFIDENCE CALIBRATION: Are any confidence scores miscalibrated? (e.g., a generic/uncertain answer marked 0.95, or a well-supported answer marked low)\n"
                    "4. FACTUAL CONCERNS: Any answers that appear fabricated or unsupported?\n\n"
                    "For each issue found, return a correction. Only return corrections for real problems.\n"
                    "Return strict JSON: {\"corrections\":[{\"id\":\"q_X\",\"action\":\"revise|flag|accept\","
                    "\"revised_answer\":\"...\",\"revised_confidence\":0.X,\"reason\":\"...\"}],\"verification_summary\":\"...\"}"
                ),
            },
            {
                "role": "user",
                "content": json.dumps(
                    {
                        "document_summary": document_summary,
                        "answers": answers_to_verify,
                        "flagged_questions": flagged_questions[:10],
                    },
                    ensure_ascii=False,
                ),
            },
        ]
        result = await _call_chat_json(messages, runtime_config, temperature=0.0, token_usage=token_usage)
        corrections = result.get("corrections", [])
        if not isinstance(corrections, list):
            corrections = []

        verification_summary = str(result.get("verification_summary", "")).strip()
        append_trace(
            trace,
            "verify",
            "completed",
            verification_summary or f"Verification found {len(corrections)} correction(s).",
            {"corrections_count": len(corrections)},
        )
        return corrections
    except Exception as exc:
        append_trace(
            trace,
            "verify",
            "error",
            f"Verification stage failed: {exc}. Proceeding with unverified answers.",
        )
        return []


async def run_contextual_fill_agent(
    *,
    file_path: Path,
    items: list[ExtractedItem],
    db: AsyncSession,
    mode: str,
    instructions: str | None = None,
    runtime_config: AgentRuntimeConfig | None = None,
) -> dict[str, Any]:
    """Run two-stage research + fill agent over extracted questions."""

    mode = normalize_agent_mode(mode)
    runtime = runtime_config or default_agent_runtime_config()
    trace: list[dict[str, Any]] = []

    if mode == AGENT_MODE_OFF:
        append_trace(trace, "agent", "skipped", "Agent mode is disabled for this job.")
        return {
            "items": items,
            "trace": trace,
            "status": "skipped",
            "summary": "Agent mode was off.",
            "stats": {
                "questions_considered": 0,
                "answers_generated": 0,
                "flags_recommended": 0,
                "overrides": 0,
            },
            "decisions": [],
        }

    if not is_agent_available(runtime):
        append_trace(
            trace,
            "agent",
            "skipped",
            "Agent was requested but API settings are incomplete.",
        )
        return {
            "items": items,
            "trace": trace,
            "status": "skipped",
            "summary": "Agent requested but not configured (missing key/model/base URL or disabled).",
            "stats": {
                "questions_considered": 0,
                "answers_generated": 0,
                "flags_recommended": 0,
                "overrides": 0,
            },
            "decisions": [],
        }

    original_answers = [item.answer_text for item in items]

    # ── Deduplication: group identical questions, process each unique question once ──
    seen_questions: dict[str, int] = {}  # normalized_question -> first work_items index
    duplicate_map: dict[int, int] = {}  # item_index -> canonical item_index
    work_items: list[dict[str, Any]] = []

    for idx, item in enumerate(items):
        norm_q = (item.question_text or "").strip().lower()
        if norm_q in seen_questions:
            # This is a duplicate — will be filled from canonical answer later
            duplicate_map[idx] = seen_questions[norm_q]
        else:
            seen_questions[norm_q] = idx
            work_items.append({
                "index": idx,
                "question_text": item.question_text,
                "current_answer": item.answer_text,
            })

    if duplicate_map:
        append_trace(
            trace,
            "agent",
            "running",
            f"Deduplicated {len(duplicate_map)} repeated question(s) — will process {len(work_items)} unique question(s).",
        )

    for work_idx, work_item in enumerate(work_items):
        work_item["work_idx"] = work_idx

    if not work_items:
        append_trace(trace, "agent", "skipped", "No questions needed agent work for this mode.")
        return {
            "items": items,
            "trace": trace,
            "status": "completed",
            "summary": "No unresolved questions needed agent reasoning.",
            "stats": {
                "questions_considered": 0,
                "answers_generated": 0,
                "flags_recommended": 0,
                "overrides": 0,
            },
            "decisions": [],
        }

    append_trace(
        trace,
        "agent",
        "running",
        "Starting contextual research + fill workflow.",
        {
            "mode": mode,
            "provider": runtime.provider,
            "model": runtime.model,
            "questions_considered": len(work_items),
        },
    )

    token_usage: dict[str, int] = {"input_tokens": 0, "output_tokens": 0, "llm_calls": 0}

    raw_context = _extract_document_context(file_path, items)
    document_context = await _summarize_document_context(raw_context, runtime, token_usage=token_usage)
    candidate_map = await _build_candidate_map(
        [work_item["question_text"] for work_item in work_items],
        db,
        top_k=3,
    )
    append_trace(
        trace,
        "research",
        "running",
        "Built knowledge-base candidate context.",
        {"question_count": len(work_items)},
    )

    # ── Confidence-based KB routing: skip LLM for high-confidence KB matches ──
    kb_direct_threshold = settings.agent_kb_direct_threshold
    kb_routed_count = 0
    llm_work_items: list[dict[str, Any]] = []

    ordered_decisions: list[dict[str, Any]] = []
    prior_answers: list[dict[str, str]] = []  # Accumulated high-confidence answers for cross-chunk context

    for work_item in work_items:
        local_idx = int(work_item["work_idx"])
        candidates = candidate_map.get(local_idx, [])
        top_sim = candidates[0]["similarity"] if candidates else 0.0

        if candidates and top_sim >= kb_direct_threshold:
            # High-confidence KB match — use directly, skip LLM
            item_index = work_item["index"]
            kb_answer = candidates[0]["answer"]
            items[item_index].answer_text = kb_answer
            items[item_index].matched_source = "kb_direct"
            items[item_index].matched_qa_id = candidates[0].get("id")
            items[item_index].confidence = top_sim

            decision_payload = {
                "id": f"q_{item_index}",
                "question": work_item["question_text"],
                "action": "answer",
                "answer": kb_answer,
                "confidence": round(top_sim, 3),
                "reason": f"Direct KB match (similarity {top_sim:.2f})",
                "issues": [],
            }
            ordered_decisions.append(decision_payload)
            prior_answers.append({
                "question": work_item["question_text"],
                "answer": kb_answer,
            })
            kb_routed_count += 1
        else:
            llm_work_items.append(work_item)

    if kb_routed_count > 0:
        append_trace(
            trace,
            "kb_routing",
            "completed",
            f"KB-routed {kb_routed_count} question(s) with similarity >= {kb_direct_threshold} (skipped LLM).",
            {"kb_routed": kb_routed_count, "remaining_for_llm": len(llm_work_items)},
        )

    # ── Early exit: all questions resolved by KB, skip LLM entirely ──
    if not llm_work_items:
        append_trace(
            trace, "fill", "completed",
            f"All {kb_routed_count} question(s) resolved via KB direct match — LLM skipped entirely.",
        )
        return {
            "items": items,
            "decisions": ordered_decisions,
            "trace": trace,
            "status": "completed",
            "summary": f"All {kb_routed_count} question(s) answered from knowledge base (no LLM needed).",
            "stats": {
                "input_tokens": token_usage["input_tokens"],
                "output_tokens": token_usage["output_tokens"],
                "llm_calls": token_usage["llm_calls"],
                "kb_routed": kb_routed_count,
            },
        }

    chunks = _chunked(llm_work_items, settings.agent_max_questions_per_call)
    failed_chunks = 0
    use_single_stage = settings.agent_single_stage

    async def _process_chunk(chunk_idx: int, chunk: list[dict[str, Any]], prior_answers_snapshot: list[dict[str, str]]) -> list[dict[str, Any]]:
        """Process a single LLM chunk — designed to run concurrently.

        Uses a snapshot of prior_answers taken before concurrent execution to
        avoid reading shared mutable state from other chunks.
        """
        payload_questions = []
        for work_item in chunk:
            global_idx = work_item["index"]
            local_idx = int(work_item["work_idx"])
            payload_questions.append(
                {
                    "id": f"q_{global_idx}",
                    "question": work_item["question_text"],
                    "current_answer": work_item["current_answer"],
                    "kb_candidates": candidate_map.get(local_idx, []),
                }
            )

        if use_single_stage:
            # ── SINGLE-STAGE MODE: one call does research + fill (saves ~50% tokens) ──
            # Build compact prior_answers summary (just Q+A, no bloat)
            prior_context = ""
            if prior_answers_snapshot:
                pa_lines = [f"- {pa['question']}: {pa['answer']}" for pa in prior_answers_snapshot[-settings.agent_max_prior_answers:]]
                prior_context = "\n\nPRIOR ANSWERS (use for consistency):\n" + "\n".join(pa_lines)

            combined_messages = [
                {
                    "role": "system",
                    "content": (
                        "You are an AI questionnaire filling agent. Analyze each question, research it using "
                        "the knowledge base candidates and document context, then decide whether to answer or flag it.\n\n"
                        "RULES:\n"
                        "1. ACTION='answer' when confidence >= 0.7 AND answer is supported by KB candidates or document context.\n"
                        "2. ACTION='flag' when: confidence < 0.7, question asks for context-specific info (names, dates, "
                        "contacts, project references), KB has no match, or any ambiguity exists.\n"
                        "3. Context-specific fields (Company Name, Contact, Date, Address) → ALWAYS flag unless prior answers "
                        "or document headers make the answer unambiguous.\n"
                        "4. Never fabricate. A wrong answer is worse than a flag.\n"
                        "5. CONFIDENCE: 0.9-1.0 = near-certain (direct KB/doc match), 0.7-0.85 = good inference, "
                        "<0.7 = uncertain (flag it).\n"
                        "6. For yes/no compliance questions with strong KB match → answer confidently.\n"
                        "7. For contact_info/date questions → almost always flag.\n"
                        "8. Include 'reason' explaining your decision. Include 'issues' array for any concerns.\n\n"
                        "Return strict JSON:\n"
                        "{\"decisions\":[{\"id\":\"q_X\",\"action\":\"answer|flag\",\"answer\":\"...\","
                        "\"confidence\":0.0,\"reason\":\"...\",\"issues\":[]}],\"summary\":\"...\"}"
                    ),
                },
                {
                    "role": "user",
                    "content": json.dumps(
                        {
                            "document_context": document_context,
                            "instructions": instructions or "",
                            "questions": payload_questions,
                        },
                        ensure_ascii=False,
                    ) + prior_context,
                },
            ]
            fill_result = await _call_chat_json(combined_messages, runtime, temperature=0.0, token_usage=token_usage)
            chunk_decisions = fill_result.get("decisions", [])
            if not isinstance(chunk_decisions, list):
                chunk_decisions = []
        else:
            # ── TWO-STAGE MODE: separate research + fill calls (original behavior) ──
            research_messages = [
                {
                    "role": "system",
                    "content": (
                        "You are ResearchAgent for questionnaire autofill. Your job is to think like a person "
                        "sitting down to fill in a questionnaire on behalf of their company.\n\n"
                        "CRITICAL CONTEXT RULES:\n"
                        "1. The DOCUMENT CONTEXT shows the questionnaire being filled. Read headers, titles, and "
                        "surrounding text to understand WHO is being asked (e.g., which company or customer).\n"
                        "2. The KB CANDIDATES contain answers from YOUR company's knowledge base — these are "
                        "answers about YOUR company's policies, practices, and information.\n"
                        "3. For GENERAL questions about your company (security policies, compliance, certifications, "
                        "processes), use KB candidates confidently when they match.\n"
                        "4. For CONTEXT-SPECIFIC questions (Company Name, Contact Person, Date, Project Name, "
                        "Client Reference, specific third-party names), these change per questionnaire. Flag these "
                        "for human review unless the answer is clearly derivable from the document context itself.\n"
                        "5. Do NOT invent facts. Do NOT guess contact details, dates, or entity-specific information.\n"
                        "6. When unsure, set needs_human_review=true and confidence below 0.5.\n\n"
                        "QUESTION TYPE CLASSIFICATION — for each question, classify its type:\n"
                        "- factual_about_company: questions about your company's policies, practices, capabilities\n"
                        "- context_specific: names, dates, contacts, project references that vary per questionnaire\n"
                        "- yes_no_compliance: binary compliance questions (Do you...? Are you...?)\n"
                        "- open_ended: narrative/descriptive questions needing detailed explanation\n"
                        "- contact_info: requesting specific contact details, addresses, phone numbers\n"
                        "- date_field: requesting specific dates or timelines\n\n"
                        "Return strict JSON: "
                        "{\"notes\":[{\"id\":\"...\",\"question_type\":\"...\",\"research_summary\":\"...\","
                        "\"proposed_answer\":\"...\",\"confidence\":0.0,\"needs_human_review\":false,"
                        "\"issues\":[\"...\"]}]}"
                    ),
                },
                {
                    "role": "user",
                    "content": json.dumps(
                        {
                            "mode": mode,
                            "instructions": instructions or "",
                            "document_context": document_context,
                            "questions": payload_questions,
                        },
                        ensure_ascii=False,
                    ),
                },
            ]
            research_result = await _call_chat_json(research_messages, runtime, temperature=0.0, token_usage=token_usage)
            research_notes = research_result.get("notes", [])
            if not isinstance(research_notes, list):
                research_notes = []

            fill_messages = [
                {
                    "role": "system",
                    "content": (
                        "You are FillAgent for questionnaire completion. Think like a professional filling in "
                        "a security or compliance questionnaire on behalf of their company.\n\n"
                        "DECISION RULES:\n"
                        "1. ACTION='answer' only when confidence >= 0.7 AND the answer is factually supported "
                        "by KB candidates or clear document context.\n"
                        "2. ACTION='flag' when: confidence < 0.7, the question is context-specific (names, "
                        "dates, contacts, project references), the KB has no relevant match, or there is any "
                        "ambiguity about which entity the question refers to.\n"
                        "3. For context-specific fields (Company Name, Contact Person, Date, Address, etc.), "
                        "ALWAYS flag unless a prior answer or document header makes the answer unambiguous.\n"
                        "4. When flagging, include a clear 'reason' explaining what information is needed so "
                        "the human reviewer can fill it quickly.\n"
                        "5. Prefer KB-backed answers for general policy/compliance/security questions.\n"
                        "6. Never fabricate information. A wrong answer is worse than a flag.\n\n"
                        "CONFIDENCE CALIBRATION:\n"
                        "- 0.9-1.0: Answer is directly backed by a KB entry with high similarity OR explicitly "
                        "stated in the document. Reserve this for near-certain answers only.\n"
                        "- 0.7-0.85: Answer is inferred from KB candidates or document context with reasonable "
                        "confidence. Good KB match but not exact.\n"
                        "- 0.5-0.7: Partial match, some uncertainty. Should be flagged.\n"
                        "- Below 0.5: Significant uncertainty or no supporting evidence.\n\n"
                        "CROSS-QUESTION CONSISTENCY:\n"
                        "- If prior_answers are provided, use them for consistency. If Q1 identified a company "
                        "name, reuse it where applicable. Do not contradict prior answers.\n\n"
                        "QUESTION TYPE AWARENESS:\n"
                        "- Use the question_type from research notes to inform your decision.\n"
                        "- 'contact_info' and 'date_field' types should almost always be flagged.\n"
                        "- 'yes_no_compliance' with strong KB match can be answered confidently.\n"
                        "- 'factual_about_company' with KB match can be answered confidently.\n\n"
                        "Return strict JSON: "
                        "{\"decisions\":[{\"id\":\"...\",\"action\":\"answer|flag\",\"answer\":\"...\","
                        "\"confidence\":0.0,\"reason\":\"...\",\"issues\":[\"...\"]}],\"summary\":\"...\"}"
                    ),
                },
                {
                    "role": "user",
                    "content": json.dumps(
                        {
                            "mode": mode,
                            "instructions": instructions or "",
                            "questions": payload_questions,
                            "research_notes": research_notes,
                            "prior_answers": prior_answers_snapshot[-settings.agent_max_prior_answers:],
                        },
                        ensure_ascii=False,
                    ),
                },
            ]
            fill_result = await _call_chat_json(fill_messages, runtime, temperature=0.0, token_usage=token_usage)
            chunk_decisions = fill_result.get("decisions", [])
            if not isinstance(chunk_decisions, list):
                chunk_decisions = []

        return chunk_decisions

    # ── Run all chunks concurrently for maximum speed ──
    append_trace(
        trace,
        "fill",
        "running",
        f"Processing {len(chunks)} chunk(s) concurrently ({len(llm_work_items)} questions total).",
        {"chunks": len(chunks), "questions": len(llm_work_items)},
    )

    # Snapshot prior_answers before concurrent execution so each chunk sees
    # KB-routed answers but not other chunks' answers (avoids race condition).
    prior_answers_snapshot = list(prior_answers)

    async def _safe_process_chunk(chunk_idx: int, chunk: list[dict[str, Any]]) -> tuple[int, list[dict[str, Any]]]:
        """Wrapper that catches exceptions per-chunk so one failure doesn't kill the batch."""
        try:
            decisions = await _process_chunk(chunk_idx, chunk, prior_answers_snapshot)
            return chunk_idx, decisions
        except Exception as chunk_exc:
            append_trace(
                trace,
                "fill",
                "error",
                f"Chunk {chunk_idx}/{len(chunks)} failed; leaving question(s) flagged.",
                {"error": str(chunk_exc)},
            )
            return chunk_idx, []

    chunk_results = await asyncio.gather(
        *[_safe_process_chunk(idx, chunk) for idx, chunk in enumerate(chunks, start=1)]
    )

    for chunk_idx, chunk_decisions in chunk_results:
        if not chunk_decisions:
            failed_chunks += 1
            continue

        for decision in chunk_decisions:
            if not isinstance(decision, dict):
                continue
            decision_id = str(decision.get("id", ""))
            if not decision_id.startswith("q_"):
                continue
            try:
                item_index = int(decision_id.split("_", maxsplit=1)[1])
            except (TypeError, ValueError):
                continue
            if item_index < 0 or item_index >= len(items):
                continue

            action = str(decision.get("action", "flag")).strip().lower()
            answer = (decision.get("answer") or "").strip()
            if action not in {"answer", "flag"}:
                action = "flag"
            if action == "answer" and not answer:
                action = "flag"

            decision_payload = {
                "id": decision_id,
                "question": items[item_index].question_text,
                "action": action,
                "answer": answer if action == "answer" else None,
                "confidence": _coerce_confidence(decision.get("confidence")),
                "reason": str(decision.get("reason") or ""),
                "issues": decision.get("issues") if isinstance(decision.get("issues"), list) else [],
            }
            ordered_decisions.append(decision_payload)

            if action == "answer":
                items[item_index].answer_text = answer
                prior_answers.append({
                    "question": items[item_index].question_text,
                    "answer": answer,
                })
            elif mode == AGENT_MODE_AGENT or not items[item_index].answer_text:
                items[item_index].answer_text = None

        append_trace(
            trace,
            "fill",
            "completed",
            f"Completed chunk {chunk_idx}/{len(chunks)}.",
            {"decisions": len(chunk_decisions)},
        )

    # ── Verification stage (conditional) ──
    if ordered_decisions and settings.agent_verification_enabled:
        # Smart skip: if average confidence is very high, skip verification to save tokens
        answered_confs = [
            d.get("confidence", 0.0) for d in ordered_decisions
            if d.get("action") == "answer" and d.get("confidence") is not None
        ]
        avg_conf = sum(answered_confs) / len(answered_confs) if answered_confs else 0.0

        if avg_conf >= settings.agent_skip_verify_threshold and len(answered_confs) > 3:
            append_trace(
                trace,
                "verify",
                "skipped",
                f"Skipped verification — average confidence {avg_conf:.2f} exceeds threshold "
                f"{settings.agent_skip_verify_threshold}.",
            )
        else:
            corrections = await _run_verification_stage(
                items=items,
                decisions=ordered_decisions,
                document_summary=document_context,
                runtime_config=runtime,
                trace=trace,
                token_usage=token_usage,
            )
            # Apply corrections
            for correction in corrections:
                if not isinstance(correction, dict):
                    continue
                cid = str(correction.get("id", ""))
                if not cid.startswith("q_"):
                    continue
                try:
                    cidx = int(cid.split("_", maxsplit=1)[1])
                except (TypeError, ValueError):
                    continue
                if cidx < 0 or cidx >= len(items):
                    continue

                c_action = str(correction.get("action", "")).strip().lower()
                if c_action == "revise":
                    revised = (correction.get("revised_answer") or "").strip()
                    if revised:
                        items[cidx].answer_text = revised
                        # Update the decision record too
                        for dec in ordered_decisions:
                            if dec.get("id") == cid:
                                dec["answer"] = revised
                                if correction.get("revised_confidence") is not None:
                                    dec["confidence"] = _coerce_confidence(correction["revised_confidence"])
                                dec["reason"] = (dec.get("reason", "") + " [Revised by verification: " +
                                                 str(correction.get("reason", "")) + "]").strip()
                                break
                elif c_action == "flag":
                    items[cidx].answer_text = None
                    for dec in ordered_decisions:
                        if dec.get("id") == cid:
                            dec["action"] = "flag"
                            dec["answer"] = None
                            dec["reason"] = (dec.get("reason", "") + " [Flagged by verification: " +
                                             str(correction.get("reason", "")) + "]").strip()
                            break

    # ── Propagate answers to duplicate questions (zero extra LLM cost) ──
    for dup_idx, canonical_idx in duplicate_map.items():
        items[dup_idx].answer_text = items[canonical_idx].answer_text

    answers_generated = 0
    overrides = 0
    flags_recommended = 0
    unresolved_questions = 0

    for idx, item in enumerate(items):
        before = original_answers[idx]
        after = item.answer_text
        if after is None:
            unresolved_questions += 1
        if before is None and after is not None:
            answers_generated += 1
        if before is not None and after is None:
            flags_recommended += 1
        if before is not None and after is not None and before.strip() != after.strip():
            overrides += 1

    summary = f"Agent processed {len(work_items)} questions, generated {answers_generated} answer(s), and left {unresolved_questions} question(s) flagged for review."
    if overrides > 0:
        summary += f" Adjusted {overrides} existing answer(s)."
    if flags_recommended > 0:
        summary += f" Added {flags_recommended} new flag(s) over prior answers."
    if failed_chunks > 0:
        summary += f" {failed_chunks} chunk(s) hit transient API issues."
    final_status = "completed" if failed_chunks == 0 else "completed_with_warnings"
    append_trace(
        trace,
        "agent",
        final_status,
        summary,
    )

    return {
        "items": items,
        "trace": trace,
        "status": final_status,
        "summary": summary,
        "stats": {
            "questions_considered": len(work_items),
            "answers_generated": answers_generated,
            "unresolved_questions": unresolved_questions,
            "flags_recommended": flags_recommended,
            "overrides": overrides,
            "failed_chunks": failed_chunks,
            "kb_routed": kb_routed_count,
            "input_tokens": token_usage["input_tokens"],
            "output_tokens": token_usage["output_tokens"],
            "llm_calls": token_usage["llm_calls"],
        },
        "decisions": ordered_decisions[:100],
    }


async def run_troubleshoot_agent(
    *,
    file_path: Path,
    profile_results: list[dict[str, Any]],
    recommended_profile: str | None,
    instructions: str | None = None,
    runtime_config: AgentRuntimeConfig | None = None,
) -> dict[str, Any]:
    """Run an optional troubleshooting-focused agent analysis."""

    runtime = runtime_config or default_agent_runtime_config()
    trace: list[dict[str, Any]] = []
    profile_labels = {
        str(profile.get("profile_name")): str(profile.get("profile_label") or profile.get("profile_name") or "")
        for profile in profile_results
        if profile.get("profile_name")
    }
    valid_profiles = set(profile_labels.keys())
    rule_recommended = (
        str(recommended_profile).strip()
        if recommended_profile and str(recommended_profile).strip() in valid_profiles
        else None
    )

    if not is_agent_available(runtime):
        append_trace(
            trace,
            "agent",
            "skipped",
            "Troubleshooting agent skipped because API settings are incomplete.",
        )
        return {
            "status": "skipped",
            "summary": "Agent diagnostics unavailable because the API key/model/base URL are not configured.",
            "trace": trace,
            "root_causes": [],
            "next_steps": [],
            "recommended_profile": rule_recommended,
            "fix_plan": {
                "type": "configuration",
                "title": "Configure AI provider to enable model troubleshooting",
                "action": "manual_follow_up",
                "can_auto_apply": False,
                "parser_profile": rule_recommended,
                "parser_profile_label": profile_labels.get(rule_recommended) if rule_recommended else None,
                "parser_hints": {},
                "steps": [
                    "Open Settings and add AI provider credentials.",
                    "Run Troubleshooting again with AI diagnostics enabled.",
                ],
            },
        }

    context = _extract_document_context(file_path, [])
    append_trace(
        trace,
        "agent",
        "running",
        "Running agent troubleshooting analysis.",
        {"provider": runtime.provider, "model": runtime.model},
    )

    payload = {
        "instructions": instructions or "",
        "recommended_profile_from_rules": recommended_profile,
        "profiles": profile_results,
        "document_context": context,
    }
    messages = [
        {
            "role": "system",
            "content": (
                "You are TroubleshootAgent for questionnaire parsing diagnostics. "
                "Analyze parser-profile outcomes and explain likely root causes, then suggest the safest system fix. "
                "Return strict JSON object: "
                "{\"summary\":\"...\",\"root_causes\":[\"...\"],\"next_steps\":[\"...\"],"
                "\"recommended_profile\":\"profile_or_null\",\"fix_type\":\"switch_profile|ocr|layout|parser_gap|none\","
                "\"fix_rationale\":\"...\","
                "\"parser_hints\":{\"question_column_index\":0,\"answer_column_index\":1,\"header_rows\":1,\"detect_row_blocks\":true}}"
            ),
        },
        {"role": "user", "content": json.dumps(payload, ensure_ascii=False)},
    ]

    result = await _call_chat_json(messages, runtime, temperature=0.0)
    summary = str(result.get("summary") or "").strip()
    root_causes = result.get("root_causes")
    next_steps = result.get("next_steps")
    recommended = result.get("recommended_profile")
    fix_type = str(result.get("fix_type") or "").strip().lower()
    fix_rationale = str(result.get("fix_rationale") or "").strip()
    parser_hints_raw = result.get("parser_hints")

    if not isinstance(root_causes, list):
        root_causes = []
    if not isinstance(next_steps, list):
        next_steps = []
    if recommended is not None:
        recommended = str(recommended).strip() or None
    if recommended and recommended not in valid_profiles:
        recommended = None
    if not recommended and rule_recommended:
        recommended = rule_recommended

    if not summary:
        summary = "Agent analysis completed."

    parser_hints: dict[str, Any] = {}
    if isinstance(parser_hints_raw, dict):
        integer_keys = {"question_column_index", "answer_column_index", "header_rows"}
        boolean_keys = {"detect_row_blocks"}
        for key in integer_keys:
            value = parser_hints_raw.get(key)
            if isinstance(value, bool):
                continue
            try:
                parsed = int(value)
            except (TypeError, ValueError):
                continue
            if parsed >= 0:
                parser_hints[key] = parsed
        for key in boolean_keys:
            value = parser_hints_raw.get(key)
            if isinstance(value, bool):
                parser_hints[key] = value

    can_auto_apply = bool(recommended)
    if can_auto_apply:
        profile_label = profile_labels.get(recommended, recommended)
        fix_title = f"Switch to '{profile_label}' parser profile for this layout"
        fix_action = "set_default_parser_profile"
        fix_steps = [
            f"Apply '{recommended}' as the default parser profile.",
            "Re-upload the problematic document and verify extracted question preview.",
            "Keep agent mode enabled so uncertain answers remain flagged.",
        ]
        if not fix_type:
            fix_type = "switch_profile"
    else:
        profile_label = None
        fix_title = "Document needs manual parsing remediation"
        fix_action = "manual_follow_up"
        fix_steps = [str(step).strip() for step in next_steps if str(step).strip()][:3]
        if not fix_steps:
            if file_path.suffix.lower() == ".pdf":
                fix_steps = [
                    "Ensure the PDF has selectable text (run OCR if scanned).",
                    "Retry troubleshooting after OCR/export.",
                ]
                if not fix_type:
                    fix_type = "ocr"
            else:
                fix_steps = [
                    "Retry with a clean export to DOCX/CSV.",
                    "If the layout is custom, add a dedicated parser profile.",
                ]
                if not fix_type:
                    fix_type = "layout"

    append_trace(
        trace,
        "agent",
        "completed",
        summary,
        data={"recommended_profile": recommended, "auto_fix": can_auto_apply},
    )
    return {
        "status": "completed",
        "summary": summary,
        "trace": trace,
        "root_causes": [str(item) for item in root_causes[:8]],
        "next_steps": [str(item) for item in next_steps[:8]],
        "recommended_profile": recommended,
        "fix_plan": {
            "type": fix_type or "none",
            "title": fix_title,
            "rationale": fix_rationale or summary,
            "action": fix_action,
            "can_auto_apply": can_auto_apply,
            "parser_profile": recommended,
            "parser_profile_label": profile_label,
            "parser_hints": parser_hints,
            "steps": fix_steps,
        },
    }
