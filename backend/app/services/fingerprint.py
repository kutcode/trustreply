"""Format fingerprinting service - remembers questionnaire layouts."""

from __future__ import annotations
import hashlib
import datetime
import logging
from pathlib import Path
from typing import Any

from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.models import FormatFingerprint

logger = logging.getLogger(__name__)


def _extract_structure(file_path: Path) -> dict[str, Any]:
    """Extract structural metadata from a document for fingerprinting."""

    suffix = file_path.suffix.lower()
    metadata: dict[str, Any] = {"file_type": suffix.lstrip(".")}

    if suffix == ".csv":
        import csv

        with open(file_path, "r", encoding="utf-8-sig") as f:
            reader = csv.reader(f)
            header_row = next(reader, None)
            if header_row:
                metadata["column_count"] = len(header_row)
                metadata["header_words"] = sorted(
                    set(w.strip().lower() for cell in header_row for w in cell.split() if w.strip())
                )
            else:
                metadata["column_count"] = 0
                metadata["header_words"] = []

    elif suffix == ".docx":
        from docx import Document as DocxDocument

        doc = DocxDocument(str(file_path))
        table_cols = []
        header_words = set()
        for table in doc.tables:
            if table.rows:
                cols = len(table.rows[0].cells) if table.rows[0].cells else 0
                table_cols.append(cols)
                for cell in table.rows[0].cells:
                    for w in cell.text.split():
                        stripped = w.strip().lower()
                        if stripped:
                            header_words.add(stripped)

        metadata["column_count"] = max(table_cols) if table_cols else 0
        metadata["table_count"] = len(doc.tables)
        metadata["header_words"] = sorted(header_words)

    elif suffix == ".pdf":
        metadata["column_count"] = 0
        metadata["header_words"] = []

    return metadata


def compute_fingerprint_hash(metadata: dict[str, Any]) -> str:
    """Compute a SHA-256 hash from structural metadata."""

    components = [
        str(metadata.get("file_type", "")),
        str(metadata.get("column_count", 0)),
        "|".join(metadata.get("header_words", [])),
    ]
    raw = "::".join(components)
    return hashlib.sha256(raw.encode()).hexdigest()


def compute_fingerprint(file_path: Path) -> tuple[str, dict[str, Any]]:
    """Extract structure and compute fingerprint hash.

    Returns (fingerprint_hash, structural_metadata).
    """

    metadata = _extract_structure(file_path)
    fp_hash = compute_fingerprint_hash(metadata)
    return fp_hash, metadata


async def find_matching_fingerprint(
    file_path: Path,
    db: AsyncSession,
) -> FormatFingerprint | None:
    """Check if a matching fingerprint exists for this file."""

    try:
        fp_hash, _ = compute_fingerprint(file_path)
    except Exception:
        logger.warning("Fingerprint lookup skipped; could not hash %s", file_path.name, exc_info=True)
        return None

    result = await db.execute(
        select(FormatFingerprint).where(FormatFingerprint.fingerprint_hash == fp_hash)
    )
    return result.scalars().first()


async def save_fingerprint(
    *,
    file_path: Path,
    parser_profile: str,
    hint_overrides: dict | None,
    original_filename: str,
    db: AsyncSession,
) -> FormatFingerprint | None:
    """Save or update a fingerprint after a successful job."""

    try:
        fp_hash, metadata = compute_fingerprint(file_path)
    except Exception:
        logger.warning("Fingerprint save skipped; could not hash %s", file_path.name, exc_info=True)
        return None

    result = await db.execute(
        select(FormatFingerprint).where(FormatFingerprint.fingerprint_hash == fp_hash)
    )
    existing = result.scalars().first()

    if existing:
        existing.success_count = (existing.success_count or 0) + 1
        existing.last_used_at = datetime.datetime.now(datetime.timezone.utc).replace(tzinfo=None)
        existing.parser_profile = parser_profile
        if hint_overrides:
            existing.hint_overrides = hint_overrides
        return existing

    header_words = metadata.get("header_words", [])
    fp = FormatFingerprint(
        fingerprint_hash=fp_hash,
        name=None,
        source_filename=original_filename,
        column_count=metadata.get("column_count"),
        header_signature=" | ".join(header_words[:20]) if header_words else None,
        structural_metadata=metadata,
        parser_profile=parser_profile,
        hint_overrides=hint_overrides,
        success_count=1,
        last_used_at=datetime.datetime.now(datetime.timezone.utc).replace(tzinfo=None),
    )
    db.add(fp)
    return fp
