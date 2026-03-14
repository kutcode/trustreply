"""Document generator — writes answers back into .docx files preserving formatting."""

from __future__ import annotations
from pathlib import Path
from copy import deepcopy

from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

from app.services.parser import ExtractedItem, RunFormat
from app.utils.questions import REVIEW_REQUIRED_PLACEHOLDER


def _apply_formatting(run, fmt: RunFormat | None) -> None:
    """Apply captured formatting to a new run."""
    if fmt is None:
        return

    if fmt.font_name:
        run.font.name = fmt.font_name
        # Also set eastAsia font for full compatibility
        r_pr = run._element.get_or_add_rPr()
        r_fonts = r_pr.find(qn('w:rFonts'))
        if r_fonts is not None:
            r_fonts.set(qn('w:eastAsia'), fmt.font_name)

    if fmt.font_size:
        run.font.size = fmt.font_size

    if fmt.bold is not None:
        run.font.bold = fmt.bold

    if fmt.italic is not None:
        run.font.italic = fmt.italic

    if fmt.underline is not None:
        run.font.underline = fmt.underline

    if fmt.color_rgb:
        run.font.color.rgb = fmt.color_rgb


def _clone_formatting_from_cell(cell) -> RunFormat | None:
    """Extract formatting from any existing content in a table cell."""
    for para in cell.paragraphs:
        for run in para.runs:
            fmt = RunFormat()
            font = run.font
            fmt.font_name = font.name
            fmt.font_size = font.size
            fmt.bold = font.bold
            fmt.italic = font.italic
            fmt.underline = font.underline
            if font.color and font.color.rgb:
                fmt.color_rgb = font.color.rgb
            return fmt
    return None


def _build_placeholder_format(base: RunFormat | None = None) -> RunFormat:
    """Style unresolved answers so they stand out during review."""

    fmt = deepcopy(base) if base is not None else RunFormat()
    fmt.italic = True
    fmt.color_rgb = RGBColor(0xCC, 0x00, 0x00)
    return fmt


def generate_filled_docx(
    source_path: Path,
    output_path: Path,
    items: list[ExtractedItem],
) -> Path:
    """Fill a .docx questionnaire with matched answers.

    For table-based questions: writes answers into the answer column cell.
    For paragraph-based questions: inserts answer as the next paragraph.
    Always preserves the original document's formatting.
    """
    doc = DocxDocument(str(source_path))

    # Index items by location type for efficient lookup
    table_items: list[ExtractedItem] = []
    para_items: dict[int, ExtractedItem] = {}

    for item in items:
        if item.item_type == "table_cell":
            table_items.append(item)
        elif item.item_type == "paragraph":
            para_items[item.location["para_idx"]] = item

    # ── Fill table cells ─────────────────────────────────────────────
    for item in table_items:
        table_idx = item.location["table_idx"]
        answer_row_idx = item.location.get("answer_row_idx", item.location.get("row_idx"))
        question_row_idx = item.location.get("q_row_idx", answer_row_idx)
        a_col_idx = item.location.get("a_col_idx", item.location.get("answer_col_idx", 1))
        q_col_idx = item.location.get("q_col_idx", item.location.get("question_col_idx", 0))

        if table_idx >= len(doc.tables):
            continue

        table = doc.tables[table_idx]
        if answer_row_idx is None or answer_row_idx >= len(table.rows):
            continue

        answer_cells = table.rows[answer_row_idx].cells
        if a_col_idx >= len(answer_cells):
            continue

        answer_cell = answer_cells[a_col_idx]

        output_text = item.answer_text or REVIEW_REQUIRED_PLACEHOLDER

        # Determine formatting: use item's captured formatting,
        # or clone from existing cell content, or from the source question cell.
        fmt = item.formatting
        if fmt is None:
            fmt = _clone_formatting_from_cell(answer_cell)
        if (
            fmt is None
            and question_row_idx is not None
            and 0 <= question_row_idx < len(table.rows)
        ):
            question_cells = table.rows[question_row_idx].cells
            if 0 <= q_col_idx < len(question_cells):
                fmt = _clone_formatting_from_cell(question_cells[q_col_idx])
        if item.answer_text is None:
            fmt = _build_placeholder_format(fmt)

        # Clear existing content in the answer cell
        for para in answer_cell.paragraphs:
            for run in para.runs:
                run.text = ""

        # Write the answer
        if answer_cell.paragraphs:
            target_para = answer_cell.paragraphs[0]
            for run in target_para.runs:
                run.text = ""
            if target_para.runs:
                target_para.runs[0].text = output_text
                _apply_formatting(target_para.runs[0], fmt)
            else:
                run = target_para.add_run(output_text)
                _apply_formatting(run, fmt)
        else:
            para = answer_cell.add_paragraph()
            run = para.add_run(output_text)
            _apply_formatting(run, fmt)

    # ── Fill paragraph-based answers ─────────────────────────────────
    # We need to insert answers AFTER the question paragraph.
    # Work backwards to avoid index shifting.
    para_indices = sorted(para_items.keys(), reverse=True)
    paragraphs = doc.paragraphs

    for p_idx in para_indices:
        if p_idx >= len(paragraphs):
            continue

        item = para_items[p_idx]
        question_para = paragraphs[p_idx]
        output_text = item.answer_text or REVIEW_REQUIRED_PLACEHOLDER
        output_format = item.formatting if item.answer_text is not None else _build_placeholder_format(item.formatting)

        # Check if the next paragraph is empty (an answer placeholder)
        if p_idx + 1 < len(paragraphs):
            next_para = paragraphs[p_idx + 1]
            next_text = next_para.text.strip()
            if len(next_text) < 3:
                # Replace the empty/placeholder paragraph with the answer
                if next_para.runs:
                    next_para.runs[0].text = output_text
                    _apply_formatting(next_para.runs[0], output_format)
                else:
                    run = next_para.add_run(output_text)
                    _apply_formatting(run, output_format)
                continue

        # Otherwise, insert a new paragraph after the question
        # We do this by adding a paragraph element after the question's XML element
        new_para_element = deepcopy(question_para._element)
        # Clear the content of the copied element
        for child in new_para_element.findall(qn('w:r')):
            new_para_element.remove(child)

        question_para._element.addnext(new_para_element)

        # Now add a run with the answer text
        from docx.oxml import OxmlElement
        r_element = OxmlElement('w:r')
        t_element = OxmlElement('w:t')
        t_element.text = output_text
        t_element.set(qn('xml:space'), 'preserve')
        r_element.append(t_element)
        new_para_element.append(r_element)

    # Save the filled document
    doc.save(str(output_path))
    return output_path


def generate_docx_from_pdf_items(
    output_path: Path,
    items: list[ExtractedItem],
) -> Path:
    """Generate a .docx file from PDF-extracted questions and their answers.

    Since PDFs can't be edited, we create a new .docx with Q&A pairs.
    """
    doc = DocxDocument()

    # Add a title
    doc.add_heading("Questionnaire Responses", level=1)
    doc.add_paragraph(
        "This document was generated from a PDF questionnaire. "
        "Answers have been auto-filled from the knowledge base."
    )
    doc.add_paragraph("")

    for i, item in enumerate(items, 1):
        # Question (bold)
        q_para = doc.add_paragraph()
        q_run = q_para.add_run(f"{i}. {item.question_text}")
        q_run.bold = True

        # Answer
        a_para = doc.add_paragraph()
        if item.answer_text:
            a_run = a_para.add_run(item.answer_text)
            a_run.italic = False
        else:
            a_run = a_para.add_run(REVIEW_REQUIRED_PLACEHOLDER)
            a_run.italic = True
            if hasattr(a_run.font.color, 'rgb'):
                a_run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

        doc.add_paragraph("")  # spacing

    doc.save(str(output_path))
    return output_path
