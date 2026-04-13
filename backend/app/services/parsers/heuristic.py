"""Heuristic parser implementation for DOCX and PDF questionnaires."""

from __future__ import annotations

import re
from pathlib import Path

import pdfplumber
from docx import Document as DocxDocument
from docx.text.run import Run

from app.services.parsers.base import BaseParser
from app.services.parsers.types import ExtractedItem, ParseOptions, ParseResult, RunFormat
from app.utils.csv_files import read_csv_rows


# Patterns that indicate a line is likely a question
QUESTION_PATTERNS = [
    re.compile(r"\?\s*$"),
    re.compile(r"^\d+[\.\)]\s+.{10,}"),
    re.compile(r"^[a-zA-Z][\.\)]\s+.{10,}"),
    re.compile(
        r"^(please\s+)?(describe|explain|provide|list|detail|specify|state|outline|indicate|confirm|identify)",
        re.IGNORECASE,
    ),
]

# Instructions/headings that look like prompts but are not answerable questions.
NON_QUESTION_PATTERNS = [
    re.compile(r"\bplease\s+provide\s+detailed\s+responses?\s+to\s+(each|all)\s+questions?\s+below\b", re.IGNORECASE),
    re.compile(r"\bdetailed\s+responses?\s+required\b", re.IGNORECASE),
    re.compile(r"\bdetailed\s+assessment\s+with\s+verbose\s+questions\b", re.IGNORECASE),
]

# Common questionnaire labels that precede an answer area
LABEL_PATTERNS = [
    re.compile(
        r"^(company\s+name|organization|contact|address|phone|email|website|date|name|title|role|department)",
        re.IGNORECASE,
    ),
    re.compile(
        r"^(policy|procedure|description|overview|summary|response|answer|comments?|notes?)\s*:?\s*$",
        re.IGNORECASE,
    ),
]

QUESTION_HEADER_PATTERNS = [
    re.compile(r"\bquestion(s)?\b", re.IGNORECASE),
    re.compile(r"\b(prompt|requirement|control|field)\b", re.IGNORECASE),
]

ANSWER_HEADER_PATTERNS = [
    re.compile(r"\banswer(s)?\b", re.IGNORECASE),
    re.compile(r"\bresponse(s)?\b", re.IGNORECASE),
    re.compile(r"\b(comment|comments|notes?)\b", re.IGNORECASE),
]

DEFAULT_MIN_QUESTION_LENGTH = 10


def _is_question(text: str, min_question_length: int = DEFAULT_MIN_QUESTION_LENGTH) -> bool:
    """Heuristic: does this text look like a question or field label?"""

    text = text.strip()
    if len(text) < min_question_length:
        return False
    for pattern in NON_QUESTION_PATTERNS:
        if pattern.search(text):
            return False
    for pattern in QUESTION_PATTERNS:
        if pattern.search(text):
            return True
    for pattern in LABEL_PATTERNS:
        if pattern.search(text):
            return True
    return False


def _extract_run_format(runs: list[Run]) -> RunFormat | None:
    """Extract formatting from the first run that has explicit formatting."""

    for run in runs:
        fmt = RunFormat()
        font = run.font
        fmt.font_name = font.name
        fmt.font_size = font.size
        fmt.bold = font.bold
        fmt.italic = font.italic
        fmt.underline = font.underline
        if font.color and font.color.rgb:
            fmt.color_rgb = font.color.rgb
        if any([fmt.font_name, fmt.font_size, fmt.bold, fmt.italic]):
            return fmt
    return RunFormat()


def _score_parse_confidence(stats: dict[str, int], items: list[ExtractedItem]) -> float:
    """Compute a simple confidence score from parse coverage."""

    if not items:
        return 0.0

    table_items = stats.get("table_items", 0) + stats.get("pdf_table_items", 0) + stats.get("csv_items", 0)
    paragraph_items = stats.get("paragraph_items", 0) + stats.get("pdf_text_items", 0)

    if table_items and paragraph_items:
        return 0.95
    if table_items:
        return 0.9
    if paragraph_items:
        return 0.75
    return 0.5


def _fallback_decision(stats: dict[str, int], confidence: float) -> tuple[bool, str | None]:
    """Recommend fallback when heuristic extraction looks weak."""

    items_total = stats.get("items_total", 0)
    table_rows = (
        stats.get("table_rows_scanned", 0)
        + stats.get("pdf_table_rows_scanned", 0)
        + stats.get("csv_rows_scanned", 0)
    )
    table_items = stats.get("table_items", 0) + stats.get("pdf_table_items", 0) + stats.get("csv_items", 0)

    if items_total == 0:
        return True, "no_questions_found"
    if confidence < 0.8:
        return True, "low_confidence_parse"
    if table_rows >= 3 and table_items == 0:
        return True, "table_layout_not_understood"
    return False, None


def _normalize_text(text: str) -> str:
    """Normalize text for row-level duplicate checks."""

    return " ".join(text.split()).strip().lower()


def _row_texts(cells) -> list[str]:
    """Extract trimmed cell text for a row."""

    return [cell.text.strip() for cell in cells]


def _row_looks_blank(cells) -> bool:
    """Heuristic: row is an answer placeholder if every cell is blank/very short."""

    texts = _row_texts(cells)
    return all(len(text.strip()) < 3 for text in texts)


def _row_has_merged_question_pattern(cells, question_text: str) -> bool:
    """Detect horizontally merged question rows exposed as repeated cell text."""

    normalized_question = _normalize_text(question_text)
    if not normalized_question:
        return False

    texts = [_normalize_text(text) for text in _row_texts(cells) if text.strip()]
    if len(texts) < 2:
        return False

    return all(text == normalized_question for text in texts)


def _matches_any_pattern(text: str, patterns: list[re.Pattern[str]]) -> bool:
    """Check whether text matches any of the supplied patterns."""

    return any(pattern.search(text) for pattern in patterns)


def _infer_table_mapping(
    rows: list[list[str]],
    options: ParseOptions,
) -> tuple[int, int, int]:
    """Infer question/answer columns from headers or wider-table row patterns."""

    question_col_idx = options.question_column_index
    answer_col_idx = options.answer_column_index
    header_rows = options.header_rows

    if not rows or not options.auto_detect_columns:
        return question_col_idx, answer_col_idx, header_rows

    first_row = [cell.strip() for cell in rows[0]]
    if first_row:
        question_headers = [
            idx for idx, text in enumerate(first_row)
            if _matches_any_pattern(text, QUESTION_HEADER_PATTERNS)
        ]
        answer_headers = [
            idx for idx, text in enumerate(first_row)
            if _matches_any_pattern(text, ANSWER_HEADER_PATTERNS)
        ]

        if question_headers and answer_headers:
            question_col_idx = question_headers[0]
            answer_to_right = [idx for idx in answer_headers if idx > question_col_idx]
            answer_col_idx = answer_to_right[0] if answer_to_right else answer_headers[0]
            return question_col_idx, answer_col_idx, max(header_rows, 1)

    if len(rows[0]) <= 2:
        return question_col_idx, answer_col_idx, header_rows

    candidate_scores: dict[tuple[int, int], int] = {}

    for row in rows[header_rows:]:
        texts = [text.strip() for text in row]
        for col_idx, q_text in enumerate(texts[:-1]):
            if not q_text or not _is_question(q_text, options.min_question_length):
                continue

            for answer_idx in range(col_idx + 1, len(texts)):
                a_text = texts[answer_idx]
                if len(a_text) >= 3:
                    continue

                score = 3
                if answer_idx == col_idx + 1:
                    score += 2
                if all(len(texts[prior_idx].strip()) < options.min_question_length for prior_idx in range(col_idx)):
                    score += 1
                if any(texts[mid_idx].strip() for mid_idx in range(col_idx + 1, answer_idx)):
                    score -= 2

                pair = (col_idx, answer_idx)
                candidate_scores[pair] = candidate_scores.get(pair, 0) + score

    if candidate_scores:
        question_col_idx, answer_col_idx = max(
            candidate_scores.items(),
            key=lambda item: (item[1], -item[0][0], -item[0][1]),
        )[0]

    return question_col_idx, answer_col_idx, header_rows


class HeuristicParser(BaseParser):
    """Current deterministic parser with configurable table mapping."""

    strategy_name = "heuristic"

    def parse(self, file_path: Path, options: ParseOptions | None = None) -> ParseResult:
        options = options or ParseOptions()
        suffix = file_path.suffix.lower()

        if suffix == ".docx":
            return self.parse_docx(file_path, options)
        if suffix == ".pdf":
            return self.parse_pdf(file_path, options)
        if suffix == ".csv":
            return self.parse_csv(file_path, options)
        if suffix in (".xlsx", ".xls"):
            return self.parse_xlsx(file_path, options)
        raise ValueError(f"Unsupported file type: {suffix}. Supported: .docx, .pdf, .csv, .xlsx, .xls")

    def parse_docx(self, file_path: Path, options: ParseOptions | None = None) -> ParseResult:
        options = options or ParseOptions()
        doc = DocxDocument(str(file_path))

        table_items, table_stats = self._parse_docx_tables(doc, options)
        if options.scan_paragraphs:
            para_items, para_stats = self._parse_docx_paragraphs(doc, options)
        else:
            para_items = []
            para_stats = {"paragraphs_scanned": len(doc.paragraphs), "paragraph_items": 0}

        table_texts = {item.question_text.lower() for item in table_items}
        items = list(table_items)
        for item in para_items:
            if item.question_text.lower() not in table_texts:
                items.append(item)

        stats = {
            "tables_scanned": table_stats["tables_scanned"],
            "table_rows_scanned": table_stats["table_rows_scanned"],
            "table_items": table_stats["table_items"],
            "paragraphs_scanned": para_stats["paragraphs_scanned"],
            "paragraph_items": para_stats["paragraph_items"],
            "items_total": len(items),
            "deduplicated_items": len(table_items) + len(para_items) - len(items),
        }

        confidence = _score_parse_confidence(stats, items)
        fallback_recommended, fallback_reason = _fallback_decision(stats, confidence)

        return ParseResult(
            items=items,
            confidence=confidence,
            stats=stats,
            profile_name=options.profile_name,
            parser_strategy=self.strategy_name,
            fallback_recommended=fallback_recommended,
            fallback_reason=fallback_reason,
        )

    def parse_pdf(self, file_path: Path, options: ParseOptions | None = None) -> ParseResult:
        options = options or ParseOptions()
        items: list[ExtractedItem] = []
        stats = {
            "pages_scanned": 0,
            "pdf_tables_scanned": 0,
            "pdf_table_rows_scanned": 0,
            "pdf_table_items": 0,
            "pdf_text_lines_scanned": 0,
            "pdf_text_items": 0,
            "items_total": 0,
        }

        with pdfplumber.open(str(file_path)) as pdf:
            stats["pages_scanned"] = len(pdf.pages)
            for page_idx, page in enumerate(pdf.pages):
                tables = page.extract_tables() or []
                stats["pdf_tables_scanned"] += len(tables)
                for t_idx, table in enumerate(tables):
                    normalized_rows = [
                        [(cell or "").strip() for cell in (row or [])]
                        for row in (table or [])
                    ]
                    q_col_idx, a_col_idx, header_rows = _infer_table_mapping(normalized_rows, options)
                    for r_idx, row in enumerate(table or []):
                        stats["pdf_table_rows_scanned"] += 1
                        if r_idx < header_rows or not row:
                            continue

                        max_idx = max(q_col_idx, a_col_idx)
                        if len(row) <= max_idx:
                            continue
                        if not options.allow_additional_columns and len(row) > max_idx + 1:
                            continue

                        q_text = (row[q_col_idx] or "").strip()

                        if not q_text or len(q_text) < options.min_question_length:
                            continue

                        if _is_question(q_text, options.min_question_length):
                            items.append(
                                ExtractedItem(
                                    question_text=q_text,
                                    item_type="pdf_table",
                                    location={
                                        "page_idx": page_idx,
                                        "table_idx": t_idx,
                                        "row_idx": r_idx,
                                        "q_col_idx": q_col_idx,
                                        "a_col_idx": a_col_idx,
                                        "question_col_idx": q_col_idx,
                                        "answer_col_idx": a_col_idx,
                                        "column_count": len(row),
                                        "header_rows": header_rows,
                                    },
                                    source_block_id=(
                                        f"pdf-page-{page_idx}-table-{t_idx}-row-{r_idx}"
                                        f"-q{q_col_idx}-a{a_col_idx}"
                                    ),
                                    confidence=0.9,
                                    parser_strategy=self.strategy_name,
                                    raw_text=q_text,
                                )
                            )
                            stats["pdf_table_items"] += 1

                text = page.extract_text() or ""
                lines = text.split("\n")
                stats["pdf_text_lines_scanned"] += len(lines)
                for line_idx, line in enumerate(lines):
                    line = line.strip()
                    if _is_question(line, options.min_question_length):
                        already_found = any(item.question_text.lower() == line.lower() for item in items)
                        if not already_found:
                            items.append(
                                ExtractedItem(
                                    question_text=line,
                                    item_type="pdf_text",
                                    location={"page_idx": page_idx, "line_idx": line_idx},
                                    source_block_id=f"pdf-page-{page_idx}-line-{line_idx}",
                                    confidence=0.7,
                                    parser_strategy=self.strategy_name,
                                    raw_text=line,
                                )
                            )
                            stats["pdf_text_items"] += 1

        stats["items_total"] = len(items)
        confidence = _score_parse_confidence(stats, items)
        fallback_recommended, fallback_reason = _fallback_decision(stats, confidence)

        return ParseResult(
            items=items,
            confidence=confidence,
            stats=stats,
            profile_name=options.profile_name,
            parser_strategy=self.strategy_name,
            fallback_recommended=fallback_recommended,
            fallback_reason=fallback_reason,
        )

    def parse_csv(self, file_path: Path, options: ParseOptions | None = None) -> ParseResult:
        options = options or ParseOptions()
        rows, _csv_format = read_csv_rows(file_path)
        items, stats = self._parse_csv_rows(rows, options)
        confidence = _score_parse_confidence(stats, items)
        fallback_recommended, fallback_reason = _fallback_decision(stats, confidence)

        return ParseResult(
            items=items,
            confidence=confidence,
            stats=stats,
            profile_name=options.profile_name,
            parser_strategy=self.strategy_name,
            fallback_recommended=fallback_recommended,
            fallback_reason=fallback_reason,
        )

    def parse_xlsx(self, file_path: Path, options: ParseOptions | None = None) -> ParseResult:
        from app.services.parsers.excel_parser import parse_excel
        return parse_excel(file_path, options)

    def _parse_docx_tables(
        self,
        doc: DocxDocument,
        options: ParseOptions,
    ) -> tuple[list[ExtractedItem], dict[str, int]]:
        items: list[ExtractedItem] = []
        stats = {"tables_scanned": len(doc.tables), "table_rows_scanned": 0, "table_items": 0}

        for t_idx, table in enumerate(doc.tables):
            raw_rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            q_col_idx, a_col_idx, header_rows = _infer_table_mapping(raw_rows, options)
            for r_idx, row in enumerate(table.rows):
                stats["table_rows_scanned"] += 1
                if r_idx < header_rows:
                    continue

                cells = row.cells
                max_idx = max(q_col_idx, a_col_idx)
                if len(cells) <= max_idx:
                    continue
                if not options.allow_additional_columns and len(cells) > max_idx + 1:
                    continue

                q_cell = cells[q_col_idx]
                a_cell = cells[a_col_idx]
                q_text = q_cell.text.strip()

                if not q_text or len(q_text) < options.min_question_length:
                    continue

                if _is_question(q_text, options.min_question_length):
                    answer_row_idx = r_idx
                    answer_col_idx = a_col_idx
                    answer_cell = a_cell

                    # When a question row is merged horizontally, python-docx
                    # often exposes the same question text in multiple cells.
                    # Prefer the following blank row as the answer area so we
                    # don't overwrite the merged question cell itself.
                    if (
                        options.detect_row_blocks
                        and _row_has_merged_question_pattern(cells, q_text)
                        and r_idx + 1 < len(table.rows)
                    ):
                        next_row = table.rows[r_idx + 1]
                        if _row_looks_blank(next_row.cells):
                            answer_row_idx = r_idx + 1
                            if len(next_row.cells) <= answer_col_idx:
                                answer_col_idx = 0
                            answer_cell = next_row.cells[answer_col_idx]

                    fmt = None
                    if answer_cell.paragraphs:
                        for para in answer_cell.paragraphs:
                            if para.runs:
                                fmt = _extract_run_format(para.runs)
                                break
                    if fmt is None and q_cell.paragraphs:
                        for para in q_cell.paragraphs:
                            if para.runs:
                                fmt = _extract_run_format(para.runs)
                                break

                    items.append(
                        ExtractedItem(
                            question_text=q_text,
                            item_type="table_cell",
                            location={
                                "table_idx": t_idx,
                                "row_idx": answer_row_idx,
                                "q_row_idx": r_idx,
                                "answer_row_idx": answer_row_idx,
                                "q_col_idx": q_col_idx,
                                "a_col_idx": answer_col_idx,
                                "question_col_idx": q_col_idx,
                                "answer_col_idx": answer_col_idx,
                                "column_count": len(cells),
                                "header_rows": header_rows,
                                "layout_hint": "row_block" if answer_row_idx != r_idx else "same_row",
                            },
                            formatting=fmt,
                            source_block_id=(
                                f"docx-table-{t_idx}-qrow-{r_idx}-arow-{answer_row_idx}"
                                f"-q{q_col_idx}-a{answer_col_idx}"
                            ),
                            confidence=0.9 if answer_row_idx == r_idx else 0.85,
                            parser_strategy=self.strategy_name,
                            raw_text=q_text,
                        )
                    )
                    stats["table_items"] += 1

        return items, stats

    def _parse_docx_paragraphs(
        self,
        doc: DocxDocument,
        options: ParseOptions,
    ) -> tuple[list[ExtractedItem], dict[str, int]]:
        items: list[ExtractedItem] = []
        paragraphs = doc.paragraphs
        stats = {"paragraphs_scanned": len(paragraphs), "paragraph_items": 0}

        for p_idx, para in enumerate(paragraphs):
            text = para.text.strip()
            if not text:
                continue

            if _is_question(text, options.min_question_length):
                fmt = None
                if p_idx + 1 < len(paragraphs):
                    next_para = paragraphs[p_idx + 1]
                    if next_para.runs:
                        fmt = _extract_run_format(next_para.runs)
                if fmt is None and para.runs:
                    fmt = _extract_run_format(para.runs)

                items.append(
                    ExtractedItem(
                        question_text=text,
                        item_type="paragraph",
                        location={"para_idx": p_idx},
                        formatting=fmt,
                        source_block_id=f"docx-paragraph-{p_idx}",
                        confidence=0.7,
                        parser_strategy=self.strategy_name,
                        raw_text=text,
                    )
                )
                stats["paragraph_items"] += 1

        return items, stats

    def _parse_csv_rows(
        self,
        rows: list[list[str]],
        options: ParseOptions,
    ) -> tuple[list[ExtractedItem], dict[str, int]]:
        items: list[ExtractedItem] = []
        stats = {"csv_rows_scanned": len(rows), "csv_items": 0, "items_total": 0}

        non_empty_rows = [row for row in rows if any(cell.strip() for cell in row)]
        if not non_empty_rows:
            return items, stats

        q_col_idx, a_col_idx, header_rows = _infer_table_mapping(non_empty_rows, options)
        seen_non_empty_rows = 0

        for r_idx, row in enumerate(rows):
            if not any(cell.strip() for cell in row):
                continue

            if seen_non_empty_rows < header_rows:
                seen_non_empty_rows += 1
                continue

            max_idx = max(q_col_idx, a_col_idx)
            if len(row) <= max_idx:
                seen_non_empty_rows += 1
                continue
            if not options.allow_additional_columns and len(row) > max_idx + 1:
                seen_non_empty_rows += 1
                continue

            q_text = (row[q_col_idx] or "").strip()

            if not q_text or len(q_text) < options.min_question_length:
                seen_non_empty_rows += 1
                continue

            if _is_question(q_text, options.min_question_length):
                items.append(
                    ExtractedItem(
                        question_text=q_text,
                        item_type="csv_row",
                        location={
                            "row_idx": r_idx,
                            "q_col_idx": q_col_idx,
                            "a_col_idx": a_col_idx,
                            "question_col_idx": q_col_idx,
                            "answer_col_idx": a_col_idx,
                            "column_count": len(row),
                            "header_rows": header_rows,
                        },
                        source_block_id=f"csv-row-{r_idx}-q{q_col_idx}-a{a_col_idx}",
                        confidence=0.88,
                        parser_strategy=self.strategy_name,
                        raw_text=q_text,
                    )
                )
                stats["csv_items"] += 1

            seen_non_empty_rows += 1

        stats["items_total"] = len(items)
        return items, stats
