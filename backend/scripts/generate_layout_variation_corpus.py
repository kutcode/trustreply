#!/usr/bin/env python3
"""Generate a second 50-document stress corpus with broader layout variations."""

from __future__ import annotations

import csv
import shutil
from dataclasses import dataclass
from pathlib import Path

from docx import Document as DocxDocument


SCRIPT_DIR = Path(__file__).resolve().parent
BACKEND_DIR = SCRIPT_DIR.parent
PROJECT_ROOT = BACKEND_DIR.parent
OUTPUT_DIR = PROJECT_ROOT / "test-data" / "Test-3"
MANIFEST_PATH = OUTPUT_DIR / "manifest.csv"
README_PATH = OUTPUT_DIR / "README.md"


@dataclass(frozen=True)
class Entry:
    text: str
    status: str  # known | unknown


THEMES: dict[str, dict[str, object]] = {
    "security": {
        "category": "Security",
        "known_questions": [
            "Describe your endpoint protection program.",
            "How do you manage privileged access across production systems?",
            "Describe your vulnerability management lifecycle.",
            "How do you monitor for suspicious activity in your environment?",
            "Describe your incident response process.",
            "How do you secure laptops and mobile devices used by employees?",
        ],
        "known_labels": [
            "Company Name:",
            "Security Contact Name:",
            "Security Contact Email:",
        ],
        "unknown_questions": [
            "Do you conduct annual phishing simulations for employees?",
            "How do you manage secrets in application code and CI pipelines?",
            "Describe your secure software development lifecycle requirements.",
            "How often do you perform external penetration testing?",
            "Do you scan source code for exposed credentials before release?",
            "Describe how firewall rule changes are reviewed and approved.",
        ],
        "unknown_labels": [
            "Security Audit Lead:",
            "Security Testing Window:",
            "Bug Bounty Contact:",
        ],
    },
    "privacy": {
        "category": "Privacy",
        "known_questions": [
            "Describe your personal data retention policy.",
            "How do you respond to data subject access requests?",
            "Describe your lawful basis review process for processing personal data.",
            "How do you manage subprocessors that access customer data?",
            "Describe your deletion workflow for customer records.",
            "How do you monitor compliance with privacy obligations?",
        ],
        "known_labels": [
            "Legal Entity Name:",
            "Privacy Contact Name:",
            "Privacy Contact Email:",
        ],
        "unknown_questions": [
            "Do you use cookies or similar tracking technologies on customer-facing sites?",
            "How do you manage consent records for marketing communications?",
            "Describe your process for handling cross-border transfer impact assessments.",
            "Do you support browser-based opt-out signals?",
            "How do you review new AI features for privacy impact?",
            "Describe your process for managing children's data.",
        ],
        "unknown_labels": [
            "DPO Name:",
            "Cookie Compliance Contact:",
            "Transfer Assessment Owner:",
        ],
    },
    "operations": {
        "category": "Operations",
        "known_questions": [
            "Describe your change management process.",
            "How do you manage production deployments?",
            "Describe your service monitoring approach.",
            "How do you track and resolve service incidents?",
            "Describe your backup verification process.",
            "How do you manage access for third-party support personnel?",
        ],
        "known_labels": [
            "Operating Entity Name:",
            "Operations Contact Name:",
            "Operations Contact Email:",
        ],
        "unknown_questions": [
            "How do you plan and announce customer maintenance windows?",
            "Describe your capacity planning process for peak traffic.",
            "Do you use canary or blue-green deployments for major releases?",
            "How are runbooks reviewed and updated?",
            "Describe your dependency upgrade planning workflow.",
            "How do you detect configuration drift across production systems?",
        ],
        "unknown_labels": [
            "Release Manager:",
            "Maintenance Window Owner:",
            "Capacity Planning Contact:",
        ],
    },
    "continuity": {
        "category": "Business Continuity",
        "known_questions": [
            "Describe your disaster recovery plan.",
            "How do you test business continuity procedures?",
            "Describe your recovery time and recovery point objectives.",
            "How do you ensure availability during infrastructure failures?",
            "Describe your crisis communication process.",
            "How do you restore critical services after a disruptive event?",
        ],
        "known_labels": [
            "Entity Name:",
            "Continuity Contact Name:",
            "Continuity Contact Email:",
        ],
        "unknown_questions": [
            "Do you maintain alternate workspace arrangements for office disruptions?",
            "How do you prioritize manual workarounds during system outages?",
            "Describe your supply chain contingency planning process.",
            "How frequently do you validate emergency contact trees?",
            "Do you maintain backup communication channels for crisis coordination?",
            "How do you assess regional concentration risk for critical services?",
        ],
        "unknown_labels": [
            "Crisis Communications Lead:",
            "Emergency Hotline Owner:",
            "Alternate Site Coordinator:",
        ],
    },
    "vendor": {
        "category": "Vendor Management",
        "known_questions": [
            "Describe your third-party risk management process.",
            "How do you review and onboard critical vendors?",
            "Describe your vendor contract security review workflow.",
            "How do you monitor vendor performance and compliance?",
            "Describe your offboarding process for third-party service providers.",
            "How do you assess vendor incidents that affect customer services?",
        ],
        "known_labels": [
            "Vendor Legal Name:",
            "Primary Contact Name:",
            "Primary Contact Email:",
        ],
        "unknown_questions": [
            "Do you require independent assurance reports from high-risk vendors?",
            "How do you review vendor concentration risk?",
            "Describe your process for tracking vendor data return certificates.",
            "How do you evaluate fourth-party risk disclosures?",
            "Do you maintain minimum security control baselines for vendors?",
            "How do you handle vendors that refuse standard contractual terms?",
        ],
        "unknown_labels": [
            "Procurement Risk Owner:",
            "Fourth-Party Review Contact:",
            "Vendor Exit Coordinator:",
        ],
    },
}


LAYOUTS: list[dict[str, str]] = [
    {"slug": "prompt_response_matrix", "label": "Prompt/response table with intro text", "profile": "default"},
    {"slug": "three_column_domain_grid", "label": "Three-column domain grid", "profile": "three_column_table"},
    {"slug": "control_reference_matrix", "label": "Four-column control reference matrix", "profile": "default"},
    {"slug": "metadata_preamble_packet", "label": "Metadata preamble followed by question table", "profile": "default"},
    {"slug": "dual_section_headers", "label": "Dual-section tables with mixed response headers", "profile": "default"},
    {"slug": "merged_followup_blocks", "label": "Merged follow-up row blocks", "profile": "row_block_questionnaire"},
    {"slug": "sectioned_followup_blocks", "label": "Sectioned merged follow-up blocks", "profile": "row_block_questionnaire"},
    {"slug": "numbered_packet", "label": "Numbered paragraph packet", "profile": "default"},
    {"slug": "hybrid_review_packet", "label": "Hybrid metadata, paragraph, and table packet", "profile": "default"},
    {"slug": "audit_registry", "label": "Audit registry table", "profile": "default"},
]


def rotate_select(items: list[str], count: int, offset: int) -> list[str]:
    return [items[(offset + index) % len(items)] for index in range(count)]


def make_entries(values: list[str], status: str) -> list[Entry]:
    return [Entry(text=value, status=status) for value in values]


def weave(*groups: list[Entry]) -> list[Entry]:
    buckets = [list(group) for group in groups]
    items: list[Entry] = []
    while any(buckets):
        for bucket in buckets:
            if bucket:
                items.append(bucket.pop(0))
    return items


def build_question_entries(theme: dict[str, object], *, known_count: int, unknown_count: int, offset: int) -> list[Entry]:
    known = make_entries(rotate_select(list(theme["known_questions"]), known_count, offset), "known")
    unknown = make_entries(rotate_select(list(theme["unknown_questions"]), unknown_count, offset * 2), "unknown")
    return weave(known, unknown)


def build_label_entries(theme: dict[str, object], *, known_count: int, unknown_count: int, offset: int) -> list[Entry]:
    known = make_entries(rotate_select(list(theme["known_labels"]), known_count, offset), "known")
    unknown = make_entries(rotate_select(list(theme["unknown_labels"]), unknown_count, offset), "unknown")
    return weave(known, unknown)


def add_title(doc: DocxDocument, title: str, subtitle: str) -> None:
    doc.add_heading(title, level=0)
    paragraph = doc.add_paragraph(subtitle)
    paragraph.style = "Intense Quote"


def render_prompt_response(doc: DocxDocument, entries: list[Entry], *, header_left: str = "Question", header_right: str = "Response") -> None:
    table = doc.add_table(rows=len(entries) + 1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = header_left
    table.rows[0].cells[1].text = header_right
    for index, entry in enumerate(entries, start=1):
        table.rows[index].cells[0].text = entry.text
        table.rows[index].cells[1].text = ""


def render_three_column(doc: DocxDocument, entries: list[Entry], category: str) -> None:
    table = doc.add_table(rows=len(entries) + 1, cols=3)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Domain"
    table.rows[0].cells[1].text = "Question"
    table.rows[0].cells[2].text = "Response"
    for index, entry in enumerate(entries, start=1):
        table.rows[index].cells[0].text = category
        table.rows[index].cells[1].text = entry.text
        table.rows[index].cells[2].text = ""


def render_four_column(doc: DocxDocument, entries: list[Entry], theme_key: str, *, reference_label: str) -> None:
    prefix = theme_key[:3].upper()
    table = doc.add_table(rows=len(entries) + 1, cols=4)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Domain"
    table.rows[0].cells[1].text = reference_label
    table.rows[0].cells[2].text = "Question"
    table.rows[0].cells[3].text = "Response"
    for index, entry in enumerate(entries, start=1):
        table.rows[index].cells[0].text = theme_key.title()
        table.rows[index].cells[1].text = f"{prefix}-{index:02d}"
        table.rows[index].cells[2].text = entry.text
        table.rows[index].cells[3].text = ""


def render_metadata_table(doc: DocxDocument, entries: list[Entry]) -> None:
    table = doc.add_table(rows=len(entries) + 1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Question"
    table.rows[0].cells[1].text = "Answer"
    for index, entry in enumerate(entries, start=1):
        table.rows[index].cells[0].text = entry.text
        table.rows[index].cells[1].text = ""


def render_row_block(doc: DocxDocument, entries: list[Entry], *, sectioned: bool) -> None:
    groups = [("Questionnaire", entries)]
    if sectioned:
        midpoint = len(entries) // 2
        groups = [("Primary Review", entries[:midpoint]), ("Secondary Review", entries[midpoint:])]

    for title, group_entries in groups:
        doc.add_heading(title, level=1)
        table = doc.add_table(rows=0, cols=2)
        table.style = "Table Grid"
        for entry in group_entries:
            question_row = table.add_row()
            merged = question_row.cells[0].merge(question_row.cells[1])
            merged.text = entry.text
            answer_row = table.add_row()
            answer_row.cells[0].text = ""
            answer_row.cells[1].text = ""


def render_numbered(doc: DocxDocument, entries: list[Entry]) -> None:
    for index, entry in enumerate(entries, start=1):
        doc.add_paragraph(f"{index}. {entry.text}")
        doc.add_paragraph("")


def build_layout_entries(theme: dict[str, object], layout_index: int) -> dict[str, list[Entry]]:
    offset = layout_index - 1

    if layout_index == 1:
        return {"main": weave(build_label_entries(theme, known_count=1, unknown_count=1, offset=offset), build_question_entries(theme, known_count=2, unknown_count=3, offset=offset))}
    if layout_index == 2:
        return {"main": weave(build_label_entries(theme, known_count=2, unknown_count=0, offset=offset), build_question_entries(theme, known_count=2, unknown_count=2, offset=offset + 1))}
    if layout_index == 3:
        return {"main": weave(build_label_entries(theme, known_count=1, unknown_count=1, offset=offset + 1), build_question_entries(theme, known_count=2, unknown_count=3, offset=offset + 2))}
    if layout_index == 4:
        return {
            "metadata": build_label_entries(theme, known_count=2, unknown_count=1, offset=offset),
            "main": build_question_entries(theme, known_count=3, unknown_count=2, offset=offset + 1),
        }
    if layout_index == 5:
        return {
            "section_a": build_question_entries(theme, known_count=2, unknown_count=2, offset=offset),
            "section_b": weave(build_label_entries(theme, known_count=1, unknown_count=0, offset=offset), build_question_entries(theme, known_count=1, unknown_count=2, offset=offset + 2)),
        }
    if layout_index == 6:
        return {"main": weave(build_label_entries(theme, known_count=1, unknown_count=1, offset=offset), build_question_entries(theme, known_count=2, unknown_count=3, offset=offset + 1))}
    if layout_index == 7:
        return {"main": weave(build_label_entries(theme, known_count=2, unknown_count=0, offset=offset), build_question_entries(theme, known_count=2, unknown_count=4, offset=offset + 2))}
    if layout_index == 8:
        return {"main": weave(build_label_entries(theme, known_count=1, unknown_count=1, offset=offset), build_question_entries(theme, known_count=2, unknown_count=3, offset=offset + 3))}
    if layout_index == 9:
        return {
            "metadata": build_label_entries(theme, known_count=1, unknown_count=1, offset=offset),
            "table": build_question_entries(theme, known_count=2, unknown_count=2, offset=offset + 1),
            "paragraph": weave(build_question_entries(theme, known_count=1, unknown_count=2, offset=offset + 2)),
        }

    return {"main": weave(build_label_entries(theme, known_count=1, unknown_count=1, offset=offset), build_question_entries(theme, known_count=3, unknown_count=2, offset=offset + 1))}


def count_status(entries: dict[str, list[Entry]]) -> tuple[int, int]:
    items = [entry for group in entries.values() for entry in group]
    known = sum(1 for entry in items if entry.status == "known")
    unknown = sum(1 for entry in items if entry.status == "unknown")
    return known, unknown


def render_document(theme_key: str, theme: dict[str, object], layout_index: int, output_path: Path) -> tuple[int, int]:
    category = str(theme["category"])
    entries = build_layout_entries(theme, layout_index)
    known_count, unknown_count = count_status(entries)

    doc = DocxDocument()
    add_title(
        doc,
        f"{category} Layout Variation Questionnaire",
        "This document mixes known and unknown prompts across alternate layout structures for parser stress testing.",
    )

    if layout_index == 1:
        doc.add_paragraph("Please complete the matrix below using concise answers.")
        render_prompt_response(doc, entries["main"], header_left="Question", header_right="Response")
    elif layout_index == 2:
        doc.add_heading("Domain Review Grid", level=1)
        render_three_column(doc, entries["main"], category)
    elif layout_index == 3:
        doc.add_heading("Control Mapping", level=1)
        render_four_column(doc, entries["main"], theme_key, reference_label="Reference")
    elif layout_index == 4:
        doc.add_heading("Document Metadata", level=1)
        render_metadata_table(doc, entries["metadata"])
        doc.add_heading("Questionnaire", level=1)
        render_prompt_response(doc, entries["main"], header_left="Question", header_right="Answer")
    elif layout_index == 5:
        doc.add_heading("Core Review", level=1)
        render_prompt_response(doc, entries["section_a"], header_left="Question", header_right="Answer")
        doc.add_heading("Follow-up Review", level=1)
        render_prompt_response(doc, entries["section_b"], header_left="Question", header_right="Response")
    elif layout_index == 6:
        render_row_block(doc, entries["main"], sectioned=False)
    elif layout_index == 7:
        render_row_block(doc, entries["main"], sectioned=True)
    elif layout_index == 8:
        doc.add_heading("Narrative Packet", level=1)
        render_numbered(doc, entries["main"])
    elif layout_index == 9:
        doc.add_heading("Metadata", level=1)
        render_metadata_table(doc, entries["metadata"])
        doc.add_heading("Quick Response Table", level=1)
        render_prompt_response(doc, entries["table"], header_left="Question", header_right="Response")
        doc.add_heading("Additional Narrative", level=1)
        render_numbered(doc, entries["paragraph"])
    else:
        doc.add_heading("Audit Registry", level=1)
        render_four_column(doc, entries["main"], theme_key, reference_label="Reference")

    doc.save(str(output_path))
    return known_count, unknown_count


def write_manifest(rows: list[dict[str, object]]) -> None:
    with MANIFEST_PATH.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.DictWriter(
            handle,
            fieldnames=[
                "file",
                "theme",
                "layout_family",
                "suggested_profile",
                "total_prompts",
                "expected_known_matches",
                "expected_unknown_flags",
            ],
        )
        writer.writeheader()
        writer.writerows(rows)


def write_readme(rows: list[dict[str, object]]) -> None:
    lines = [
        "# Layout Variation Stress-Test Corpus",
        "",
        "This folder contains 50 generated `.docx` questionnaires in a fresh layout-variation batch.",
        "",
        "What this corpus is designed to test:",
        "- Known prompts that should resolve against the sample knowledge base",
        "- Unknown prompts that should be flagged for review",
        "- Broader table, metadata, numbered-paragraph, and row-block document structures",
        "- Alternate document packaging that still keeps answer cells blank for upload testing",
        "",
        "Suggested parser profiles:",
        "- `default`: standard tables, metadata packets, numbered paragraphs, and hybrid documents",
        "- `three_column_table`: three-column grids with the question in the middle column",
        "- `row_block_questionnaire`: merged question rows with blank answer rows beneath",
        "",
        "See `manifest.csv` for per-file expected counts.",
        "",
        "| File | Layout Family | Theme | Suggested Profile | Expected Known | Expected Unknown |",
        "|---|---|---|---|---:|---:|",
    ]

    for row in rows:
        lines.append(
            f"| `{row['file']}` | {row['layout_family']} | {row['theme']} | `{row['suggested_profile']}` | {row['expected_known_matches']} | {row['expected_unknown_flags']} |"
        )

    README_PATH.write_text("\n".join(lines) + "\n", encoding="utf-8")


def main() -> None:
    if OUTPUT_DIR.exists():
        shutil.rmtree(OUTPUT_DIR)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    manifest_rows: list[dict[str, object]] = []
    theme_order = ["security", "privacy", "operations", "continuity", "vendor"]

    file_number = 1
    for theme_key in theme_order:
        theme = THEMES[theme_key]
        for layout_index, layout in enumerate(LAYOUTS, start=1):
            filename = f"{file_number:02d}_{layout['slug']}_{theme_key}.docx"
            output_path = OUTPUT_DIR / filename
            known_count, unknown_count = render_document(theme_key, theme, layout_index, output_path)
            manifest_rows.append(
                {
                    "file": filename,
                    "theme": str(theme["category"]),
                    "layout_family": layout["label"],
                    "suggested_profile": layout["profile"],
                    "total_prompts": known_count + unknown_count,
                    "expected_known_matches": known_count,
                    "expected_unknown_flags": unknown_count,
                }
            )
            file_number += 1

    write_manifest(manifest_rows)
    write_readme(manifest_rows)
    print(f"Generated {len(manifest_rows)} documents in {OUTPUT_DIR}")


if __name__ == "__main__":
    main()
