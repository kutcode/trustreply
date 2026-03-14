#!/usr/bin/env python3
"""Generate a 50-document questionnaire corpus with mixed known and unknown prompts."""

from __future__ import annotations

import csv
import shutil
from dataclasses import dataclass
from pathlib import Path

from docx import Document as DocxDocument


SCRIPT_DIR = Path(__file__).resolve().parent
BACKEND_DIR = SCRIPT_DIR.parent
PROJECT_ROOT = BACKEND_DIR.parent
OUTPUT_DIR = PROJECT_ROOT / "test-data" / "Test-2"
MANIFEST_PATH = OUTPUT_DIR / "manifest.csv"
README_PATH = OUTPUT_DIR / "README.md"


@dataclass(frozen=True)
class Entry:
    text: str
    status: str  # "known" or "unknown"


THEMES: dict[str, dict[str, object]] = {
    "security": {
        "category": "Security",
        "known_questions": [
            "Describe your endpoint protection program.",
            "How do you manage privileged access across production systems?",
            "Describe your vulnerability management lifecycle.",
            "How do you monitor for suspicious activity in your environment?",
            "Describe your incident response process.",
            "How do you secure laptops and mobile devices used by employees?",
        ],
        "known_labels": [
            "Company Name:",
            "Security Contact Name:",
            "Security Contact Email:",
        ],
        "unknown_questions": [
            "Do you conduct annual phishing simulations for employees?",
            "How do you manage secrets in application code and CI pipelines?",
            "Describe your secure software development lifecycle requirements.",
            "How often do you perform external penetration testing?",
            "Do you scan source code for exposed credentials before release?",
            "Describe how firewall rule changes are reviewed and approved.",
        ],
        "unknown_labels": [
            "Security Audit Lead:",
            "Security Testing Window:",
            "Bug Bounty Contact:",
        ],
    },
    "privacy": {
        "category": "Privacy",
        "known_questions": [
            "Describe your personal data retention policy.",
            "How do you respond to data subject access requests?",
            "Describe your lawful basis review process for processing personal data.",
            "How do you manage subprocessors that access customer data?",
            "Describe your deletion workflow for customer records.",
            "How do you monitor compliance with privacy obligations?",
        ],
        "known_labels": [
            "Legal Entity Name:",
            "Privacy Contact Name:",
            "Privacy Contact Email:",
        ],
        "unknown_questions": [
            "Do you use cookies or similar tracking technologies on customer-facing sites?",
            "How do you manage consent records for marketing communications?",
            "Describe your process for handling cross-border transfer impact assessments.",
            "Do you support browser-based opt-out signals?",
            "How do you review new AI features for privacy impact?",
            "Describe your process for managing children's data.",
        ],
        "unknown_labels": [
            "DPO Name:",
            "Cookie Compliance Contact:",
            "Transfer Assessment Owner:",
        ],
    },
    "operations": {
        "category": "Operations",
        "known_questions": [
            "Describe your change management process.",
            "How do you manage production deployments?",
            "Describe your service monitoring approach.",
            "How do you track and resolve service incidents?",
            "Describe your backup verification process.",
            "How do you manage access for third-party support personnel?",
        ],
        "known_labels": [
            "Operating Entity Name:",
            "Operations Contact Name:",
            "Operations Contact Email:",
        ],
        "unknown_questions": [
            "How do you plan and announce customer maintenance windows?",
            "Describe your capacity planning process for peak traffic.",
            "Do you use canary or blue-green deployments for major releases?",
            "How are runbooks reviewed and updated?",
            "Describe your dependency upgrade planning workflow.",
            "How do you detect configuration drift across production systems?",
        ],
        "unknown_labels": [
            "Release Manager:",
            "Maintenance Window Owner:",
            "Capacity Planning Contact:",
        ],
    },
    "continuity": {
        "category": "Business Continuity",
        "known_questions": [
            "Describe your disaster recovery plan.",
            "How do you test business continuity procedures?",
            "Describe your recovery time and recovery point objectives.",
            "How do you ensure availability during infrastructure failures?",
            "Describe your crisis communication process.",
            "How do you restore critical services after a disruptive event?",
        ],
        "known_labels": [
            "Entity Name:",
            "Continuity Contact Name:",
            "Continuity Contact Email:",
        ],
        "unknown_questions": [
            "Do you maintain alternate workspace arrangements for office disruptions?",
            "How do you prioritize manual workarounds during system outages?",
            "Describe your supply chain contingency planning process.",
            "How frequently do you validate emergency contact trees?",
            "Do you maintain backup communication channels for crisis coordination?",
            "How do you assess regional concentration risk for critical services?",
        ],
        "unknown_labels": [
            "Crisis Communications Lead:",
            "Emergency Hotline Owner:",
            "Alternate Site Coordinator:",
        ],
    },
    "vendor": {
        "category": "Vendor Management",
        "known_questions": [
            "Describe your third-party risk management process.",
            "How do you review and onboard critical vendors?",
            "Describe your vendor contract security review workflow.",
            "How do you monitor vendor performance and compliance?",
            "Describe your offboarding process for third-party service providers.",
            "How do you assess vendor incidents that affect customer services?",
        ],
        "known_labels": [
            "Vendor Legal Name:",
            "Primary Contact Name:",
            "Primary Contact Email:",
        ],
        "unknown_questions": [
            "Do you require independent assurance reports from high-risk vendors?",
            "How do you review vendor concentration risk?",
            "Describe your process for tracking vendor data return certificates.",
            "How do you evaluate fourth-party risk disclosures?",
            "Do you maintain minimum security control baselines for vendors?",
            "How do you handle vendors that refuse standard contractual terms?",
        ],
        "unknown_labels": [
            "Procurement Risk Owner:",
            "Fourth-Party Review Contact:",
            "Vendor Exit Coordinator:",
        ],
    },
}


LAYOUTS: list[dict[str, str]] = [
    {
        "slug": "balanced_two_column_mix",
        "label": "Balanced two-column table",
        "profile": "default",
    },
    {
        "slug": "two_column_header_mix",
        "label": "Two-column table with header row",
        "profile": "default",
    },
    {
        "slug": "multi_section_unknown_heavy",
        "label": "Multiple section tables with heavier unknown mix",
        "profile": "default",
    },
    {
        "slug": "strict_two_column_known_heavy",
        "label": "Compact two-column table with heavier known mix",
        "profile": "strict_two_column",
    },
    {
        "slug": "three_column_middle_mix",
        "label": "Three-column table with question in the middle column",
        "profile": "three_column_table",
    },
    {
        "slug": "four_column_metadata_mix",
        "label": "Four-column metadata table",
        "profile": "default",
    },
    {
        "slug": "merged_row_block_mix",
        "label": "Merged row-block questionnaire",
        "profile": "row_block_questionnaire",
    },
    {
        "slug": "sectioned_row_block_unknown_heavy",
        "label": "Sectioned row-block questionnaire with more unknown prompts",
        "profile": "row_block_questionnaire",
    },
    {
        "slug": "numbered_paragraph_unknown_heavy",
        "label": "Numbered paragraph questionnaire",
        "profile": "default",
    },
    {
        "slug": "mixed_layout_partial_kb",
        "label": "Mixed table, paragraph, and row-block layout",
        "profile": "default",
    },
]


def rotate_select(items: list[str], count: int, offset: int) -> list[str]:
    """Select items cyclically so prompts vary across documents."""

    return [items[(offset + index) % len(items)] for index in range(count)]


def make_entries(values: list[str], status: str) -> list[Entry]:
    return [Entry(text=value, status=status) for value in values]


def weave(*groups: list[Entry]) -> list[Entry]:
    """Interleave prompt groups while preserving group order."""

    buckets = [list(group) for group in groups]
    items: list[Entry] = []
    while any(buckets):
        for bucket in buckets:
            if bucket:
                items.append(bucket.pop(0))
    return items


def build_question_entries(theme: dict[str, object], *, known_count: int, unknown_count: int, offset: int) -> list[Entry]:
    known = make_entries(rotate_select(list(theme["known_questions"]), known_count, offset), "known")
    unknown = make_entries(rotate_select(list(theme["unknown_questions"]), unknown_count, offset * 2), "unknown")
    return weave(known, unknown)


def build_label_entries(theme: dict[str, object], *, known_count: int, unknown_count: int, offset: int) -> list[Entry]:
    known = make_entries(rotate_select(list(theme["known_labels"]), known_count, offset), "known")
    unknown = make_entries(rotate_select(list(theme["unknown_labels"]), unknown_count, offset), "unknown")
    return weave(known, unknown)


def add_title(doc: DocxDocument, title: str, subtitle: str) -> None:
    doc.add_heading(title, level=0)
    paragraph = doc.add_paragraph(subtitle)
    paragraph.style = "Intense Quote"


def render_two_column(doc: DocxDocument, entries: list[Entry], *, include_header: bool) -> None:
    rows = len(entries) + (1 if include_header else 0)
    table = doc.add_table(rows=rows, cols=2)
    table.style = "Table Grid"

    start_row = 0
    if include_header:
        table.rows[0].cells[0].text = "Question"
        table.rows[0].cells[1].text = "Answer"
        start_row = 1

    for index, entry in enumerate(entries, start=start_row):
        table.rows[index].cells[0].text = entry.text
        table.rows[index].cells[1].text = ""


def render_multi_section(doc: DocxDocument, entries: list[Entry], category: str) -> None:
    split = len(entries) // 2
    sections = [
        ("Core Controls", entries[:split]),
        (f"{category} Follow-up", entries[split:]),
    ]
    for section_title, section_entries in sections:
        doc.add_heading(section_title, level=1)
        table = doc.add_table(rows=len(section_entries) + 1, cols=2)
        table.style = "Table Grid"
        table.rows[0].cells[0].text = "Question"
        table.rows[0].cells[1].text = "Response"
        for index, entry in enumerate(section_entries, start=1):
            table.rows[index].cells[0].text = entry.text
            table.rows[index].cells[1].text = ""


def render_three_column(doc: DocxDocument, entries: list[Entry], category: str) -> None:
    table = doc.add_table(rows=len(entries) + 1, cols=3)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Domain"
    table.rows[0].cells[1].text = "Question"
    table.rows[0].cells[2].text = "Response"
    for index, entry in enumerate(entries, start=1):
        table.rows[index].cells[0].text = category
        table.rows[index].cells[1].text = entry.text
        table.rows[index].cells[2].text = ""


def render_four_column(doc: DocxDocument, entries: list[Entry], theme_key: str) -> None:
    code_prefix = theme_key[:3].upper()
    table = doc.add_table(rows=len(entries) + 1, cols=4)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Domain"
    table.rows[0].cells[1].text = "Reference"
    table.rows[0].cells[2].text = "Question"
    table.rows[0].cells[3].text = "Response"
    for index, entry in enumerate(entries, start=1):
        table.rows[index].cells[0].text = theme_key.title()
        table.rows[index].cells[1].text = f"{code_prefix}-{index:02d}"
        table.rows[index].cells[2].text = entry.text
        table.rows[index].cells[3].text = ""


def render_row_block(doc: DocxDocument, entries: list[Entry], *, sectioned: bool) -> None:
    if sectioned:
        split = len(entries) // 2
        groups = [
            ("Section A", entries[:split]),
            ("Section B", entries[split:]),
        ]
    else:
        groups = [("Questionnaire", entries)]

    for group_title, group_entries in groups:
        doc.add_heading(group_title, level=1)
        table = doc.add_table(rows=0, cols=2)
        table.style = "Table Grid"
        for entry in group_entries:
            question_row = table.add_row()
            merged = question_row.cells[0].merge(question_row.cells[1])
            merged.text = entry.text
            answer_row = table.add_row()
            answer_row.cells[0].text = ""
            answer_row.cells[1].text = ""


def render_numbered_paragraphs(doc: DocxDocument, entries: list[Entry]) -> None:
    for index, entry in enumerate(entries, start=1):
        doc.add_paragraph(f"{index}. {entry.text}")
        doc.add_paragraph("")


def render_mixed_layout(doc: DocxDocument, category: str, table_entries: list[Entry], paragraph_entries: list[Entry], row_block_entries: list[Entry]) -> None:
    doc.add_heading(f"{category} Quick Checks", level=1)
    render_two_column(doc, table_entries, include_header=True)

    doc.add_heading("Narrative Responses", level=1)
    for index, entry in enumerate(paragraph_entries, start=1):
        doc.add_paragraph(f"{index}. {entry.text}")
        doc.add_paragraph("")

    doc.add_heading("Follow-up Rows", level=1)
    render_row_block(doc, row_block_entries, sectioned=False)


def build_layout_entries(theme: dict[str, object], layout_index: int) -> dict[str, list[Entry]]:
    """Return prompt groups for the selected layout family."""

    offset = layout_index - 1

    if layout_index == 1:
        return {"main": build_question_entries(theme, known_count=3, unknown_count=3, offset=offset)}
    if layout_index == 2:
        return {"main": build_question_entries(theme, known_count=3, unknown_count=3, offset=offset + 1)}
    if layout_index == 3:
        return {"main": build_question_entries(theme, known_count=2, unknown_count=4, offset=offset + 2)}
    if layout_index == 4:
        return {"main": build_question_entries(theme, known_count=4, unknown_count=2, offset=offset + 1)}
    if layout_index == 5:
        main = build_question_entries(theme, known_count=3, unknown_count=2, offset=offset)
        labels = build_label_entries(theme, known_count=1, unknown_count=0, offset=offset)
        return {"main": weave(labels, main)}
    if layout_index == 6:
        labels = build_label_entries(theme, known_count=1, unknown_count=1, offset=offset)
        questions = build_question_entries(theme, known_count=2, unknown_count=3, offset=offset)
        return {"main": weave(labels, questions)}
    if layout_index == 7:
        return {"main": build_question_entries(theme, known_count=3, unknown_count=3, offset=offset + 3)}
    if layout_index == 8:
        return {"main": build_question_entries(theme, known_count=2, unknown_count=4, offset=offset + 4)}
    if layout_index == 9:
        return {"main": build_question_entries(theme, known_count=2, unknown_count=4, offset=offset + 5)}

    known_pool = rotate_select(list(theme["known_questions"]), 4, offset)
    unknown_pool = rotate_select(list(theme["unknown_questions"]), 4, offset + 1)
    table_entries = weave(
        make_entries(known_pool[:2], "known"),
        make_entries(unknown_pool[:1], "unknown"),
    )
    paragraph_entries = weave(
        make_entries(known_pool[2:3], "known"),
        make_entries(unknown_pool[1:3], "unknown"),
    )
    row_block_entries = weave(
        make_entries(known_pool[3:4], "known"),
        make_entries(unknown_pool[3:4], "unknown"),
    )
    return {
        "table": table_entries,
        "paragraph": paragraph_entries,
        "row_block": row_block_entries,
    }


def count_status(entries: dict[str, list[Entry]]) -> tuple[int, int]:
    items = [entry for group in entries.values() for entry in group]
    known = sum(1 for entry in items if entry.status == "known")
    unknown = sum(1 for entry in items if entry.status == "unknown")
    return known, unknown


def render_document(theme_key: str, theme: dict[str, object], layout_index: int, output_path: Path) -> tuple[int, int]:
    category = str(theme["category"])
    entries = build_layout_entries(theme, layout_index)
    known_count, unknown_count = count_status(entries)

    doc = DocxDocument()
    title = f"{category} Stress Test Questionnaire"
    subtitle = "This document intentionally mixes prompts that exist in the sample knowledge base with prompts that do not."
    add_title(doc, title, subtitle)

    if layout_index == 1:
        render_two_column(doc, entries["main"], include_header=False)
    elif layout_index == 2:
        render_two_column(doc, entries["main"], include_header=True)
    elif layout_index == 3:
        render_multi_section(doc, entries["main"], category)
    elif layout_index == 4:
        render_two_column(doc, entries["main"], include_header=False)
    elif layout_index == 5:
        render_three_column(doc, entries["main"], category)
    elif layout_index == 6:
        render_four_column(doc, entries["main"], theme_key)
    elif layout_index == 7:
        render_row_block(doc, entries["main"], sectioned=False)
    elif layout_index == 8:
        render_row_block(doc, entries["main"], sectioned=True)
    elif layout_index == 9:
        render_numbered_paragraphs(doc, entries["main"])
    else:
        render_mixed_layout(doc, category, entries["table"], entries["paragraph"], entries["row_block"])

    doc.save(str(output_path))
    return known_count, unknown_count


def write_manifest(rows: list[dict[str, object]]) -> None:
    with MANIFEST_PATH.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.DictWriter(
            handle,
            fieldnames=[
                "file",
                "theme",
                "layout_family",
                "suggested_profile",
                "total_prompts",
                "expected_known_matches",
                "expected_unknown_flags",
            ],
        )
        writer.writeheader()
        writer.writerows(rows)


def write_readme(rows: list[dict[str, object]]) -> None:
    lines = [
        "# Mixed Coverage Stress-Test Corpus",
        "",
        "This folder contains 50 generated `.docx` questionnaire files for upload testing.",
        "",
        "What this corpus is designed to test:",
        "- A mix of prompts that already exist in the sample knowledge base from `test-data/Test-1`",
        "- A mix of prompts that are intentionally missing and should be flagged for review",
        "- Multiple supported document layouts to stress parser selection and write-back behavior",
        "",
        "Important note:",
        "- The `expected_known_matches` counts assume you still have the `Test-1` sample knowledge base loaded into the app.",
        "- The `expected_unknown_flags` counts are deliberate gaps and should remain unresolved unless you add more Q&A pairs.",
        "",
        "Suggested parser profiles:",
        "- `default`: standard two-column, paragraph, and mixed layouts",
        "- `strict_two_column`: compact two-column layouts with no extra metadata columns",
        "- `three_column_table`: three-column tables where the question is in the middle column",
        "- `row_block_questionnaire`: merged question rows with blank answer rows below",
        "",
        "See `manifest.csv` for a machine-readable summary.",
        "",
        "| File | Layout Family | Theme | Suggested Profile | Expected Known | Expected Unknown |",
        "|---|---|---|---|---:|---:|",
    ]

    for row in rows:
        lines.append(
            f"| `{row['file']}` | {row['layout_family']} | {row['theme']} | `{row['suggested_profile']}` | {row['expected_known_matches']} | {row['expected_unknown_flags']} |"
        )

    README_PATH.write_text("\n".join(lines) + "\n", encoding="utf-8")


def main() -> None:
    if OUTPUT_DIR.exists():
        shutil.rmtree(OUTPUT_DIR)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    manifest_rows: list[dict[str, object]] = []
    theme_order = ["security", "privacy", "operations", "continuity", "vendor"]

    file_number = 1
    for theme_key in theme_order:
        theme = THEMES[theme_key]
        for layout_index, layout in enumerate(LAYOUTS, start=1):
            filename = f"{file_number:02d}_{layout['slug']}_{theme_key}.docx"
            output_path = OUTPUT_DIR / filename
            known_count, unknown_count = render_document(theme_key, theme, layout_index, output_path)

            manifest_rows.append(
                {
                    "file": filename,
                    "theme": str(theme["category"]),
                    "layout_family": layout["label"],
                    "suggested_profile": layout["profile"],
                    "total_prompts": known_count + unknown_count,
                    "expected_known_matches": known_count,
                    "expected_unknown_flags": unknown_count,
                }
            )
            file_number += 1

    write_manifest(manifest_rows)
    write_readme(manifest_rows)
    print(f"Generated {len(manifest_rows)} documents in {OUTPUT_DIR}")


if __name__ == "__main__":
    main()
