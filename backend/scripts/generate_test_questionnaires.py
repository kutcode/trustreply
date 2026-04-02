"""Generate diverse test questionnaires in DOCX and PDF formats.

Usage:
    python scripts/generate_test_questionnaires.py [output_dir]

Creates ~30 files with varied layouts, topics, and formats for stress-testing.
"""

from __future__ import annotations

import os
import sys
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF

OUTPUT_DIR = Path(sys.argv[1]) if len(sys.argv) > 1 else Path(__file__).resolve().parent.parent.parent / "test_questionnaires"

# ─── Question banks by topic ────────────────────────────────────────

SECURITY_QS = [
    ("Do you encrypt data at rest?", "Yes, all data at rest is encrypted using AES-256 encryption."),
    ("Do you encrypt data in transit?", "Yes, all data in transit is encrypted using TLS 1.2 or higher."),
    ("Do you have a vulnerability management program?", "Yes, we perform quarterly vulnerability scans and annual penetration testing."),
    ("How do you handle security incidents?", "We have a documented incident response plan that is tested annually and includes 24/7 on-call rotation."),
    ("Do you perform background checks on employees?", "Yes, all employees undergo background checks prior to employment."),
    ("What authentication mechanisms do you support?", "We support SSO via SAML 2.0 and OIDC, MFA via TOTP and hardware keys, and password-based auth with complexity requirements."),
    ("Do you have a SOC 2 Type II report?", "Yes, our most recent SOC 2 Type II report covers the period from January to December 2025."),
    ("How do you manage access control?", "We implement role-based access control (RBAC) with least-privilege principles and quarterly access reviews."),
    ("Do you have a data retention policy?", "Yes, data is retained for the minimum period required by contract and applicable regulations, then securely deleted."),
    ("What is your disaster recovery RTO/RPO?", "Our RTO is 4 hours and RPO is 1 hour for critical systems."),
    ("Do you conduct security awareness training?", "Yes, all employees complete security awareness training upon hire and annually thereafter."),
    ("How do you secure your development lifecycle?", "We follow secure SDLC practices including code reviews, SAST/DAST scanning, and dependency vulnerability checks."),
]

COMPLIANCE_QS = [
    ("Are you GDPR compliant?", "Yes, we are fully GDPR compliant and have appointed a Data Protection Officer."),
    ("Do you comply with CCPA?", "Yes, we comply with CCPA and honor all consumer data rights requests within the required timeframe."),
    ("Are you HIPAA compliant?", "Yes, we maintain HIPAA compliance and can execute Business Associate Agreements (BAAs)."),
    ("Do you have ISO 27001 certification?", "Yes, our information security management system is ISO 27001:2022 certified."),
    ("How do you handle data subject access requests?", "We have an automated workflow for processing DSARs within the 30-day GDPR deadline."),
    ("Do you comply with PCI DSS?", "Yes, we maintain PCI DSS Level 1 compliance, validated by an annual QSA audit."),
    ("What is your data processing legal basis?", "We process data based on contractual necessity, legitimate interest, and explicit consent where required."),
    ("Do you conduct privacy impact assessments?", "Yes, we conduct Data Protection Impact Assessments for all new processing activities involving personal data."),
    ("How do you handle cross-border data transfers?", "We use Standard Contractual Clauses (SCCs) and ensure adequate safeguards for international data transfers."),
    ("Do you have a privacy policy?", "Yes, our privacy policy is publicly available and updated at least annually."),
]

VENDOR_QS = [
    ("What is your company's financial stability?", "We have been profitable for the last 5 years with year-over-year revenue growth exceeding 20%."),
    ("Do you have business continuity plans?", "Yes, our BCP is tested annually and covers all critical business functions."),
    ("What are your SLA commitments?", "We offer 99.95% uptime SLA with financial credits for any downtime below the threshold."),
    ("Do you carry cyber insurance?", "Yes, we maintain cyber liability insurance with coverage of $10M per occurrence."),
    ("How many employees do you have?", "We have approximately 500 employees across 3 offices globally."),
    ("What is your customer support model?", "We provide 24/7 support via email, chat, and phone with dedicated account managers for enterprise clients."),
    ("Do you use subprocessors?", "Yes, we maintain a list of subprocessors that is available upon request and updated with 30-day advance notice."),
    ("What is your change management process?", "All changes follow an ITIL-based change management process with CAB review for significant changes."),
]

INFRASTRUCTURE_QS = [
    ("Where is your data hosted?", "Our primary infrastructure is hosted on AWS in us-east-1 and us-west-2 regions with failover capabilities."),
    ("Do you use multi-tenant or single-tenant architecture?", "We offer both options. Our standard offering is multi-tenant with logical data isolation."),
    ("How do you handle backups?", "We perform automated daily backups with 30-day retention, stored in a geographically separate region."),
    ("What is your network security architecture?", "We use VPCs with network segmentation, WAF, DDoS protection, and IDS/IPS systems."),
    ("Do you monitor your systems 24/7?", "Yes, we have a 24/7 NOC/SOC with automated alerting and escalation procedures."),
    ("How do you handle patch management?", "Critical patches are applied within 24 hours, high within 7 days, and medium within 30 days."),
    ("Do you have redundant systems?", "Yes, all critical systems are deployed in an active-active configuration across multiple availability zones."),
    ("What logging and monitoring tools do you use?", "We use a SIEM platform for centralized logging with 12-month retention and real-time alerting."),
]

AI_GOVERNANCE_QS = [
    ("Do you use AI/ML in your product?", "Yes, we use machine learning for anomaly detection and natural language processing features."),
    ("How do you ensure AI fairness?", "We conduct regular bias audits and use diverse training datasets with fairness metrics tracking."),
    ("Do you have an AI ethics policy?", "Yes, we have a published AI ethics policy reviewed by our ethics advisory board."),
    ("How do you handle AI model transparency?", "We provide model cards for all production models documenting capabilities, limitations, and performance metrics."),
    ("What data is used to train your models?", "Models are trained on anonymized, aggregated data with explicit consent and opt-out mechanisms."),
    ("Do you perform AI risk assessments?", "Yes, all AI systems undergo risk assessments aligned with the NIST AI RMF before deployment."),
]

THIRD_PARTY_QS = [
    ("How do you assess third-party risk?", "We conduct annual vendor risk assessments including security questionnaires, SOC 2 reviews, and on-site audits for critical vendors."),
    ("Do you have a vendor management program?", "Yes, our vendor management program includes risk tiering, ongoing monitoring, and contractual security requirements."),
    ("How many third-party vendors have access to customer data?", "Approximately 12 vendors have access to customer data, each under strict DPAs and security requirements."),
    ("Do you require vendors to carry insurance?", "Yes, critical vendors must maintain cyber liability insurance with minimum $5M coverage."),
]


# ─── DOCX generators ────────────────────────────────────────────────

def _style_heading(doc: Document, text: str, level: int = 1):
    h = doc.add_heading(text, level=level)
    return h


def make_table_qa(filename: str, title: str, questions: list[tuple[str, str]], has_header: bool = True):
    """Classic table format: Question | Answer columns."""
    doc = Document()
    _style_heading(doc, title)
    doc.add_paragraph(f"Vendor: [Company Name]  |  Date: 2026-03-31  |  Version: 1.0")
    doc.add_paragraph("")

    table = doc.add_table(rows=1 if has_header else 0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    if has_header:
        hdr = table.rows[0].cells
        hdr[0].text = "Question"
        hdr[1].text = "Response"
        for cell in hdr:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True

    for q, _ in questions:
        row = table.add_row().cells
        row[0].text = q
        row[1].text = ""  # leave blank for filling

    doc.save(str(OUTPUT_DIR / filename))


def make_table_qa_prefilled(filename: str, title: str, questions: list[tuple[str, str]]):
    """Table format with some answers pre-filled (for testing partial fills)."""
    doc = Document()
    _style_heading(doc, title)
    doc.add_paragraph("Please review and complete any blank responses.")
    doc.add_paragraph("")

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "#"
    hdr[1].text = "Question"
    hdr[2].text = "Answer"
    for cell in hdr:
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True

    for i, (q, a) in enumerate(questions):
        row = table.add_row().cells
        row[0].text = str(i + 1)
        row[1].text = q
        row[2].text = a if i % 3 == 0 else ""  # every 3rd pre-filled

    doc.save(str(OUTPUT_DIR / filename))


def make_numbered_list(filename: str, title: str, questions: list[tuple[str, str]]):
    """Numbered list format with blank lines for answers."""
    doc = Document()
    _style_heading(doc, title)
    doc.add_paragraph(f"Organization: _______________    Date: _______________")
    doc.add_paragraph("")

    for i, (q, _) in enumerate(questions):
        doc.add_paragraph(f"{i + 1}. {q}", style="List Number")
        doc.add_paragraph("Answer: ", style="Normal")
        doc.add_paragraph("")

    doc.save(str(OUTPUT_DIR / filename))


def make_sectioned(filename: str, title: str, sections: dict[str, list[tuple[str, str]]]):
    """Multi-section format with headings and sub-questions."""
    doc = Document()
    _style_heading(doc, title)
    doc.add_paragraph("Complete all sections below. Provide detailed responses where applicable.")
    doc.add_paragraph("")

    for section_name, qs in sections.items():
        _style_heading(doc, section_name, level=2)
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Requirement"
        hdr[1].text = "Vendor Response"
        for cell in hdr:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True

        for q, _ in qs:
            row = table.add_row().cells
            row[0].text = q
            row[1].text = ""

        doc.add_paragraph("")

    doc.save(str(OUTPUT_DIR / filename))


def make_yes_no_format(filename: str, title: str, questions: list[tuple[str, str]]):
    """Yes/No/NA table with comments column."""
    doc = Document()
    _style_heading(doc, title)
    doc.add_paragraph("Mark Yes, No, or N/A and provide additional comments where needed.")
    doc.add_paragraph("")

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "#"
    hdr[1].text = "Control Question"
    hdr[2].text = "Yes/No/NA"
    hdr[3].text = "Comments / Evidence"
    for cell in hdr:
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True

    for i, (q, _) in enumerate(questions):
        row = table.add_row().cells
        row[0].text = str(i + 1)
        row[1].text = q
        row[2].text = ""
        row[3].text = ""

    doc.save(str(OUTPUT_DIR / filename))


def make_free_text(filename: str, title: str, questions: list[tuple[str, str]]):
    """Free-text paragraph format (no tables)."""
    doc = Document()
    _style_heading(doc, title)
    doc.add_paragraph("")

    for i, (q, _) in enumerate(questions):
        p = doc.add_paragraph()
        run = p.add_run(f"Q{i + 1}: {q}")
        run.bold = True
        doc.add_paragraph("A: ")
        doc.add_paragraph("")

    doc.save(str(OUTPUT_DIR / filename))


def make_multi_table_sections(filename: str, title: str, sections: dict[str, list[tuple[str, str]]]):
    """Multiple separate tables, one per section, with different widths."""
    doc = Document()
    _style_heading(doc, title)
    doc.add_paragraph(f"Completed by: _______________    Role: _______________")
    doc.add_paragraph("")

    for idx, (section, qs) in enumerate(sections.items()):
        _style_heading(doc, f"Section {idx + 1}: {section}", level=2)

        if idx % 2 == 0:
            # 2-column table
            table = doc.add_table(rows=1, cols=2)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            hdr[0].text = "Question"
            hdr[1].text = "Response"
        else:
            # 3-column table with ref number
            table = doc.add_table(rows=1, cols=3)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            hdr[0].text = "Ref"
            hdr[1].text = "Requirement"
            hdr[2].text = "Response"

        for cell in hdr:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True

        for j, (q, _) in enumerate(qs):
            row = table.add_row().cells
            if idx % 2 == 0:
                row[0].text = q
                row[1].text = ""
            else:
                row[0].text = f"{idx + 1}.{j + 1}"
                row[1].text = q
                row[2].text = ""

        doc.add_paragraph("")

    doc.save(str(OUTPUT_DIR / filename))


def make_rfi_style(filename: str, title: str, questions: list[tuple[str, str]]):
    """RFI/RFP style with category codes and weighted scoring columns."""
    doc = Document()
    _style_heading(doc, title)
    doc.add_paragraph("Please respond to each item. Responses will be scored based on completeness.")
    doc.add_paragraph("")

    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "ID"
    hdr[1].text = "Category"
    hdr[2].text = "Question"
    hdr[3].text = "Response"
    hdr[4].text = "Score (1-5)"
    for cell in hdr:
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True

    categories = ["Security", "Compliance", "Operations", "Technical"]
    for i, (q, _) in enumerate(questions):
        row = table.add_row().cells
        row[0].text = f"RFI-{i + 1:03d}"
        row[1].text = categories[i % len(categories)]
        row[2].text = q
        row[3].text = ""
        row[4].text = ""

    doc.save(str(OUTPUT_DIR / filename))


def make_matrix_format(filename: str, title: str, questions: list[tuple[str, str]]):
    """Matrix/checklist format with multiple assessment columns."""
    doc = Document()
    _style_heading(doc, title)
    doc.add_paragraph("For each control, indicate implementation status and provide evidence.")
    doc.add_paragraph("")

    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Control"
    hdr[1].text = "Fully\nImplemented"
    hdr[2].text = "Partially\nImplemented"
    hdr[3].text = "Not\nImplemented"
    hdr[4].text = "Evidence / Notes"
    for cell in hdr:
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)

    for q, _ in questions:
        row = table.add_row().cells
        row[0].text = q
        row[1].text = ""
        row[2].text = ""
        row[3].text = ""
        row[4].text = ""

    doc.save(str(OUTPUT_DIR / filename))


# ─── PDF generators ─────────────────────────────────────────────────

class QuestionnairePDF(FPDF):
    def header(self):
        self.set_font("Helvetica", "B", 14)
        self.cell(0, 10, self._title_text, new_x="LMARGIN", new_y="NEXT", align="C")
        self.ln(5)

    def footer(self):
        self.set_y(-15)
        self.set_font("Helvetica", "I", 8)
        self.cell(0, 10, f"Page {self.page_no()}/{{nb}}", align="C")


def make_pdf_table(filename: str, title: str, questions: list[tuple[str, str]]):
    """PDF with table layout."""
    pdf = QuestionnairePDF()
    pdf._title_text = title
    pdf.alias_nb_pages()
    pdf.add_page()
    pdf.set_font("Helvetica", "", 10)

    pdf.cell(0, 8, f"Date: 2026-03-31    Organization: _______________", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(5)

    # Header row
    pdf.set_font("Helvetica", "B", 10)
    pdf.set_fill_color(50, 50, 80)
    pdf.set_text_color(255, 255, 255)
    pdf.cell(10, 8, "#", border=1, fill=True, align="C")
    pdf.cell(90, 8, "Question", border=1, fill=True)
    pdf.cell(90, 8, "Response", border=1, fill=True, new_x="LMARGIN", new_y="NEXT")

    pdf.set_text_color(0, 0, 0)
    pdf.set_font("Helvetica", "", 9)

    for i, (q, _) in enumerate(questions):
        row_h = max(12, len(q) // 40 * 6 + 12)
        pdf.cell(10, row_h, str(i + 1), border=1, align="C")
        pdf.cell(90, row_h, q[:80], border=1)
        pdf.cell(90, row_h, "", border=1, new_x="LMARGIN", new_y="NEXT")

    pdf.output(str(OUTPUT_DIR / filename))


def make_pdf_list(filename: str, title: str, questions: list[tuple[str, str]]):
    """PDF with numbered list format."""
    pdf = QuestionnairePDF()
    pdf._title_text = title
    pdf.alias_nb_pages()
    pdf.add_page()

    pdf.set_font("Helvetica", "I", 10)
    pdf.cell(0, 8, "Please provide detailed responses to each question below.", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(5)

    for i, (q, _) in enumerate(questions):
        pdf.set_font("Helvetica", "B", 10)
        pdf.multi_cell(0, 6, f"{i + 1}. {q}", new_x="LMARGIN", new_y="NEXT")
        pdf.set_font("Helvetica", "", 10)
        pdf.cell(0, 6, "Answer: _____________________________________________", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(4)

    pdf.output(str(OUTPUT_DIR / filename))


def make_pdf_sectioned(filename: str, title: str, sections: dict[str, list[tuple[str, str]]]):
    """PDF with multiple sections and tables."""
    pdf = QuestionnairePDF()
    pdf._title_text = title
    pdf.alias_nb_pages()
    pdf.add_page()

    for section, qs in sections.items():
        pdf.set_font("Helvetica", "B", 12)
        pdf.set_fill_color(240, 240, 245)
        pdf.cell(0, 10, section, new_x="LMARGIN", new_y="NEXT", fill=True)
        pdf.ln(3)

        pdf.set_font("Helvetica", "", 9)
        for i, (q, _) in enumerate(qs):
            pdf.set_font("Helvetica", "B", 9)
            pdf.multi_cell(0, 5, f"  {i + 1}. {q}", new_x="LMARGIN", new_y="NEXT")
            pdf.set_font("Helvetica", "", 9)
            pdf.cell(0, 5, "     Response: ", new_x="LMARGIN", new_y="NEXT")
            pdf.ln(3)

        pdf.ln(5)

    pdf.output(str(OUTPUT_DIR / filename))


def make_pdf_two_column(filename: str, title: str, questions: list[tuple[str, str]]):
    """PDF with a two-column layout."""
    pdf = QuestionnairePDF()
    pdf._title_text = title
    pdf.alias_nb_pages()
    pdf.add_page()
    pdf.set_font("Helvetica", "", 9)

    col_w = 92
    x_start = [10, 105]

    for i, (q, _) in enumerate(questions):
        col = i % 2
        if col == 0 and i > 0:
            pdf.ln(15)

        x = x_start[col]
        y = pdf.get_y()
        pdf.set_xy(x, y)
        pdf.set_font("Helvetica", "B", 9)
        pdf.cell(col_w, 5, f"Q{i + 1}: {q[:60]}", new_x="LMARGIN", new_y="NEXT")
        pdf.set_xy(x, pdf.get_y())
        pdf.set_font("Helvetica", "", 9)
        pdf.cell(col_w, 5, "A: ____________________", new_x="LMARGIN", new_y="NEXT")

    pdf.output(str(OUTPUT_DIR / filename))


def make_pdf_yes_no(filename: str, title: str, questions: list[tuple[str, str]]):
    """PDF with Yes/No checkbox style."""
    pdf = QuestionnairePDF()
    pdf._title_text = title
    pdf.alias_nb_pages()
    pdf.add_page()

    pdf.set_font("Helvetica", "B", 10)
    pdf.set_fill_color(50, 50, 80)
    pdf.set_text_color(255, 255, 255)
    pdf.cell(10, 8, "#", border=1, fill=True, align="C")
    pdf.cell(100, 8, "Question", border=1, fill=True)
    pdf.cell(15, 8, "Yes", border=1, fill=True, align="C")
    pdf.cell(15, 8, "No", border=1, fill=True, align="C")
    pdf.cell(15, 8, "N/A", border=1, fill=True, align="C")
    pdf.cell(35, 8, "Comments", border=1, fill=True, new_x="LMARGIN", new_y="NEXT")

    pdf.set_text_color(0, 0, 0)
    pdf.set_font("Helvetica", "", 8)

    for i, (q, _) in enumerate(questions):
        fill = i % 2 == 0
        if fill:
            pdf.set_fill_color(245, 245, 250)
        pdf.cell(10, 10, str(i + 1), border=1, align="C", fill=fill)
        pdf.cell(100, 10, q[:70], border=1, fill=fill)
        pdf.cell(15, 10, "", border=1, align="C", fill=fill)
        pdf.cell(15, 10, "", border=1, align="C", fill=fill)
        pdf.cell(15, 10, "", border=1, align="C", fill=fill)
        pdf.cell(35, 10, "", border=1, fill=fill, new_x="LMARGIN", new_y="NEXT")

    pdf.output(str(OUTPUT_DIR / filename))


# ─── Generate all test files ────────────────────────────────────────

def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    print(f"Generating test questionnaires in {OUTPUT_DIR} ...\n")

    # ── DOCX files ──

    # 1. Basic security table
    make_table_qa("01_security_assessment.docx", "Information Security Assessment", SECURITY_QS)

    # 2. Compliance table with 3 columns
    make_table_qa_prefilled("02_compliance_review.docx", "Regulatory Compliance Questionnaire", COMPLIANCE_QS)

    # 3. Vendor risk numbered list
    make_numbered_list("03_vendor_risk_questionnaire.docx", "Third-Party Vendor Risk Questionnaire", VENDOR_QS)

    # 4. Multi-section security + compliance
    make_sectioned("04_comprehensive_security_review.docx", "Comprehensive Security & Compliance Review", {
        "Access Control & Authentication": SECURITY_QS[:4],
        "Data Protection": SECURITY_QS[4:8],
        "Compliance & Regulatory": COMPLIANCE_QS[:5],
        "Business Continuity": VENDOR_QS[:4],
    })

    # 5. Yes/No format
    make_yes_no_format("05_security_controls_checklist.docx", "Security Controls Assessment Checklist", SECURITY_QS[:10])

    # 6. Free-text format
    make_free_text("06_privacy_impact_assessment.docx", "Privacy Impact Assessment Questionnaire", COMPLIANCE_QS[:8])

    # 7. Multi-table sections
    make_multi_table_sections("07_vendor_due_diligence.docx", "Vendor Due Diligence Package", {
        "Company Information": VENDOR_QS[:4],
        "Technical Security": SECURITY_QS[:4],
        "Data Handling": COMPLIANCE_QS[:4],
    })

    # 8. RFI style
    make_rfi_style("08_rfp_security_requirements.docx", "RFP — Security Requirements Response", SECURITY_QS + COMPLIANCE_QS[:4])

    # 9. Matrix format
    make_matrix_format("09_control_assessment_matrix.docx", "Control Assessment Matrix — ISO 27001", SECURITY_QS[:8])

    # 10. Infrastructure focused
    make_table_qa("10_infrastructure_questionnaire.docx", "Cloud Infrastructure Security Questionnaire", INFRASTRUCTURE_QS)

    # 11. AI governance
    make_table_qa("11_ai_governance_assessment.docx", "AI / ML Governance Assessment", AI_GOVERNANCE_QS)

    # 12. Third-party risk
    make_yes_no_format("12_third_party_risk_assessment.docx", "Third-Party Risk Assessment Form", THIRD_PARTY_QS + VENDOR_QS[:4])

    # 13. Large comprehensive (all topics)
    make_sectioned("13_enterprise_security_questionnaire.docx", "Enterprise Security Questionnaire (Full)", {
        "Information Security": SECURITY_QS,
        "Regulatory Compliance": COMPLIANCE_QS,
        "Vendor Management": VENDOR_QS,
        "Infrastructure": INFRASTRUCTURE_QS,
        "AI Governance": AI_GOVERNANCE_QS,
    })

    # 14. Short/minimal (3 questions)
    make_table_qa("14_quick_security_check.docx", "Quick Security Verification", SECURITY_QS[:3])

    # 15. No header row table
    make_table_qa("15_no_header_table.docx", "Security Questionnaire (Simple)", SECURITY_QS[:6], has_header=False)

    # 16. Mixed format
    make_free_text("16_incident_response_assessment.docx", "Incident Response Readiness Assessment", SECURITY_QS[3:7] + INFRASTRUCTURE_QS[4:7])

    # 17. Numbered compliance
    make_numbered_list("17_gdpr_compliance_checklist.docx", "GDPR Compliance Verification Checklist", COMPLIANCE_QS)

    # 18. RFI with infrastructure
    make_rfi_style("18_cloud_services_rfp.docx", "Cloud Services RFP — Technical Requirements", INFRASTRUCTURE_QS + SECURITY_QS[:4])

    # 19. Matrix for compliance
    make_matrix_format("19_compliance_maturity_assessment.docx", "Compliance Maturity Assessment", COMPLIANCE_QS[:8])

    # 20. Prefilled vendor assessment
    make_table_qa_prefilled("20_vendor_self_assessment.docx", "Vendor Self-Assessment (Partially Completed)", VENDOR_QS + SECURITY_QS[:4])

    # ── PDF files ──

    # 21. PDF table
    make_pdf_table("21_security_questionnaire.pdf", "Security Questionnaire", SECURITY_QS)

    # 22. PDF list
    make_pdf_list("22_compliance_questionnaire.pdf", "Compliance Assessment Questionnaire", COMPLIANCE_QS)

    # 23. PDF sectioned
    make_pdf_sectioned("23_comprehensive_assessment.pdf", "Comprehensive Vendor Assessment", {
        "Security Controls": SECURITY_QS[:6],
        "Regulatory Compliance": COMPLIANCE_QS[:5],
        "Infrastructure": INFRASTRUCTURE_QS[:4],
    })

    # 24. PDF two-column
    make_pdf_two_column("24_quick_assessment.pdf", "Rapid Security Assessment", SECURITY_QS[:8])

    # 25. PDF yes/no
    make_pdf_yes_no("25_security_controls_checklist.pdf", "Security Controls Compliance Checklist", SECURITY_QS[:10])

    # 26. PDF large
    make_pdf_table("26_full_vendor_questionnaire.pdf", "Full Vendor Security Questionnaire", SECURITY_QS + COMPLIANCE_QS[:6])

    # 27. PDF infrastructure
    make_pdf_list("27_infrastructure_assessment.pdf", "Infrastructure Security Assessment", INFRASTRUCTURE_QS)

    # 28. PDF AI governance
    make_pdf_yes_no("28_ai_risk_assessment.pdf", "AI Risk Assessment Checklist", AI_GOVERNANCE_QS)

    # 29. PDF sectioned all topics
    make_pdf_sectioned("29_enterprise_assessment.pdf", "Enterprise-Wide Security Assessment", {
        "Information Security": SECURITY_QS[:6],
        "Compliance": COMPLIANCE_QS[:5],
        "Vendor Management": VENDOR_QS[:4],
        "Infrastructure": INFRASTRUCTURE_QS[:4],
        "AI Governance": AI_GOVERNANCE_QS[:3],
    })

    # 30. PDF minimal
    make_pdf_table("30_mini_security_check.pdf", "Quick Security Check", SECURITY_QS[:3])

    # List generated files
    files = sorted(OUTPUT_DIR.iterdir())
    print(f"Generated {len(files)} test questionnaires:\n")
    for f in files:
        size_kb = f.stat().st_size / 1024
        print(f"  {f.name:50s} {size_kb:6.1f} KB")

    print(f"\nAll files saved to: {OUTPUT_DIR}")


if __name__ == "__main__":
    main()
