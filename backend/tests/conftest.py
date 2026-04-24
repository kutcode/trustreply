"""Shared test fixtures."""

import pytest
import pytest_asyncio
import csv
from pathlib import Path
from httpx import AsyncClient, ASGITransport
from sqlalchemy.ext.asyncio import create_async_engine, async_sessionmaker, AsyncSession
from docx import Document as DocxDocument

import app.database as _db_module
import app.routers.upload as _upload_module
import app.routers.qa as _qa_module
import app.routers.flagged as _flagged_module
from app.database import Base, get_db
import app.models  # noqa: F401 — ensure all models are registered with Base.metadata before create_all
from app.main import app
from app.config import settings

TEST_DB_URL = "sqlite+aiosqlite://"  # in-memory


_SETTINGS_DEFAULTS = {
    # Agent defaults (LLM disabled during tests unless a test opts in)
    "agent_enabled": False,
    "agent_provider": "openai",
    "agent_api_base": "https://api.openai.com/v1",
    "agent_api_key": "",
    "agent_model": "gpt-4.1-nano",
    "agent_default_mode": "agent",
    # Per-provider keys are checked by the settings/models endpoint fallback;
    # clearing them keeps "no API key provided" paths reachable in tests.
    "agent_openai_api_key": "",
    "agent_anthropic_api_key": "",
    # Auth off: tests exercise endpoints directly without Supabase or API keys.
    # Without this, a local .env populated with auth values turns every request
    # into a 401 and 41 API-level tests fail on a fresh clone.
    "supabase_url": "",
    "supabase_anon_key": "",
    "supabase_jwt_secret": "",
    "api_key": "",
    "allowed_email_domains": [],
}


@pytest.fixture(autouse=True)
def isolate_runtime_settings():
    """Keep tests deterministic regardless of local .env overrides."""
    original = {key: getattr(settings, key) for key in _SETTINGS_DEFAULTS}
    for key, value in _SETTINGS_DEFAULTS.items():
        setattr(settings, key, value)
    try:
        yield
    finally:
        for key, value in original.items():
            setattr(settings, key, value)


@pytest_asyncio.fixture
async def db_engine():
    """Create a fresh in-memory database engine for each test."""
    engine = create_async_engine(TEST_DB_URL, echo=False)
    async with engine.begin() as conn:
        await conn.run_sync(Base.metadata.create_all)
    yield engine
    async with engine.begin() as conn:
        await conn.run_sync(Base.metadata.drop_all)
    await engine.dispose()


@pytest_asyncio.fixture
async def db_session(db_engine):
    """Yield a DB session bound to the test engine."""
    session_factory = async_sessionmaker(db_engine, class_=AsyncSession, expire_on_commit=False)
    async with session_factory() as session:
        yield session


@pytest_asyncio.fixture
async def client(db_engine, db_session):
    """HTTP test client with overridden DB dependency.

    Patches both the FastAPI dependency *and* the module-level engine/session
    so that background tasks (which use ``async_session()`` directly) also
    hit the test database.
    """
    async def override_get_db():
        yield db_session

    # Patch module-level engine + session factory so background tasks use the test DB
    test_session_factory = async_sessionmaker(db_engine, class_=AsyncSession, expire_on_commit=False)

    originals = {
        "db_engine": _db_module.engine,
        "db_session": _db_module.async_session,
        "upload_session": _upload_module.async_session,
        "qa_session": _qa_module.async_session,
        "flagged_session": _flagged_module.async_session,
    }
    _db_module.engine = db_engine
    _db_module.async_session = test_session_factory
    _upload_module.async_session = test_session_factory
    _qa_module.async_session = test_session_factory
    _flagged_module.async_session = test_session_factory

    app.dependency_overrides[get_db] = override_get_db
    transport = ASGITransport(app=app)
    async with AsyncClient(transport=transport, base_url="http://test") as c:
        yield c
    app.dependency_overrides.clear()

    # Restore originals
    _db_module.engine = originals["db_engine"]
    _db_module.async_session = originals["db_session"]
    _upload_module.async_session = originals["upload_session"]
    _qa_module.async_session = originals["qa_session"]
    _flagged_module.async_session = originals["flagged_session"]


@pytest.fixture
def make_docx(tmp_path):
    """Factory fixture to create test .docx files with table-based questions."""
    def _make(questions_and_answers: list[tuple[str, str]], filename: str = "test.docx") -> Path:
        doc = DocxDocument()
        if questions_and_answers:
            table = doc.add_table(rows=len(questions_and_answers), cols=2)
            for i, (q, a) in enumerate(questions_and_answers):
                table.rows[i].cells[0].text = q
                table.rows[i].cells[1].text = a
        path = tmp_path / filename
        doc.save(str(path))
        return path
    return _make


@pytest.fixture
def make_paragraph_docx(tmp_path):
    """Factory fixture to create test .docx files with paragraph-based questions."""
    def _make(questions: list[str], filename: str = "test_para.docx") -> Path:
        doc = DocxDocument()
        for q in questions:
            doc.add_paragraph(q)
            doc.add_paragraph("")  # empty answer placeholder
        path = tmp_path / filename
        doc.save(str(path))
        return path
    return _make


@pytest.fixture
def make_csv(tmp_path):
    """Factory fixture to create questionnaire CSV files."""

    def _make(rows: list[list[str]], filename: str = "test.csv") -> Path:
        path = tmp_path / filename
        with path.open("w", newline="", encoding="utf-8") as handle:
            writer = csv.writer(handle)
            writer.writerows(rows)
        return path

    return _make
