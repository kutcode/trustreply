"""Tests for the document generator module."""

from pathlib import Path
from docx import Document as DocxDocument

from app.services.parser import ExtractedItem, RunFormat
from app.services.generator import generate_filled_csv, generate_filled_docx, generate_docx_from_pdf_items
from app.utils.questions import REVIEW_REQUIRED_PLACEHOLDER


class TestGenerateFilledDocx:
    """Tests for writing answers into .docx table cells."""

    def test_fills_table_cell(self, make_docx, tmp_path):
        source = make_docx([
            ("What is your company name?", ""),
            ("Describe your security policy.", ""),
        ])
        output = tmp_path / "filled.docx"

        items = [
            ExtractedItem(
                question_text="What is your company name?",
                item_type="table_cell",
                location={"table_idx": 0, "row_idx": 0, "q_col_idx": 0, "a_col_idx": 1},
                answer_text="Acme Corporation",
            ),
        ]

        generate_filled_docx(source, output, items)
        assert output.exists()

        # Verify the answer was written
        doc = DocxDocument(str(output))
        table = doc.tables[0]
        answer_cell = table.rows[0].cells[1]
        assert "Acme Corporation" in answer_cell.text

    def test_marks_unanswered_items_for_review(self, make_docx, tmp_path):
        source = make_docx([
            ("What is your company name?", ""),
            ("Describe your security policy.", ""),
        ])
        output = tmp_path / "filled.docx"

        items = [
            ExtractedItem(
                question_text="What is your company name?",
                item_type="table_cell",
                location={"table_idx": 0, "row_idx": 0, "q_col_idx": 0, "a_col_idx": 1},
                answer_text=None,  # not matched
            ),
        ]

        generate_filled_docx(source, output, items)
        doc = DocxDocument(str(output))
        table = doc.tables[0]
        answer_cell = table.rows[0].cells[1]
        assert REVIEW_REQUIRED_PLACEHOLDER in answer_cell.text

    def test_fills_multiple_rows(self, make_docx, tmp_path):
        source = make_docx([
            ("What is your company name?", ""),
            ("Describe your security policy.", ""),
        ])
        output = tmp_path / "filled.docx"

        items = [
            ExtractedItem(
                question_text="What is your company name?",
                item_type="table_cell",
                location={"table_idx": 0, "row_idx": 0, "q_col_idx": 0, "a_col_idx": 1},
                answer_text="Acme Corp",
            ),
            ExtractedItem(
                question_text="Describe your security policy.",
                item_type="table_cell",
                location={"table_idx": 0, "row_idx": 1, "q_col_idx": 0, "a_col_idx": 1},
                answer_text="ISO 27001 compliant",
            ),
        ]

        generate_filled_docx(source, output, items)
        doc = DocxDocument(str(output))
        table = doc.tables[0]
        assert "Acme Corp" in table.rows[0].cells[1].text
        assert "ISO 27001 compliant" in table.rows[1].cells[1].text

    def test_fills_configured_answer_column(self, tmp_path):
        doc = DocxDocument()
        table = doc.add_table(rows=1, cols=3)
        table.rows[0].cells[0].text = "ID-1"
        table.rows[0].cells[1].text = "What is your company name?"
        table.rows[0].cells[2].text = ""

        source = tmp_path / "three_col.docx"
        output = tmp_path / "three_col_filled.docx"
        doc.save(str(source))

        items = [
            ExtractedItem(
                question_text="What is your company name?",
                item_type="table_cell",
                location={
                    "table_idx": 0,
                    "row_idx": 0,
                    "q_col_idx": 1,
                    "a_col_idx": 2,
                    "question_col_idx": 1,
                    "answer_col_idx": 2,
                },
                answer_text="Acme Corporation",
            ),
        ]

        generate_filled_docx(source, output, items)
        filled = DocxDocument(str(output))
        assert "Acme Corporation" in filled.tables[0].rows[0].cells[2].text

    def test_fills_row_block_answer_below_merged_question(self, tmp_path):
        doc = DocxDocument()
        table = doc.add_table(rows=2, cols=3)
        table.rows[0].cells[0].merge(table.rows[0].cells[2]).text = "Describe your security policy."
        table.rows[1].cells[0].merge(table.rows[1].cells[2]).text = ""

        source = tmp_path / "row_block.docx"
        output = tmp_path / "row_block_filled.docx"
        doc.save(str(source))

        items = [
            ExtractedItem(
                question_text="Describe your security policy.",
                item_type="table_cell",
                location={
                    "table_idx": 0,
                    "row_idx": 1,
                    "q_row_idx": 0,
                    "answer_row_idx": 1,
                    "q_col_idx": 0,
                    "a_col_idx": 1,
                    "question_col_idx": 0,
                    "answer_col_idx": 1,
                    "layout_hint": "row_block",
                },
                answer_text="We follow ISO 27001 standards.",
            ),
        ]

        generate_filled_docx(source, output, items)
        filled = DocxDocument(str(output))
        assert "Describe your security policy." in filled.tables[0].rows[0].cells[0].text
        assert "We follow ISO 27001 standards." in filled.tables[0].rows[1].cells[0].text


class TestGenerateFilledDocxParagraphs:
    """Tests for filling paragraph-based questions."""

    def test_fills_empty_paragraph(self, make_paragraph_docx, tmp_path):
        source = make_paragraph_docx([
            "1. What is your company name?",
        ])
        output = tmp_path / "filled_para.docx"

        items = [
            ExtractedItem(
                question_text="1. What is your company name?",
                item_type="paragraph",
                location={"para_idx": 0},
                answer_text="Acme Corporation",
            ),
        ]

        generate_filled_docx(source, output, items)
        doc = DocxDocument(str(output))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Acme Corporation" in all_text

    def test_marks_unanswered_paragraph_for_review(self, make_paragraph_docx, tmp_path):
        source = make_paragraph_docx([
            "1. What is your company name?",
        ])
        output = tmp_path / "filled_para_review.docx"

        items = [
            ExtractedItem(
                question_text="1. What is your company name?",
                item_type="paragraph",
                location={"para_idx": 0},
                answer_text=None,
            ),
        ]

        generate_filled_docx(source, output, items)
        doc = DocxDocument(str(output))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert REVIEW_REQUIRED_PLACEHOLDER in all_text


class TestGenerateFilledCsv:
    """Tests for writing answers into CSV questionnaires."""

    def test_fills_csv_answer_cells(self, make_csv, tmp_path):
        source = make_csv([
            ["Question", "Answer"],
            ["What is your company name?", ""],
            ["Describe your security policy.", ""],
        ])
        output = tmp_path / "filled.csv"

        items = [
            ExtractedItem(
                question_text="What is your company name?",
                item_type="csv_row",
                location={"row_idx": 1, "q_col_idx": 0, "a_col_idx": 1},
                answer_text="Acme Corporation",
            ),
        ]

        generate_filled_csv(source, output, items)
        lines = output.read_text(encoding="utf-8").splitlines()
        assert "Acme Corporation" in lines[1]

    def test_marks_unanswered_csv_rows_for_review(self, make_csv, tmp_path):
        source = make_csv([
            ["Question", "Answer"],
            ["Describe your security policy.", ""],
        ], filename="review.csv")
        output = tmp_path / "review_filled.csv"

        items = [
            ExtractedItem(
                question_text="Describe your security policy.",
                item_type="csv_row",
                location={"row_idx": 1, "q_col_idx": 0, "a_col_idx": 1},
                answer_text=None,
            ),
        ]

        generate_filled_csv(source, output, items)
        assert REVIEW_REQUIRED_PLACEHOLDER in output.read_text(encoding="utf-8")


class TestGenerateDocxFromPdfItems:
    """Tests for generating a new .docx from PDF-extracted items."""

    def test_creates_docx_with_answers(self, tmp_path):
        output = tmp_path / "from_pdf.docx"

        items = [
            ExtractedItem(
                question_text="What is your company name?",
                item_type="pdf_text",
                location={"page_idx": 0, "line_idx": 0},
                answer_text="Acme Corporation",
            ),
            ExtractedItem(
                question_text="Describe your security policy.",
                item_type="pdf_text",
                location={"page_idx": 0, "line_idx": 1},
                answer_text=None,  # flagged
            ),
        ]

        generate_docx_from_pdf_items(output, items)
        assert output.exists()

        doc = DocxDocument(str(output))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Acme Corporation" in all_text
        assert REVIEW_REQUIRED_PLACEHOLDER in all_text
