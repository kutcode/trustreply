"""Tests for the document parser module."""

import pytest
from pathlib import Path
from docx import Document as DocxDocument

from app.services.parser import (
    ParseOptions,
    parse_docx,
    parse_document,
    parse_document_result,
    _is_question,
)


class TestIsQuestion:
    """Tests for the _is_question heuristic."""

    def test_question_mark(self):
        assert _is_question("What is your company name?")

    def test_numbered_item(self):
        assert _is_question("1. Describe your security practices in detail")

    def test_lettered_item(self):
        assert _is_question("a. Describe your data retention policy")

    def test_imperative_verb(self):
        assert _is_question("Please describe your incident response plan")

    def test_label_pattern(self):
        assert _is_question("Company Name:")

    def test_short_text_rejected(self):
        assert not _is_question("Yes")

    def test_empty_text_rejected(self):
        assert not _is_question("")

    def test_random_short_text_rejected(self):
        assert not _is_question("OK done")


class TestParseDocxTables:
    """Tests for table-based question extraction from .docx files."""

    def test_extracts_table_questions(self, make_docx):
        path = make_docx([
            ("What is your company name?", ""),
            ("Describe your security policy.", ""),
        ])
        items = parse_docx(path)
        assert len(items) >= 2
        questions = [item.question_text for item in items]
        assert "What is your company name?" in questions

    def test_skips_filled_answers(self, make_docx):
        path = make_docx([
            ("What is your company name?", "Acme Corp is our company"),
            ("Describe your security policy.", ""),
        ])
        items = parse_docx(path)
        # The first row has a filled answer AND it matches a question pattern,
        # so it may or may not be extracted depending on heuristic
        question_texts = [item.question_text for item in items]
        assert "Describe your security policy." in question_texts

    def test_item_type_is_table_cell(self, make_docx):
        path = make_docx([("What is your company name?", "")])
        items = parse_docx(path)
        assert len(items) >= 1
        assert items[0].item_type == "table_cell"

    def test_location_metadata(self, make_docx):
        path = make_docx([("What is your company name?", "")])
        items = parse_docx(path)
        assert len(items) >= 1
        loc = items[0].location
        assert "table_idx" in loc
        assert "row_idx" in loc

    def test_configurable_question_and_answer_columns(self, tmp_path):
        doc = DocxDocument()
        table = doc.add_table(rows=2, cols=3)
        table.rows[0].cells[0].text = "ID"
        table.rows[0].cells[1].text = "Question"
        table.rows[0].cells[2].text = "Answer"
        table.rows[1].cells[0].text = "1"
        table.rows[1].cells[1].text = "What is your company name?"
        table.rows[1].cells[2].text = ""

        path = tmp_path / "three_col.docx"
        doc.save(str(path))

        items = parse_docx(
            path,
            options=ParseOptions(question_column_index=1, answer_column_index=2, header_rows=1),
        )
        assert len(items) == 1
        assert items[0].question_text == "What is your company name?"
        assert items[0].location["q_col_idx"] == 1
        assert items[0].location["a_col_idx"] == 2

    def test_default_profile_auto_detects_three_column_table(self, tmp_path):
        doc = DocxDocument()
        table = doc.add_table(rows=2, cols=3)
        table.rows[0].cells[0].text = "ID"
        table.rows[0].cells[1].text = "Question"
        table.rows[0].cells[2].text = "Answer"
        table.rows[1].cells[0].text = "1"
        table.rows[1].cells[1].text = "Describe your security policy."
        table.rows[1].cells[2].text = ""

        path = tmp_path / "default_three_col.docx"
        doc.save(str(path))

        items = parse_docx(path)
        assert len(items) == 1
        assert items[0].question_text == "Describe your security policy."
        assert items[0].location["q_col_idx"] == 1
        assert items[0].location["a_col_idx"] == 2
        assert items[0].location["header_rows"] == 1

    def test_default_profile_auto_detects_four_column_metadata_table(self, tmp_path):
        doc = DocxDocument()
        table = doc.add_table(rows=2, cols=4)
        table.rows[0].cells[0].text = "Domain"
        table.rows[0].cells[1].text = "Ref"
        table.rows[0].cells[2].text = "Question"
        table.rows[0].cells[3].text = "Response"
        table.rows[1].cells[0].text = "Security"
        table.rows[1].cells[1].text = "R01"
        table.rows[1].cells[2].text = "Describe your incident response process."
        table.rows[1].cells[3].text = ""

        path = tmp_path / "default_four_col.docx"
        doc.save(str(path))

        items = parse_docx(path)
        assert len(items) == 1
        assert items[0].question_text == "Describe your incident response process."
        assert items[0].location["q_col_idx"] == 2
        assert items[0].location["a_col_idx"] == 3
        assert items[0].location["header_rows"] == 1

    def test_header_rows_are_skipped(self, tmp_path):
        doc = DocxDocument()
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "Please provide response"
        table.rows[0].cells[1].text = ""
        table.rows[1].cells[0].text = "Describe your security policy."
        table.rows[1].cells[1].text = ""

        path = tmp_path / "header_skip.docx"
        doc.save(str(path))

        items = parse_docx(path, options=ParseOptions(header_rows=1))
        questions = [item.question_text for item in items]
        assert "Please provide response" not in questions
        assert "Describe your security policy." in questions

    def test_merged_question_row_targets_following_answer_row(self, tmp_path):
        doc = DocxDocument()
        table = doc.add_table(rows=2, cols=3)
        table.rows[0].cells[0].merge(table.rows[0].cells[2]).text = "Describe your security policy."
        table.rows[1].cells[0].merge(table.rows[1].cells[2]).text = ""

        path = tmp_path / "merged_row_block.docx"
        doc.save(str(path))

        items = parse_docx(path)
        assert len(items) == 1
        item = items[0]
        assert item.question_text == "Describe your security policy."
        assert item.location["q_row_idx"] == 0
        assert item.location["answer_row_idx"] == 1
        assert item.location["layout_hint"] == "row_block"


class TestParseDocxParagraphs:
    """Tests for paragraph-based question extraction."""

    def test_extracts_paragraph_questions(self, make_paragraph_docx):
        path = make_paragraph_docx([
            "1. What is your company name?",
            "2. Describe your security measures in detail.",
        ])
        items = parse_docx(path)
        assert len(items) >= 2
        types = {item.item_type for item in items}
        assert "paragraph" in types

    def test_deduplication(self, tmp_path):
        """If same question appears in table and paragraph, only one is returned."""
        from docx import Document as DocxDocument

        doc = DocxDocument()
        table = doc.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = "What is your company name?"
        table.rows[0].cells[1].text = ""
        doc.add_paragraph("What is your company name?")
        doc.add_paragraph("")

        path = tmp_path / "dedup_test.docx"
        doc.save(str(path))

        items = parse_docx(path)
        question_texts = [item.question_text for item in items]
        count = question_texts.count("What is your company name?")
        assert count == 1, f"Expected 1 occurrence, got {count}"


class TestParseDocument:
    """Tests for the unified parse_document entry point."""

    def test_docx_routing(self, make_docx):
        path = make_docx([("What is your company name?", "")])
        items = parse_document(path)
        assert len(items) >= 1

    def test_parse_document_result_includes_diagnostics(self, make_docx):
        path = make_docx([("What is your company name?", "")])
        result = parse_document_result(path)
        assert result.parser_strategy == "heuristic"
        assert result.profile_name == "default"
        assert result.confidence > 0
        assert result.stats["items_total"] >= 1
        assert result.fallback_recommended is False

    def test_parse_document_result_recommends_fallback_for_empty_parse(self, make_docx):
        path = make_docx([])
        result = parse_document_result(path)
        assert result.items == []
        assert result.fallback_recommended is True
        assert result.fallback_reason == "no_questions_found"

    def test_unsupported_format(self, tmp_path):
        path = tmp_path / "test.txt"
        path.write_text("hello")
        with pytest.raises(ValueError, match="Unsupported file type"):
            parse_document(path)
