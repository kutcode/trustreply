#!/usr/bin/env python3
"""Generate 50 test questionnaire files (.docx and .pdf) for Trust Reply testing.

Creates files that exercise:
  - KB-matched questions (should auto-fill)
  - Unknown questions (should get flagged)
  - Duplicate / near-duplicate questions across files (tests dedup)
  - Various document layouts (2-col tables, 3-col tables, paragraph-based)
"""

import os
import random
import datetime
import json
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF

OUTPUT_DIR = Path(__file__).parent / "sample-questionnaires"
OUTPUT_DIR.mkdir(exist_ok=True)

# ─── KNOWLEDGE BASE SEED DATA ────────────────────────────────────
# These are the Q&A pairs that should be imported into the KB first.
# The test files will reference these questions (with variations).

KB_ENTRIES = [
    # ── Company Information ──
    ("Company Information", "What is the full legal name of your company?",
     "Trust Reply Technologies Inc., a Delaware C-Corporation incorporated in 2021."),
    ("Company Information", "What is your company's primary business address?",
     "1200 Innovation Drive, Suite 400, San Francisco, CA 94107, United States."),
    ("Company Information", "How many employees does your organization have?",
     "We currently employ 247 full-time employees across 3 offices globally."),
    ("Company Information", "What is your company's annual revenue?",
     "Our annual recurring revenue (ARR) for FY2025 was $18.4 million."),
    ("Company Information", "Who is the CEO of the company?",
     "Sarah Chen serves as our Chief Executive Officer since the company's founding in 2021."),
    ("Company Information", "What industry does your company operate in?",
     "Enterprise SaaS, specifically AI-powered document processing and compliance automation."),
    ("Company Information", "What year was the company founded?",
     "The company was founded in 2021."),
    ("Company Information", "What countries does your company operate in?",
     "We operate in the United States, United Kingdom, and Germany."),

    # ── Security & Compliance ──
    ("Security & Compliance", "Does your organization have a formal information security policy?",
     "Yes. We maintain a comprehensive Information Security Policy reviewed annually by our CISO and approved by the Board. Last review date: January 2025."),
    ("Security & Compliance", "Are you SOC 2 Type II certified?",
     "Yes. We achieved SOC 2 Type II certification in March 2024, covering Security, Availability, and Confidentiality trust service criteria. Our latest audit report is available under NDA."),
    ("Security & Compliance", "Do you comply with GDPR?",
     "Yes. We are fully GDPR compliant. We have appointed a Data Protection Officer (DPO), maintain Records of Processing Activities (ROPA), and conduct annual Data Protection Impact Assessments (DPIAs)."),
    ("Security & Compliance", "How do you handle data encryption?",
     "All data is encrypted at rest using AES-256 and in transit using TLS 1.3. Database-level encryption is enabled via AWS RDS encryption. Encryption keys are managed through AWS KMS with automatic rotation every 365 days."),
    ("Security & Compliance", "What is your incident response plan?",
     "We maintain a documented Incident Response Plan (IRP) aligned with NIST SP 800-61. It includes identification, containment, eradication, recovery, and lessons learned phases. Our target response time for critical incidents is under 1 hour."),
    ("Security & Compliance", "Do you perform regular penetration testing?",
     "Yes. We engage an independent third-party firm to conduct annual penetration tests. Our last pentest was completed in November 2024 with no critical findings. Remediation of all medium-severity findings was completed within 30 days."),
    ("Security & Compliance", "What access control mechanisms do you use?",
     "We implement Role-Based Access Control (RBAC) with the principle of least privilege. All access requires MFA via hardware security keys or TOTP. Privileged access is managed through a PAM solution with just-in-time provisioning."),
    ("Security & Compliance", "Do you have a Business Continuity Plan?",
     "Yes. Our Business Continuity Plan (BCP) is reviewed and tested annually. It covers disaster recovery, data backup procedures, and communication protocols. Our RTO is 4 hours and RPO is 1 hour for critical systems."),
    ("Security & Compliance", "How do you manage third-party vendor risk?",
     "We maintain a formal Vendor Risk Management program. All vendors undergo security assessments before onboarding and are re-evaluated annually. Critical vendors are subject to enhanced due diligence including SOC 2 report review."),
    ("Security & Compliance", "Is your infrastructure hosted in the cloud?",
     "Yes. Our infrastructure is hosted on AWS (Amazon Web Services) in the us-east-1 and eu-west-1 regions. We use a multi-AZ deployment for high availability."),

    # ── Data Privacy ──
    ("Data Privacy", "How do you handle personally identifiable information (PII)?",
     "PII is classified as Confidential under our data classification policy. It is encrypted at rest and in transit, access is logged and audited, and retention follows our data retention schedule (maximum 3 years unless legally required)."),
    ("Data Privacy", "Do you have a data retention policy?",
     "Yes. Our Data Retention Policy specifies retention periods by data category. Customer data is retained for the duration of the contract plus 90 days. Logs are retained for 1 year. PII is purged upon request or contract termination."),
    ("Data Privacy", "Can customers request deletion of their data?",
     "Yes. We support data subject access requests (DSARs) including the right to erasure. Deletion requests are processed within 30 days and confirmed in writing. Backups are purged within 90 days."),
    ("Data Privacy", "Where is customer data stored?",
     "Customer data is stored in AWS data centers. US customers' data resides in us-east-1 (N. Virginia). EU customers' data resides in eu-west-1 (Ireland). No data is transferred across regions without explicit consent."),
    ("Data Privacy", "Do you share customer data with third parties?",
     "We do not sell or share customer data with third parties for marketing purposes. Data may be shared with sub-processors (listed in our DPA) solely for service delivery. All sub-processors are contractually bound to equivalent data protection standards."),

    # ── Technical Architecture ──
    ("Technical Architecture", "What technology stack do you use?",
     "Our backend runs on Python (FastAPI) with PostgreSQL. The frontend is built with Next.js (React). We use Redis for caching, Celery for async task processing, and Docker/Kubernetes for orchestration."),
    ("Technical Architecture", "How do you ensure high availability?",
     "We deploy across multiple AWS Availability Zones with auto-scaling groups. Our architecture includes load balancers, health checks, and automatic failover. Historical uptime is 99.95% over the past 12 months."),
    ("Technical Architecture", "What is your disaster recovery strategy?",
     "We maintain automated daily backups with cross-region replication. Database snapshots are taken every 6 hours. Full DR failover can be executed within our 4-hour RTO target. DR drills are conducted quarterly."),
    ("Technical Architecture", "Do you use a microservices architecture?",
     "Yes. Our platform uses a microservices architecture with 12 core services communicating via gRPC and async message queues (AWS SQS). Each service is independently deployable and horizontally scalable."),
    ("Technical Architecture", "How do you handle API security?",
     "All API endpoints require authentication via OAuth 2.0 / JWT tokens. Rate limiting is enforced at 1000 requests/minute per client. API keys are rotated quarterly. We use API gateways with WAF rules for additional protection."),

    # ── HR & Operations ──
    ("HR & Operations", "Do you conduct background checks on employees?",
     "Yes. All employees undergo comprehensive background checks prior to their start date, including criminal history, education verification, and employment history verification through a certified screening provider."),
    ("HR & Operations", "What security awareness training do you provide?",
     "All employees complete mandatory security awareness training during onboarding and annually thereafter. Training covers phishing, social engineering, data handling, and incident reporting. Completion rate is tracked and reported to management."),
    ("HR & Operations", "Do you have a code of conduct?",
     "Yes. All employees sign our Code of Conduct upon hiring. It covers ethical behavior, conflicts of interest, data protection responsibilities, and acceptable use of company resources. Violations are subject to disciplinary action."),
    ("HR & Operations", "What is your employee offboarding process?",
     "Our offboarding process includes immediate revocation of all system access, return of company equipment, exit interview, and NDA reminder. Access revocation is completed within 4 hours of termination notification."),

    # ── Financial & Insurance ──
    ("Financial & Insurance", "Do you carry cyber liability insurance?",
     "Yes. We maintain a cyber liability insurance policy with a $10 million aggregate limit, covering data breaches, business interruption, and regulatory fines. Our policy is underwritten by a top-rated carrier."),
    ("Financial & Insurance", "What is your D&O insurance coverage?",
     "We carry Directors & Officers liability insurance with a $5 million limit, providing coverage for claims arising from management decisions and corporate governance matters."),
    ("Financial & Insurance", "Are your financial statements audited?",
     "Yes. Our financial statements are audited annually by a Big Four accounting firm. The most recent audit (FY2024) resulted in an unqualified opinion."),

    # ── Product & Service ──
    ("Product & Service", "What SLA do you offer?",
     "We offer a 99.9% uptime SLA for our Enterprise tier. Service credits are issued for any month where uptime falls below the committed level: 10% credit for 99.0-99.9%, 25% credit for 95.0-99.0%, and 50% credit below 95.0%."),
    ("Product & Service", "What support channels do you offer?",
     "We provide 24/7 email and ticket-based support for all tiers. Enterprise customers additionally receive dedicated Slack channel support, a named Customer Success Manager, and phone support with a 1-hour response SLA for critical issues."),
    ("Product & Service", "Do you offer a sandbox or trial environment?",
     "Yes. We provide a fully functional 14-day free trial environment. Enterprise prospects can request an extended 30-day proof of concept with dedicated solution engineering support."),
    ("Product & Service", "What integrations do you support?",
     "We offer native integrations with Salesforce, HubSpot, Slack, Microsoft Teams, Jira, and Google Workspace. We also provide a REST API and webhooks for custom integrations. An SDK is available for Python and JavaScript."),
    ("Product & Service", "How do you handle product updates and releases?",
     "We follow a continuous deployment model with releases every 2 weeks. All updates are backward-compatible. Major releases are announced 30 days in advance. Customers can opt into a staging environment to preview changes."),
]

# ─── DUPLICATE / NEAR-DUPLICATE QUESTION VARIANTS ────────────────
# These rephrase KB questions to test semantic matching and dedup.

DUPLICATE_VARIANTS = [
    # Near-duplicates of "What is the full legal name of your company?"
    "What is your company's legal name?",
    "Please provide the full legal name of your organization.",
    "Company legal name?",
    "What is the registered name of your company?",

    # Near-duplicates of "Are you SOC 2 Type II certified?"
    "Do you have SOC 2 Type II certification?",
    "Have you achieved SOC2 Type 2 compliance?",
    "Is your organization SOC 2 certified?",
    "Please confirm your SOC 2 Type II audit status.",

    # Near-duplicates of "How do you handle data encryption?"
    "What encryption standards do you use?",
    "Describe your data encryption approach.",
    "How is data encrypted at rest and in transit?",
    "What encryption methods does your company employ?",

    # Near-duplicates of "What technology stack do you use?"
    "Describe your technology stack.",
    "What technologies power your platform?",
    "What programming languages and frameworks do you use?",

    # Near-duplicates of "Do you comply with GDPR?"
    "Is your company GDPR compliant?",
    "How do you ensure GDPR compliance?",
    "Are you compliant with EU data protection regulations?",

    # Near-duplicates of "What SLA do you offer?"
    "What is your uptime SLA?",
    "Describe your service level agreement.",
    "What availability guarantees do you provide?",

    # Near-duplicates of "Do you perform regular penetration testing?"
    "How often do you conduct penetration tests?",
    "Describe your penetration testing program.",
    "When was your last third-party pentest?",

    # Near-duplicates of "What is your incident response plan?"
    "Describe your incident response process.",
    "How do you respond to security incidents?",
    "What steps do you take when a security breach occurs?",
]

# ─── UNKNOWN QUESTIONS (should get flagged) ───────────────────────
# Questions NOT in the KB — these should trigger flagged items.

UNKNOWN_QUESTIONS = [
    "What is your policy on open-source software usage?",
    "Do you have a responsible AI/ML ethics policy?",
    "How do you handle software supply chain security?",
    "What container security scanning tools do you use?",
    "Describe your approach to zero-trust network architecture.",
    "Do you participate in any bug bounty programs?",
    "What is your policy on remote work and BYOD?",
    "How do you handle cryptographic key management?",
    "Do you have a formal change management process?",
    "What logging and monitoring tools do you use?",
    "How do you ensure compliance with CCPA?",
    "What is your approach to API versioning?",
    "Do you have an AI governance framework?",
    "How do you handle multi-tenancy data isolation?",
    "What is your patch management process?",
    "Do you perform regular vulnerability scanning?",
    "What is your data classification scheme?",
    "How do you handle secure software development lifecycle (SSDLC)?",
    "Do you have a formal risk management framework?",
    "What network segmentation controls do you use?",
    "How do you handle secrets management in your CI/CD pipeline?",
    "Do you support customer-managed encryption keys (CMEK)?",
    "What is your approach to identity federation?",
    "Do you have ISO 27001 certification?",
    "What database backup verification procedures do you follow?",
    "How do you handle denial of service (DDoS) attacks?",
    "What is your mobile device management (MDM) policy?",
    "Do you have a formal asset management program?",
    "How do you ensure code review quality?",
    "What is your approach to data loss prevention (DLP)?",
]

# ─── HELPERS ──────────────────────────────────────────────────────

def pick_kb_questions(n, with_answers=False):
    """Pick n random KB questions. Returns list of (question, answer?) tuples."""
    items = random.sample(KB_ENTRIES, min(n, len(KB_ENTRIES)))
    if with_answers:
        return [(q, a) for _, q, a in items]
    return [q for _, q, _ in items]


def pick_duplicates(n):
    """Pick n duplicate variant questions."""
    return random.sample(DUPLICATE_VARIANTS, min(n, len(DUPLICATE_VARIANTS)))


def pick_unknowns(n):
    """Pick n unknown questions."""
    return random.sample(UNKNOWN_QUESTIONS, min(n, len(UNKNOWN_QUESTIONS)))


def build_question_set(n_kb, n_dup, n_unknown):
    """Build a mixed question list, shuffled."""
    questions = pick_kb_questions(n_kb) + pick_duplicates(n_dup) + pick_unknowns(n_unknown)
    random.shuffle(questions)
    return questions


# ─── DOCX GENERATORS ─────────────────────────────────────────────

def make_docx_two_column(filepath, title, questions):
    """Standard 2-column table: Question | Answer (blank)."""
    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph(f"Date: {datetime.date.today().isoformat()}")
    doc.add_paragraph("Please complete all questions below.")
    doc.add_paragraph("")

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr = table.rows[0].cells
    for cell, text in zip(hdr, ["Question", "Answer"]):
        cell.text = text
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)

    for q in questions:
        row = table.add_row().cells
        row[0].text = q
        row[1].text = ""

    # Set column widths
    for row in table.rows:
        row.cells[0].width = Inches(4.5)
        row.cells[1].width = Inches(3.0)

    doc.save(str(filepath))


def make_docx_three_column(filepath, title, questions):
    """3-column table: # | Question | Response."""
    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph(f"Prepared: {datetime.date.today().strftime('%B %d, %Y')}")
    doc.add_paragraph("")

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'

    hdr = table.rows[0].cells
    for cell, text in zip(hdr, ["#", "Question", "Response"]):
        cell.text = text
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True

    for i, q in enumerate(questions, 1):
        row = table.add_row().cells
        row[0].text = str(i)
        row[1].text = q
        row[2].text = ""

    for row in table.rows:
        row.cells[0].width = Inches(0.5)
        row.cells[1].width = Inches(4.5)
        row.cells[2].width = Inches(2.5)

    doc.save(str(filepath))


def make_docx_sectioned(filepath, title, category_questions):
    """Multiple tables grouped by section/category."""
    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph(f"Vendor: _________________    Date: {datetime.date.today().isoformat()}")
    doc.add_paragraph("")

    for section_name, questions in category_questions.items():
        doc.add_heading(section_name, level=2)
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'

        hdr = table.rows[0].cells
        hdr[0].text = "Question"
        hdr[1].text = "Your Response"
        for cell in hdr:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True

        for q in questions:
            row = table.add_row().cells
            row[0].text = q
            row[1].text = ""

        doc.add_paragraph("")

    doc.save(str(filepath))


def make_docx_paragraph_style(filepath, title, questions):
    """Paragraph-based questionnaire (numbered questions, no table)."""
    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph("Please provide detailed responses to each question below.")
    doc.add_paragraph("")

    for i, q in enumerate(questions, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. {q}")
        run.bold = True
        run.font.size = Pt(11)
        # Add blank response area
        doc.add_paragraph("Response: _______________________________________________")
        doc.add_paragraph("")

    doc.save(str(filepath))


def make_docx_row_block(filepath, title, questions):
    """Row-block style: question row (merged) + blank answer row."""
    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph("")

    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'

    for q in questions:
        # Question row (merged)
        row = table.add_row()
        row.cells[0].merge(row.cells[1])
        row.cells[0].text = q
        for p in row.cells[0].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x99)

        # Answer row
        ans_row = table.add_row()
        ans_row.cells[0].merge(ans_row.cells[1])
        ans_row.cells[0].text = ""
        # Set min height for answer area
        from docx.oxml.ns import qn
        tr = ans_row._tr
        trPr = tr.get_or_add_trPr()
        trHeight = trPr.makeelement(qn('w:trHeight'), {})
        trHeight.set(qn('w:val'), '720')  # ~0.5 inch
        trHeight.set(qn('w:hRule'), 'atLeast')
        trPr.append(trHeight)

    doc.save(str(filepath))


# ─── PDF GENERATORS ──────────────────────────────────────────────

def make_pdf_table(filepath, title, questions):
    """PDF with a 2-column table layout."""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font('Helvetica', 'B', 16)
    pdf.cell(0, 12, title, new_x="LMARGIN", new_y="NEXT", align='C')
    pdf.set_font('Helvetica', '', 10)
    pdf.cell(0, 8, f"Date: {datetime.date.today().isoformat()}", new_x="LMARGIN", new_y="NEXT", align='C')
    pdf.ln(8)

    # Table header
    pdf.set_font('Helvetica', 'B', 10)
    pdf.set_fill_color(230, 230, 240)
    pdf.cell(120, 8, 'Question', border=1, fill=True)
    pdf.cell(70, 8, 'Answer', border=1, fill=True, new_x="LMARGIN", new_y="NEXT")

    pdf.set_font('Helvetica', '', 9)
    for q in questions:
        # Calculate needed height
        lines = max(1, len(q) // 50 + 1)
        h = max(8, lines * 6)
        x = pdf.get_x()
        y = pdf.get_y()
        if y + h > 270:
            pdf.add_page()
            y = pdf.get_y()
        pdf.multi_cell(120, 6, q, border=1)
        end_y = pdf.get_y()
        pdf.set_xy(x + 120, y)
        pdf.cell(70, end_y - y, '', border=1, new_x="LMARGIN", new_y="NEXT")
        pdf.set_y(end_y)

    pdf.output(str(filepath))


def make_pdf_numbered(filepath, title, questions):
    """PDF with numbered questions (paragraph style)."""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font('Helvetica', 'B', 16)
    pdf.cell(0, 12, title, new_x="LMARGIN", new_y="NEXT", align='C')
    pdf.ln(8)

    pdf.set_font('Helvetica', '', 10)
    for i, q in enumerate(questions, 1):
        if pdf.get_y() > 260:
            pdf.add_page()
        pdf.set_font('Helvetica', 'B', 10)
        pdf.multi_cell(0, 6, f"{i}. {q}")
        pdf.set_font('Helvetica', '', 10)
        pdf.cell(0, 8, "Answer: _________________________________________________", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(4)

    pdf.output(str(filepath))


# ─── FILE GENERATION PLAN ────────────────────────────────────────

VENDOR_NAMES = [
    "Acme Corp", "GlobalTech Solutions", "Pinnacle Industries", "Vertex Systems",
    "Meridian Partners", "Catalyst Group", "Horizon Technologies", "Atlas Enterprises",
    "Nexus Digital", "Summit Analytics", "Forge Security", "Quantum Networks",
    "Apex Consulting", "Stellar Software", "Prism Technologies", "Vanguard Systems",
    "Eclipse Data", "Zenith Solutions", "Citadel IT", "Beacon Corp",
]

QUESTIONNAIRE_TYPES = [
    "Security Assessment Questionnaire",
    "Vendor Due Diligence Questionnaire",
    "Third-Party Risk Assessment",
    "Information Security Questionnaire",
    "Compliance Assessment Form",
    "Vendor Onboarding Questionnaire",
    "Security & Privacy Questionnaire",
    "Supplier Risk Evaluation",
    "IT Security Assessment",
    "Data Protection Questionnaire",
]

def generate_files():
    """Generate 50 test questionnaire files."""
    random.seed(42)  # Reproducible
    files = []

    # ── Batch 1: Standard 2-column DOCX (15 files) ──
    for i in range(15):
        vendor = random.choice(VENDOR_NAMES)
        qtype = random.choice(QUESTIONNAIRE_TYPES)
        title = f"{vendor} - {qtype}"
        # Mix: 60% KB, 15% duplicates, 25% unknown
        n_total = random.randint(8, 20)
        n_kb = int(n_total * 0.6)
        n_dup = int(n_total * 0.15)
        n_unk = n_total - n_kb - n_dup
        questions = build_question_set(n_kb, n_dup, n_unk)

        fname = f"{i+1:02d}_{vendor.replace(' ', '_')}_{qtype.split()[0]}.docx"
        fpath = OUTPUT_DIR / fname
        make_docx_two_column(fpath, title, questions)
        files.append((fname, "docx-2col", len(questions), n_kb, n_dup, n_unk))

    # ── Batch 2: 3-column DOCX (8 files) ──
    for i in range(8):
        vendor = random.choice(VENDOR_NAMES)
        title = f"{vendor} - Vendor Assessment (3-Column)"
        n_total = random.randint(10, 18)
        n_kb = int(n_total * 0.5)
        n_dup = int(n_total * 0.2)
        n_unk = n_total - n_kb - n_dup
        questions = build_question_set(n_kb, n_dup, n_unk)

        fname = f"{i+16:02d}_{vendor.replace(' ', '_')}_3col.docx"
        fpath = OUTPUT_DIR / fname
        make_docx_three_column(fpath, title, questions)
        files.append((fname, "docx-3col", len(questions), n_kb, n_dup, n_unk))

    # ── Batch 3: Sectioned DOCX with category headers (7 files) ──
    categories_pool = {
        "Company Information": [q for c, q, _ in KB_ENTRIES if c == "Company Information"],
        "Security & Compliance": [q for c, q, _ in KB_ENTRIES if c == "Security & Compliance"],
        "Data Privacy": [q for c, q, _ in KB_ENTRIES if c == "Data Privacy"],
        "Technical Architecture": [q for c, q, _ in KB_ENTRIES if c == "Technical Architecture"],
        "HR & Operations": [q for c, q, _ in KB_ENTRIES if c == "HR & Operations"],
    }
    for i in range(7):
        vendor = random.choice(VENDOR_NAMES)
        title = f"{vendor} - Comprehensive Security Review"
        cat_questions = {}
        total_kb = total_dup = total_unk = 0
        for cat, cat_qs in categories_pool.items():
            n_from_kb = random.randint(1, min(3, len(cat_qs)))
            selected = random.sample(cat_qs, n_from_kb)
            # Add some unknowns and duplicates per section
            extras = pick_unknowns(random.randint(0, 2)) + pick_duplicates(random.randint(0, 2))
            n_dup_here = min(len(extras), 2)
            n_unk_here = len(extras) - n_dup_here
            cat_questions[cat] = selected + extras
            random.shuffle(cat_questions[cat])
            total_kb += n_from_kb
            total_dup += n_dup_here
            total_unk += n_unk_here

        fname = f"{i+24:02d}_{vendor.replace(' ', '_')}_sectioned.docx"
        fpath = OUTPUT_DIR / fname
        make_docx_sectioned(fpath, title, cat_questions)
        n_total = sum(len(v) for v in cat_questions.values())
        files.append((fname, "docx-sectioned", n_total, total_kb, total_dup, total_unk))

    # ── Batch 4: Paragraph-style DOCX (5 files) ──
    for i in range(5):
        vendor = random.choice(VENDOR_NAMES)
        title = f"{vendor} - Quick Security Checklist"
        n_total = random.randint(6, 12)
        n_kb = int(n_total * 0.5)
        n_dup = int(n_total * 0.2)
        n_unk = n_total - n_kb - n_dup
        questions = build_question_set(n_kb, n_dup, n_unk)

        fname = f"{i+31:02d}_{vendor.replace(' ', '_')}_paragraph.docx"
        fpath = OUTPUT_DIR / fname
        make_docx_paragraph_style(fpath, title, questions)
        files.append((fname, "docx-paragraph", len(questions), n_kb, n_dup, n_unk))

    # ── Batch 5: Row-block DOCX (5 files) ──
    for i in range(5):
        vendor = random.choice(VENDOR_NAMES)
        title = f"{vendor} - Detailed Assessment"
        n_total = random.randint(8, 14)
        n_kb = int(n_total * 0.55)
        n_dup = int(n_total * 0.15)
        n_unk = n_total - n_kb - n_dup
        questions = build_question_set(n_kb, n_dup, n_unk)

        fname = f"{i+36:02d}_{vendor.replace(' ', '_')}_rowblock.docx"
        fpath = OUTPUT_DIR / fname
        make_docx_row_block(fpath, title, questions)
        files.append((fname, "docx-rowblock", len(questions), n_kb, n_dup, n_unk))

    # ── Batch 6: PDF table-style (5 files) ──
    for i in range(5):
        vendor = random.choice(VENDOR_NAMES)
        title = f"{vendor} - Security Questionnaire"
        n_total = random.randint(10, 16)
        n_kb = int(n_total * 0.6)
        n_dup = int(n_total * 0.15)
        n_unk = n_total - n_kb - n_dup
        questions = build_question_set(n_kb, n_dup, n_unk)

        fname = f"{i+41:02d}_{vendor.replace(' ', '_')}_table.pdf"
        fpath = OUTPUT_DIR / fname
        make_pdf_table(fpath, title, questions)
        files.append((fname, "pdf-table", len(questions), n_kb, n_dup, n_unk))

    # ── Batch 7: PDF numbered (5 files) ──
    for i in range(5):
        vendor = random.choice(VENDOR_NAMES)
        title = f"{vendor} - Data Protection Assessment"
        n_total = random.randint(8, 14)
        n_kb = int(n_total * 0.5)
        n_dup = int(n_total * 0.25)
        n_unk = n_total - n_kb - n_dup
        questions = build_question_set(n_kb, n_dup, n_unk)

        fname = f"{i+46:02d}_{vendor.replace(' ', '_')}_numbered.pdf"
        fpath = OUTPUT_DIR / fname
        make_pdf_numbered(fpath, title, questions)
        files.append((fname, "pdf-numbered", len(questions), n_kb, n_dup, n_unk))

    # ─── Print summary ───
    total_q = sum(f[2] for f in files)
    total_kb = sum(f[3] for f in files)
    total_dup = sum(f[4] for f in files)
    total_unk = sum(f[5] for f in files)

    print(f"\n{'='*70}")
    print(f"  Generated {len(files)} test files in: {OUTPUT_DIR}")
    print(f"{'='*70}")
    print(f"  Total questions across all files: {total_q}")
    print(f"    - KB-matchable:  {total_kb}")
    print(f"    - Duplicates:    {total_dup}")
    print(f"    - Unknown/Flag:  {total_unk}")
    print(f"{'='*70}")
    print(f"\n  Format breakdown:")
    from collections import Counter
    fmt_counts = Counter(f[1] for f in files)
    for fmt, count in sorted(fmt_counts.items()):
        ext = "docx" if fmt.startswith("docx") else "pdf"
        print(f"    {fmt:<20s}: {count} files (.{ext})")
    print()

    # ─── Also export the KB seed as JSON for import ───
    kb_seed_path = OUTPUT_DIR / "_KB_SEED_DATA.json"
    kb_data = [{"category": c, "question": q, "answer": a} for c, q, a in KB_ENTRIES]
    with open(kb_seed_path, "w") as f:
        json.dump(kb_data, f, indent=2, ensure_ascii=False)
    print(f"  KB seed data exported to: {kb_seed_path}")
    print(f"  ({len(KB_ENTRIES)} Q&A pairs across {len(set(c for c,_,_ in KB_ENTRIES))} categories)")

    # ─── Export file manifest ───
    manifest_path = OUTPUT_DIR / "_MANIFEST.json"
    manifest = []
    for fname, fmt, nq, nkb, ndup, nunk in files:
        manifest.append({
            "filename": fname,
            "format": fmt,
            "total_questions": nq,
            "kb_matchable": nkb,
            "duplicate_variants": ndup,
            "unknown_flaggable": nunk,
        })
    with open(manifest_path, "w") as f:
        json.dump(manifest, f, indent=2)
    print(f"  Manifest exported to: {manifest_path}")
    print()


if __name__ == "__main__":
    generate_files()
