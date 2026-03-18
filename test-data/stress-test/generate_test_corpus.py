#!/usr/bin/env python3
"""Generate a comprehensive stress-test corpus for TrustReply."""

import csv
import os
import random
import textwrap
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from fpdf import FPDF

OUT = Path(__file__).parent

# ── Question banks ──────────────────────────────────────────────────

# Questions that SHOULD match KB (security, privacy, general)
KB_MATCH_QUESTIONS = [
    "What is your company name?",
    "Where is your company headquartered?",
    "How many employees does your organization have?",
    "Do you have a dedicated information security team?",
    "What security certifications does your organization hold?",
    "Do you encrypt data at rest?",
    "Do you encrypt data in transit?",
    "How do you manage access control?",
    "Do you perform penetration testing?",
    "What is your incident response process?",
    "Do you have a business continuity and disaster recovery plan?",
    "How do you handle security awareness training?",
    "What is your password policy?",
    "How do you handle personal data / PII?",
    "Are you GDPR compliant?",
    "Do you have a Data Protection Officer?",
    "What is your data retention policy?",
    "Do you share data with third parties?",
    "Where is customer data stored?",
    "Do you perform background checks on employees?",
    "Do you have a written information security policy?",
    "How do you handle vulnerability management?",
    "What products or services does your company offer?",
    "What is your company website?",
    "Who is the CEO of your company?",
]

# Questions that should be FLAGGED (context-specific / not in KB)
FLAG_QUESTIONS = [
    "Company Name:",
    "Contact Person:",
    "Date of Completion:",
    "Vendor Reference Number:",
    "Primary Contact Email:",
    "Primary Contact Phone Number:",
    "Project Name or Engagement Title:",
    "Contract Start Date:",
    "Contract End Date:",
    "Name of Assessor:",
    "Department Requesting Service:",
    "Client Organization Name:",
    "Billing Address:",
    "Shipping Address:",
    "Tax Identification Number:",
    "DUNS Number:",
    "Unique Entity Identifier (UEI):",
    "Authorized Signatory Name:",
    "Authorized Signatory Title:",
    "Date of Last Audit:",
]

# Questions that are NOT in KB at all (novel topics)
NO_MATCH_QUESTIONS = [
    "Describe your approach to AI/ML model governance.",
    "What quantum-safe cryptography measures are you evaluating?",
    "How do you handle software bill of materials (SBOM)?",
    "What is your approach to zero-trust architecture?",
    "Describe your container security scanning process.",
    "How do you manage secrets rotation in CI/CD pipelines?",
    "What is your policy on open-source software usage?",
    "Describe your approach to API rate limiting and abuse prevention.",
    "How do you handle multi-tenancy isolation?",
    "What edge computing security measures do you employ?",
    "Describe your supply chain security assessment process.",
    "How do you manage shadow IT risks?",
    "What is your approach to data sovereignty requirements?",
    "Describe your chaos engineering practices.",
    "How do you handle security in serverless architectures?",
]

# Rephrased KB questions (should still match via semantic similarity)
REPHRASED_QUESTIONS = [
    "Tell us about your encryption standards for stored data.",
    "Explain your process for responding to security incidents.",
    "What access management framework do you use?",
    "Describe your employee security training program.",
    "What compliance certifications have you achieved?",
    "How is personally identifiable information protected?",
    "Do you comply with European data protection regulations?",
    "Outline your disaster recovery capabilities.",
    "What vulnerability scanning tools do you use?",
    "How long do you keep customer information?",
]

UNICODE_QUESTIONS = [
    "Décrivez votre politique de sécurité des données.",
    "データ暗号化の方法を説明してください。",
    "Beschreiben Sie Ihre Zugriffskontrollrichtlinien.",
    "¿Cómo manejan la protección de datos personales?",
    "Опишите вашу политику информационной безопасности.",
    "귀사의 보안 인증을 나열해 주십시오.",
    "Descreva sua política de retenção de dados.",
    "Come gestite la risposta agli incidenti di sicurezza?",
]


def make_docx_two_column_table(path, title, questions, with_answers=False):
    """Standard two-column Q&A table in DOCX."""
    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph(f"Vendor Security Assessment — {title}")
    doc.add_paragraph("")
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    hdr[0].text = "Question"
    hdr[1].text = "Answer"
    for q in questions:
        row = table.add_row().cells
        row[0].text = q
        row[1].text = "Sample answer here." if with_answers else ""
    doc.save(str(path))


def make_docx_three_column_table(path, title, questions):
    """Three-column table: ID | Question | Answer."""
    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph("Complete all fields. Leave blank if not applicable.")
    table = doc.add_table(rows=1, cols=3)
    hdr = table.rows[0].cells
    hdr[0].text = "Ref #"
    hdr[1].text = "Question"
    hdr[2].text = "Response"
    for i, q in enumerate(questions, 1):
        row = table.add_row().cells
        row[0].text = f"Q{i:03d}"
        row[1].text = q
        row[2].text = ""
    doc.save(str(path))


def make_docx_sections_with_paragraphs(path, title, sections):
    """DOCX with headed sections and paragraph-style questions."""
    doc = Document()
    doc.add_heading(title, level=0)
    doc.add_paragraph("Please provide detailed responses to each question below.")
    for section_name, questions in sections.items():
        doc.add_heading(section_name, level=1)
        for q in questions:
            doc.add_paragraph(q, style="List Number")
            doc.add_paragraph("")  # blank line for answer
    doc.save(str(path))


def make_docx_mixed_layout(path, title, table_qs, paragraph_qs, flag_qs):
    """Mixed: metadata fields at top, table in middle, paragraphs at bottom."""
    doc = Document()
    doc.add_heading(title, level=0)
    doc.add_paragraph("Section A: Vendor Information")
    for q in flag_qs:
        p = doc.add_paragraph()
        run = p.add_run(q + " ")
        run.bold = True
        p.add_run("____________________")
    doc.add_paragraph("")
    doc.add_heading("Section B: Security Controls", level=1)
    table = doc.add_table(rows=1, cols=2)
    hdr = table.rows[0].cells
    hdr[0].text = "Question"
    hdr[1].text = "Response"
    for q in table_qs:
        row = table.add_row().cells
        row[0].text = q
        row[1].text = ""
    doc.add_paragraph("")
    doc.add_heading("Section C: Additional Information", level=1)
    for q in paragraph_qs:
        doc.add_paragraph(q, style="List Number")
        doc.add_paragraph("")
    doc.save(str(path))


def make_docx_row_block(path, title, questions):
    """Row-block style: merged question row, then answer row below."""
    doc = Document()
    doc.add_heading(title, level=1)
    table = doc.add_table(rows=0, cols=2)
    for i, q in enumerate(questions):
        # Question row (merged)
        row = table.add_row()
        row.cells[0].merge(row.cells[1])
        row.cells[0].text = f"{i+1}. {q}"
        for paragraph in row.cells[0].paragraphs:
            for run in paragraph.runs:
                run.bold = True
        # Answer row
        ans_row = table.add_row()
        ans_row.cells[0].merge(ans_row.cells[1])
        ans_row.cells[0].text = ""
    doc.save(str(path))


def make_docx_multi_table(path, title, sections_dict):
    """Multiple tables, each with its own header."""
    doc = Document()
    doc.add_heading(title, level=0)
    for section, questions in sections_dict.items():
        doc.add_heading(section, level=1)
        table = doc.add_table(rows=1, cols=2)
        hdr = table.rows[0].cells
        hdr[0].text = "Item"
        hdr[1].text = "Answer"
        for q in questions:
            row = table.add_row().cells
            row[0].text = q
            row[1].text = ""
        doc.add_paragraph("")
    doc.save(str(path))


def make_docx_four_column(path, title, questions):
    """Four columns: #, Category, Question, Answer."""
    doc = Document()
    doc.add_heading(title, level=1)
    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text = "#"
    hdr[1].text = "Category"
    hdr[2].text = "Question"
    hdr[3].text = "Response"
    categories = ["General", "Security", "Privacy", "Compliance", "Operations"]
    for i, q in enumerate(questions, 1):
        row = table.add_row().cells
        row[0].text = str(i)
        row[1].text = random.choice(categories)
        row[2].text = q
        row[3].text = ""
    doc.save(str(path))


def make_csv_simple(path, questions, with_header=True):
    """Simple two-column CSV."""
    with open(path, "w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        if with_header:
            writer.writerow(["Question", "Answer"])
        for q in questions:
            writer.writerow([q, ""])


def make_csv_three_column(path, questions):
    """CSV with ID, Question, Answer columns."""
    with open(path, "w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerow(["ID", "Question", "Answer"])
        for i, q in enumerate(questions, 1):
            writer.writerow([f"Q{i:03d}", q, ""])


def make_csv_with_metadata(path, flag_qs, regular_qs):
    """CSV that mixes metadata fields and regular questions."""
    with open(path, "w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerow(["Field", "Value"])
        for q in flag_qs:
            writer.writerow([q, ""])
        writer.writerow(["", ""])
        writer.writerow(["Security Questions", ""])
        for q in regular_qs:
            writer.writerow([q, ""])


def make_csv_semicolon(path, questions):
    """Semicolon-delimited CSV (edge case)."""
    with open(path, "w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f, delimiter=";")
        writer.writerow(["Question", "Answer"])
        for q in questions:
            writer.writerow([q, ""])


def make_csv_with_existing_answers(path, questions, answers):
    """CSV with some answers pre-filled (for Agent Full override testing)."""
    with open(path, "w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerow(["Question", "Answer"])
        for q, a in zip(questions, answers):
            writer.writerow([q, a])


def make_pdf_simple(path, title, questions):
    """Simple PDF with questions listed."""
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 10, title, new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 10)
    pdf.cell(0, 8, "Please answer all questions below.", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(5)
    for i, q in enumerate(questions, 1):
        pdf.set_font("Helvetica", "B", 10)
        safe_q = q.encode("latin-1", "replace").decode("latin-1")
        pdf.cell(0, 7, f"{i}. {safe_q}", new_x="LMARGIN", new_y="NEXT")
        pdf.set_font("Helvetica", "", 10)
        pdf.cell(0, 7, "Answer: _______________________________________", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(3)
    pdf.output(str(path))


def make_pdf_multi_page(path, title, questions):
    """Multi-page PDF with many questions."""
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 18)
    pdf.cell(0, 12, title, new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 9)
    pdf.cell(0, 6, "Comprehensive Vendor Assessment - All Sections", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(5)
    sections = {
        "General Information": questions[:10],
        "Security Controls": questions[10:25],
        "Privacy & Compliance": questions[25:35],
        "Operations & Governance": questions[35:],
    }
    for section, qs in sections.items():
        if not qs:
            continue
        pdf.set_font("Helvetica", "B", 13)
        pdf.cell(0, 10, section, new_x="LMARGIN", new_y="NEXT")
        pdf.ln(2)
        for i, q in enumerate(qs, 1):
            pdf.set_font("Helvetica", "B", 9)
            safe_q = q.encode("latin-1", "replace").decode("latin-1")
            pdf.cell(0, 6, f"{i}. {safe_q}", new_x="LMARGIN", new_y="NEXT")
            pdf.set_font("Helvetica", "", 9)
            pdf.cell(0, 6, "Response: ", new_x="LMARGIN", new_y="NEXT")
            pdf.ln(4)
    pdf.output(str(path))


def make_docx_empty(path):
    """Empty DOCX — no questions at all."""
    doc = Document()
    doc.add_heading("Vendor Assessment", level=1)
    doc.add_paragraph("This document is intentionally left blank for testing.")
    doc.save(str(path))


def make_csv_empty(path):
    """CSV with headers only."""
    with open(path, "w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerow(["Question", "Answer"])


def make_docx_huge(path, count=200):
    """Large DOCX with many questions to test batching."""
    all_qs = []
    pool = KB_MATCH_QUESTIONS + FLAG_QUESTIONS + NO_MATCH_QUESTIONS + REPHRASED_QUESTIONS
    for i in range(count):
        q = pool[i % len(pool)]
        all_qs.append(f"{q}")
    make_docx_two_column_table(path, "Mega Questionnaire (200 Questions)", all_qs)


def make_csv_huge(path, count=150):
    """Large CSV with many questions."""
    pool = KB_MATCH_QUESTIONS + NO_MATCH_QUESTIONS + REPHRASED_QUESTIONS + FLAG_QUESTIONS
    qs = [pool[i % len(pool)] for i in range(count)]
    make_csv_simple(path, qs)


# ── Generate everything ─────────────────────────────────────────────

def main():
    print("Generating stress-test corpus...\n")

    # --- DOCX FILES ---

    # 01: All KB-matched security questions (should get high match rate)
    make_docx_two_column_table(
        OUT / "01_security_full_match.docx",
        "Vendor Security Questionnaire",
        KB_MATCH_QUESTIONS[:15],
    )
    print("  01_security_full_match.docx — 15 KB-matched security questions")

    # 02: All flaggable context-specific questions (should all be flagged)
    make_docx_two_column_table(
        OUT / "02_context_fields_all_flagged.docx",
        "Vendor Registration Form",
        FLAG_QUESTIONS,
    )
    print("  02_context_fields_all_flagged.docx — 20 context-specific fields (should flag)")

    # 03: Mix of KB match + flags + no-match
    make_docx_mixed_layout(
        OUT / "03_mixed_layout_vendor_assessment.docx",
        "Third-Party Vendor Assessment",
        table_qs=KB_MATCH_QUESTIONS[5:15],
        paragraph_qs=NO_MATCH_QUESTIONS[:5],
        flag_qs=FLAG_QUESTIONS[:6],
    )
    print("  03_mixed_layout_vendor_assessment.docx — mixed layout: flags + KB + novel")

    # 04: Three-column table format
    make_docx_three_column_table(
        OUT / "04_three_column_compliance.docx",
        "Annual Compliance Questionnaire",
        KB_MATCH_QUESTIONS[10:] + NO_MATCH_QUESTIONS[:5],
    )
    print("  04_three_column_compliance.docx — three-column table layout")

    # 05: Row-block format
    make_docx_row_block(
        OUT / "05_row_block_due_diligence.docx",
        "Due Diligence Assessment",
        KB_MATCH_QUESTIONS[:8] + FLAG_QUESTIONS[:4],
    )
    print("  05_row_block_due_diligence.docx — row-block (merged cells) layout")

    # 06: Rephrased questions (tests semantic similarity)
    make_docx_two_column_table(
        OUT / "06_rephrased_questions.docx",
        "Alternative Wording Security Review",
        REPHRASED_QUESTIONS,
    )
    print("  06_rephrased_questions.docx — 10 rephrased versions of KB questions")

    # 07: Novel questions not in KB (agent should flag or attempt)
    make_docx_two_column_table(
        OUT / "07_novel_questions_no_kb.docx",
        "Advanced Technology Risk Assessment",
        NO_MATCH_QUESTIONS,
    )
    print("  07_novel_questions_no_kb.docx — 15 questions with zero KB coverage")

    # 08: Multi-section document with multiple tables
    make_docx_multi_table(
        OUT / "08_multi_section_assessment.docx",
        "Comprehensive Vendor Risk Assessment",
        {
            "Company Overview": FLAG_QUESTIONS[:5] + KB_MATCH_QUESTIONS[:3],
            "Information Security": KB_MATCH_QUESTIONS[11:18],
            "Privacy & Data Protection": KB_MATCH_QUESTIONS[24:] if len(KB_MATCH_QUESTIONS) > 24 else REPHRASED_QUESTIONS[5:],
            "Advanced Topics": NO_MATCH_QUESTIONS[:5],
        },
    )
    print("  08_multi_section_assessment.docx — multi-table, multi-section document")

    # 09: Four-column layout
    make_docx_four_column(
        OUT / "09_four_column_categorized.docx",
        "Categorized Vendor Questionnaire",
        KB_MATCH_QUESTIONS + FLAG_QUESTIONS[:5],
    )
    print("  09_four_column_categorized.docx — four-column categorized layout")

    # 10: Sections with paragraph-style questions
    make_docx_sections_with_paragraphs(
        OUT / "10_paragraph_style_questionnaire.docx",
        "Detailed Security & Privacy Questionnaire",
        {
            "General Information": FLAG_QUESTIONS[:4] + KB_MATCH_QUESTIONS[:3],
            "Security Policies": KB_MATCH_QUESTIONS[11:17],
            "Data Privacy": [
                "How do you handle personal data / PII?",
                "Are you GDPR compliant?",
                "Do you have a Data Protection Officer?",
                "What is your data retention policy?",
            ],
            "Emerging Risks": NO_MATCH_QUESTIONS[:4],
        },
    )
    print("  10_paragraph_style_questionnaire.docx — paragraph/list-style layout")

    # 11: Pre-filled answers (for Agent Full mode override testing)
    doc = Document()
    doc.add_heading("Pre-Filled Vendor Questionnaire", level=1)
    doc.add_paragraph("Some answers have been pre-filled. Please verify and correct.")
    table = doc.add_table(rows=1, cols=2)
    hdr = table.rows[0].cells
    hdr[0].text = "Question"
    hdr[1].text = "Answer"
    prefilled = [
        ("What is your company name?", "WRONG COMPANY NAME LLC"),
        ("Do you encrypt data at rest?", "Yes we use encryption."),
        ("How many employees does your organization have?", "About 50 people"),
        ("Do you perform penetration testing?", ""),
        ("What is your incident response process?", ""),
        ("Are you GDPR compliant?", "Not sure, maybe?"),
        ("Where is customer data stored?", "In the cloud somewhere"),
        ("Do you have a business continuity and disaster recovery plan?", ""),
    ]
    for q, a in prefilled:
        row = table.add_row().cells
        row[0].text = q
        row[1].text = a
    doc.save(str(OUT / "11_prefilled_needs_override.docx"))
    print("  11_prefilled_needs_override.docx — pre-filled with wrong/vague answers (Agent Full test)")

    # 12: Huge questionnaire (200 questions — tests batching)
    make_docx_huge(OUT / "12_mega_200_questions.docx", count=200)
    print("  12_mega_200_questions.docx — 200 questions (batch chunking test)")

    # 13: Empty document
    make_docx_empty(OUT / "13_empty_document.docx")
    print("  13_empty_document.docx — empty document (zero questions edge case)")

    # --- CSV FILES ---

    # 14: Simple CSV, all KB matches
    make_csv_simple(OUT / "14_csv_security_basics.csv", KB_MATCH_QUESTIONS[:12])
    print("  14_csv_security_basics.csv — 12 KB-matched questions")

    # 15: CSV three-column
    make_csv_three_column(OUT / "15_csv_three_column_compliance.csv", KB_MATCH_QUESTIONS[10:] + REPHRASED_QUESTIONS[:5])
    print("  15_csv_three_column_compliance.csv — three-column CSV")

    # 16: CSV with metadata fields mixed in
    make_csv_with_metadata(
        OUT / "16_csv_metadata_mixed.csv",
        FLAG_QUESTIONS[:8],
        KB_MATCH_QUESTIONS[:10],
    )
    print("  16_csv_metadata_mixed.csv — metadata fields + security questions")

    # 17: CSV all novel questions
    make_csv_simple(OUT / "17_csv_novel_advanced.csv", NO_MATCH_QUESTIONS)
    print("  17_csv_novel_advanced.csv — all novel questions (no KB match)")

    # 18: CSV with pre-filled answers
    qs = KB_MATCH_QUESTIONS[:8]
    ans = [
        "We are Acme Corp", "San Francisco", "About 800", "Yes",
        "SOC 2", "Yes AES-256", "Yes TLS", "RBAC"
    ]
    make_csv_with_existing_answers(OUT / "18_csv_prefilled_answers.csv", qs, ans)
    print("  18_csv_prefilled_answers.csv — pre-filled (Agent Full override test)")

    # 19: Large CSV (150 questions)
    make_csv_huge(OUT / "19_csv_huge_150_questions.csv", count=150)
    print("  19_csv_huge_150_questions.csv — 150 questions (batch test)")

    # 20: Empty CSV
    make_csv_empty(OUT / "20_csv_empty.csv")
    print("  20_csv_empty.csv — headers only, no questions")

    # 21: Semicolon-delimited CSV
    make_csv_semicolon(OUT / "21_csv_semicolon_delimited.csv", KB_MATCH_QUESTIONS[:8])
    print("  21_csv_semicolon_delimited.csv — semicolon delimiter edge case")

    # 22: CSV no header
    make_csv_simple(OUT / "22_csv_no_header.csv", KB_MATCH_QUESTIONS[:6], with_header=False)
    print("  22_csv_no_header.csv — no header row edge case")

    # --- PDF FILES ---

    # 23: Simple PDF
    make_pdf_simple(
        OUT / "23_pdf_security_review.pdf",
        "Vendor Security Review",
        KB_MATCH_QUESTIONS[:12],
    )
    print("  23_pdf_security_review.pdf — 12 KB-matched questions")

    # 24: PDF with flags + KB mix
    make_pdf_simple(
        OUT / "24_pdf_vendor_registration.pdf",
        "Vendor Registration & Assessment",
        FLAG_QUESTIONS[:6] + KB_MATCH_QUESTIONS[:8],
    )
    print("  24_pdf_vendor_registration.pdf — mix of flag + KB questions")

    # 25: PDF with novel questions
    make_pdf_simple(
        OUT / "25_pdf_advanced_risk.pdf",
        "Advanced Risk & Technology Assessment",
        NO_MATCH_QUESTIONS,
    )
    print("  25_pdf_advanced_risk.pdf — all novel questions")

    # 26: Multi-page PDF (40+ questions)
    all_pdf_qs = (
        FLAG_QUESTIONS[:8]
        + KB_MATCH_QUESTIONS
        + NO_MATCH_QUESTIONS[:8]
        + REPHRASED_QUESTIONS
    )
    make_pdf_multi_page(
        OUT / "26_pdf_comprehensive_multi_page.pdf",
        "Comprehensive Multi-Page Assessment",
        all_pdf_qs,
    )
    print("  26_pdf_comprehensive_multi_page.pdf — 40+ questions across sections")

    # 27: PDF rephrased only
    make_pdf_simple(
        OUT / "27_pdf_rephrased_semantic.pdf",
        "Alternative Wording Assessment",
        REPHRASED_QUESTIONS,
    )
    print("  27_pdf_rephrased_semantic.pdf — rephrased questions (semantic matching test)")

    # --- SPECIAL EDGE CASES ---

    # 28: Unicode / multilingual questions
    make_csv_simple(OUT / "28_csv_unicode_multilingual.csv", UNICODE_QUESTIONS)
    print("  28_csv_unicode_multilingual.csv — 8 questions in different languages")

    # 29: Duplicate questions in one file
    dupes = KB_MATCH_QUESTIONS[:5] * 3  # same 5 questions repeated 3 times
    make_docx_two_column_table(
        OUT / "29_duplicate_questions.docx",
        "Questionnaire with Duplicate Questions",
        dupes,
    )
    print("  29_duplicate_questions.docx — 15 questions (5 repeated 3x)")

    # 30: Very long question text
    long_qs = [
        f"Please provide a detailed and comprehensive description of your organization's approach to {topic}, including any relevant policies, procedures, tools, frameworks, and third-party services that you use, as well as how these are monitored, audited, and updated over time."
        for topic in [
            "data encryption and key management",
            "incident response and breach notification",
            "identity and access management",
            "vulnerability management and patching",
            "business continuity and disaster recovery",
        ]
    ]
    make_docx_two_column_table(
        OUT / "30_very_long_questions.docx",
        "Detailed Assessment with Verbose Questions",
        long_qs,
    )
    print("  30_very_long_questions.docx — 5 extremely long questions")

    # 31: Single question file
    make_csv_simple(OUT / "31_csv_single_question.csv", ["Do you encrypt data at rest?"])
    print("  31_csv_single_question.csv — single question only")

    # 32: Questions with special characters
    special_qs = [
        'Do you encrypt data at rest? (AES-256 / AES-128)',
        'What % of employees complete security training annually?',
        'Have you achieved SOC 2 Type II & ISO 27001 certification?',
        'Data retention: how long (in months/years) do you keep PII?',
        'Do you use MFA for all access? [Yes/No/Partial]',
        'Describe your DR/BCP: RTO < 4hrs? RPO < 1hr?',
    ]
    make_docx_two_column_table(
        OUT / "32_special_characters.docx",
        "Questionnaire with Special Characters",
        special_qs,
    )
    print("  32_special_characters.docx — questions with symbols, brackets, slashes")

    # 33: Mixed format in one doc — YES/NO questions vs free-text
    yn_qs = [
        "Do you encrypt data at rest? (Yes/No)",
        "Do you perform annual penetration testing? (Yes/No)",
        "Is MFA required for remote access? (Yes/No)",
        "Do you have a formal change management process? (Yes/No)",
        "Are audit logs retained for at least 12 months? (Yes/No)",
    ]
    free_qs = [
        "Describe your incident response process in detail.",
        "Explain how you manage third-party risk.",
        "What is your approach to data classification?",
    ]
    make_docx_sections_with_paragraphs(
        OUT / "33_yes_no_plus_freetext.docx",
        "Hybrid Yes/No and Free-Text Questionnaire",
        {
            "Yes/No Controls Assessment": yn_qs,
            "Detailed Responses Required": free_qs,
        },
    )
    print("  33_yes_no_plus_freetext.docx — mix of yes/no and free-text questions")

    # 34: Questionnaire that looks like a real SIG Lite
    sig_sections = {
        "A. Enterprise Risk Management": [
            "Does your organization have a formal risk management program?",
            "How often is the risk assessment updated?",
            "Who is responsible for enterprise risk management?",
        ],
        "B. Security Policy": [
            "Do you have a written information security policy?",
            "How often is the security policy reviewed?",
            "How are policy exceptions handled?",
        ],
        "C. Organizational Security": [
            "Do you have a dedicated information security team?",
            "Do you perform background checks on employees?",
            "How do you handle security awareness training?",
        ],
        "D. Asset Management": [
            "Do you maintain an asset inventory?",
            "How are assets classified?",
            "What is your process for secure asset disposal?",
        ],
        "E. Access Control": [
            "How do you manage access control?",
            "What is your password policy?",
            "Do you use multi-factor authentication?",
        ],
        "F. Cryptography": [
            "Do you encrypt data at rest?",
            "Do you encrypt data in transit?",
            "How do you manage encryption keys?",
        ],
        "G. Incident Management": [
            "What is your incident response process?",
            "How quickly do you notify affected parties after a breach?",
            "Do you conduct post-incident reviews?",
        ],
    }
    make_docx_multi_table(
        OUT / "34_sig_lite_style.docx",
        "Standardized Information Gathering (SIG) Lite — Vendor Assessment",
        sig_sections,
    )
    print("  34_sig_lite_style.docx — SIG Lite-style multi-section (21 questions)")

    # 35: CAIQ-style cloud security
    caiq_qs = [
        "Are audit logs protected from unauthorized access, modification, and deletion?",
        "Are information system audit tools separated from production systems?",
        "Is a risk-based corrective action plan maintained?",
        "Are encryption keys stored securely and access restricted?",
        "Is customer data encrypted at rest using AES-256 or equivalent?",
        "Are data centers reviewed for physical security controls annually?",
        "Is a formal governance and risk management program established?",
        "Are security roles and responsibilities formally documented?",
        "Do you support identity federation using SAML or OIDC?",
        "Is a vulnerability management program in place?",
        "Are container images scanned for vulnerabilities before deployment?",
        "Is network traffic between tenants isolated?",
    ]
    make_csv_three_column(OUT / "35_csv_caiq_cloud_security.csv", caiq_qs)
    print("  35_csv_caiq_cloud_security.csv — CAIQ-style cloud security (12 questions)")

    print(f"\nDone! Generated 35 test files in {OUT}")
    print("\nTest coverage summary:")
    print("  DOCX: 15 files (two-col, three-col, four-col, row-block, multi-table, mixed, paragraph, edge cases)")
    print("  CSV:  12 files (simple, three-col, metadata, semicolon, no-header, pre-filled, huge, edge cases)")
    print("  PDF:   5 files (simple, mixed, novel, multi-page, rephrased)")
    print("  Edge:  empty, huge, unicode, duplicates, long questions, special chars, single question")


if __name__ == "__main__":
    main()
